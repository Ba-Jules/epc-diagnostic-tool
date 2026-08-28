"""Local workshop diagnosis engine.

Run with: python app.py
The application uses only the Python standard library and SQLite.  It is meant
to be a dependable local-first starting point, not a simulated real-time app.
"""
from __future__ import annotations

import csv
import base64
import hashlib
import hmac
import json
import math
import os
import re
import secrets
import shutil
import sqlite3
import sys
import time
import unicodedata
import uuid
import zipfile
from datetime import datetime, timezone, timedelta
from http import HTTPStatus
from http.cookies import SimpleCookie
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from io import StringIO
from io import BytesIO
from pathlib import Path
from urllib.parse import parse_qs, urlparse
import urllib.request
import urllib.error
from xml.etree import ElementTree as ET

try:  # Used only to generate the downloadable Excel template; the app stays local.
    import xlsxwriter
except ImportError:
    xlsxwriter = None
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None
    Inches = None
try:  # Used to draw the report chart images (Word and Excel both embed the same PNGs); reports degrade to text/tables without it.
    from PIL import Image as PILImage, ImageDraw, ImageFont
except ImportError:
    PILImage = None

ROOT = Path(__file__).resolve().parent
STATIC = ROOT / "static"
DATABASE = Path(os.environ["DATA_DIR"]) / "workshops.sqlite3" if os.environ.get("DATA_DIR") else ROOT / "data" / "workshops.sqlite3"
IMPORTS = {}
MATRIX_COLUMNS = ["Domaine", "Ordre domaine", "Code indicateur", "Indicateur", "Description", "Ordre indicateur", "Type réponse", "Obligatoire", "Actif"]
PARAMETERS = ["Nom questionnaire", "Description", "Version", "Type d'échelle", "Valeur minimum", "Valeur maximum", "Libellés des valeurs", "Nombre de priorités par domaine"]


def now() -> str:
    return datetime.now(timezone.utc).isoformat()


def slugify(text: str) -> str:
    text = unicodedata.normalize("NFKD", text or "").encode("ascii", "ignore").decode("ascii")
    text = re.sub(r"[^A-Za-z0-9]+", "-", text).strip("-")
    return text or "sans-titre"


def export_filename(*parts, ext: str) -> str:
    slug = "_".join(slugify(p) for p in parts if p)
    return f"{slug}_{datetime.now().strftime('%Y-%m-%d')}.{ext}"


# Questionnaire EPC/SENEVAL de reference : source de verite absolue = le document
# local "5 Diagnostic de SenEval selon EPC PAB (1).docx" (extrait mecaniquement via
# python-docx, aucune reformulation). 7 domaines x 10 indicateurs = 70. Chaque
# indicateur porte deux textes distincts du DOCX : sa Reference courte et son enonce
# complet ("Indicateurs qualitatifs ou Capacites"), stockes respectivement dans les
# colonnes indicators.code et indicators.label.
EPC_DOMAINS = [
    ('grh', 'Gestion des Ressources Humaines', [
        ('Formation au personnel', 'Nous offrons régulièrement la formation au personnel'),
        ('Priorités de notre Organisation', 'Notre formation du personnel contribue directement à l’exécution des priorités de notre Organisation'),
        ('Compétences pour notre mission', 'Notre personnel a les compétences appropriées pour exécuter notre mission'),
        ('Nombre du personnel', 'Le nombre de personnel est approprié pour exécuter notre mission'),
        ('Diversité de nos bénéficiaires', 'Notre personnel reflète la diversité de nos bénéficiaires'),
        ('Recrutement des membres', 'Le système de recrutement des membres favorise l’adhésion à l’organisation'),
        ('Evaluation du Personnel', 'L’évaluation du Personnel encourage la fixation des membres'),
        ('Résolution des conflits', 'La politique de résolution des griefs et des conflits contribue à la rétention du personnel :'),
        ('Allocation des tâches', 'L’allocation des tâches et des responsabilités participent à la rétention du personnel :'),
        ('Pratiques de supervision', 'Les pratiques de supervision améliorent la capacité de notre personnel à atteindre les objectifs de l’Organisation.'),
    ]),
    ('grf', 'Gestion des Ressources Financières', [
        ('Equilibre des recettes et des dépenses', 'Nous utilisons régulièrement les procédures établies pour maintenir nos recettes et nos dépenses équilibrées'),
        ('Allocation des fonds', 'Le processus de budgétisation nous amène à allouer des fonds d’une manière qui reflète étroitement nos priorités organisationnelles'),
        ('Prévisions financières', 'Nos prévisions financières sont exactes'),
        ('Modification des dépenses', 'Nous modifions nos dépenses sur une base régulière chaque fois que nous avons des déficits de revenus'),
        ('Evitement des perturbations', 'Notre système financier et nos procédures nous évitent des perturbations opérationnelles'),
        ('Décaissements périodiques', 'Nos procédures de gestion des liquidités conduisent à des décaissements périodiques de fonds'),
        ('Appui financier des bailleurs', 'Le niveau de l’appui financier de la part des bailleurs reste stable ou croissant'),
        ('Moins de  dépendance', 'Nous prenons des mesures concrètes pour rendre notre organisation moins dépendante de quelques sources de financement'),
        ('Ressources pour les activités', 'Le niveau de nos ressources disponibles pour les activités du projet est approprié pour accomplir notre mission'),
        ('Ressources pour l’équipement', 'Le niveau de nos ressources disponibles pour l’équipement (bureaux, fournitures) est approprié pour accomplir notre mission'),
    ]),
    ('parteq', 'Gestion de la Participation Equitable', [
        ('Evaluation des besoins', 'Les niveaux de participation dans l’évaluation des besoins sont élevés'),
        ('Conception des projets', 'Les niveaux de participation dans la conception des projets sont élevés'),
        ('Mise en œuvre des projets', 'Les niveaux de participation dans la mise en œuvre des projets sont élevés'),
        ('Suivi et évaluation des projets', 'Les niveaux de participation dans le suivi et l’évaluation des projets sont élevés'),
        ('Groupes sous-représentés et accès aux activités', 'Les groupes sous-représentés ont un accès équitable aux activités du projet'),
        ('Groupes sous-représentés et bénéfice', 'Les groupes sous-représentés tirent un bénéfice équitable des activités du projet'),
        ('Promotion de l’équité', 'Nos projets font constamment la promotion de l’équité à tous les niveaux de la conception et de la mise en œuvre des projets'),
        ('Evaluation des changements', 'Nous examinons régulièrement les besoins des participants au projet pour évaluer s’ils changent.'),
        ('Besoins changeants des participants', 'Nous modifions les projets pour refléter les besoins changeant des participants'),
        ('Dialogue pour le développement équitable', 'Nous engageons régulièrement les décideurs et les institutions pertinents dans un dialogue qui contribue au développement équitable et participatif'),
    ]),
    ('dur', 'Gestion de la Durabilité des Acquis de l’organisation', [
        ('Conception et D. environnementale', 'Au moment de la conception de nos projets, nous accordons une attention adéquate à la durabilité environnementale'),
        ('Conception et D. économique', 'A la durabilité économique'),
        ('Conception et D. institutionnelle', 'A la durabilité institutionnelle'),
        ('Mise en œuvre et D. environnementale', 'En exécutant les projets, nous accordons une attention adéquate à la durabilité environnementale'),
        ('Mise en œuvre et D. économique', 'A la durabilité économique'),
        ('Mise en œuvre et D. institutionnelle', 'A la durabilité institutionnelle'),
        ('Evaluation et D. environnementale', 'En faisant le suivi du projet et l’évaluation de l’impact, nous accordions une attention adéquate à la durabilité environnementale'),
        ('Evaluation et D. économique', 'A la durabilité économique'),
        ('Evaluation et D. institutionnelle', 'A la durabilité institutionnelle'),
        ('Appui technique et durabilité', 'La qualité de l’appui technique pour nos activités de terrain contribue à la durabilité du projet'),
    ]),
    ('partn', 'Gestion du Partenariat', [
        ('Liens avec les décideurs politiques', 'Nous établissons de nouveaux liens précieux avec les décideurs politiques pertinents'),
        ('Liens avec le secteur privé', 'Nous établissons de nouveaux liens précieux avec les représentants du secteur privé'),
        ('Partenariats avec d’autres Organisations', 'Nous nous engageons activement dans des partenariats productifs avec d’autres Organisations'),
        ('Suivi de nos partenariats', 'Nous faisons le suivi de l’efficacité de nos partenariats avec les autres Organisations'),
        ('Avantages financiers', 'A travers le partenariat nous obtenons des avantages financiers qui améliorent notre capacité pour accomplir notre mission'),
        ('Compétences techniques', 'Et aussi des compétences techniques qui améliorent notre capacité à accomplir notre mission'),
        ('Nouveaux réseaux et relations', 'Et aussi de nouveaux réseaux et des relations qui améliorent notre capacité à accomplir notre mission'),
        ('Confiance et coopération', 'Les partenariats ont des mécanismes en place pour renforcer la confiance et la coopération'),
        ('Contribution aux objectifs partagés', 'Les partenaires individuels contribuent de manière appropriée aux objectifs partagés'),
        ('Effort de coopération', 'Les partenaires individuels participent aux bénéfices de l’effort de coopération'),
    ]),
    ('apporg', 'Gestion de l’Apprentissage Organisationnel', [
        ('Evaluation des projets', 'Nous faisons régulièrement le suivi et l’évaluation de la mise en œuvre de nos projets'),
        ('Implication des structures dans les défis', 'Nous impliquons régulièrement les structures dans la satisfaction des défis organisationnels majeurs'),
        ('Interdépendance des structures', 'Nous reconnaissons l’interdépendance des différentes structures de notre Organisation lorsque nous analysons les problèmes'),
        ('Informations pour le travail', 'Les membres ont régulièrement les informations dont ils ont besoin pour faire efficacement leur travail'),
        ('Informations pour les priorités', 'Nous disposons d’informations appropriées pour répondre à nos priorités'),
        ('Travail d’équipe pour les défis', 'Nous utilisons efficacement le travail d’équipe pour répondre aux défis organisationnels'),
        ('Travail d’équipe pour les défis organisationnels', 'Les responsables utilisent efficacement la travail d’équipe pour répondre aux défis organisationnels'),
        ('Réunions et apprentissage organisationnel', 'Nos réunions de personnel contribuent directement à l’apprentissage organisationnel.'),
        ('Expression libre lors des réunions', 'Les membres se sentent généralement à l’aise pour s’exprimer lors des réunions de personnel'),
        ('Prise de risque pour les innovateurs', 'Notre Organisation est une place sûre de prise de risque pour les innovateurs'),
    ]),
    ('gouv', 'Gestion Stratégique et Gouvernance', [
        ('Rapportage pour les bailleurs', 'Notre système de rapportage pour les bailleurs de fonds montre une compréhension claire de leurs besoins et exigences'),
        ('CC et mobilisation des fonds', 'Notre Comité de coordination a contribué à l’exécution des fonctions comme la mobilisation des fonds'),
        ('CC et relations publiques', 'Comme les relations publiques'),
        ('CC et plaidoyer', 'Comme le plaidoyer'),
        ('CC et définition de politique', 'Comme la définition de politique'),
        ('Représentation des bénéficiaires', 'Notre Comité de coordination a une représentation appropriée de nos principales bénéficiaires'),
        ('Engagement et décisions prises', 'L’engagement par rapport à notre mission, à nos objectifs et à nos valeurs est systématiquement reflété dans les décisions prises par le Comité de coordination et le personnel'),
        ('Planification stratégique et extérieur', 'Nous utilisons une planification stratégique pour nous examiner nous-mêmes par rapport à notre environnement externe'),
        ('Initiatives et plans stratégiques', 'Nos initiatives sont élaborées et mises en œuvre conformément à nos plans stratégiques et opérationnels'),
        ('Suivi du progrès', 'Régulièrement nous faisons le suivi du progrès dans l’accomplissement de nos objectifs stratégiques'),
    ]),
]

GRADING = [(0, 22, 5), (23, 32, 10), (33, 39, 15), (40, 45, 20), (46, 50, 25), (51, 55, 30), (56, 59, 35), (60, 63, 40), (64, 67, 45), (68, 71, 50), (72, 74, 55), (75, 78, 60), (79, 81, 65), (82, 84, 70), (85, 87, 75), (88, 89, 80), (90, 92, 85), (93, 95, 90), (96, 98, 95), (99, 100, 100)]

# ==================================================
# ASSISTANT IA OPTIONNEL — couche multi-fournisseur
# ==================================================
# Aucune dépendance externe : requêtes HTTPS via urllib.request (stdlib).
# La clé API ne quitte jamais ce module : jamais renvoyée par une route GET,
# jamais journalisée, jamais écrite dans un export.

AI_PROVIDERS = {
    "gemini": {"label": "Google Gemini", "pricing": "GRATUIT", "family": "gemini",
        "models": [("gemini-2.0-flash", "Recommandé"), ("gemini-1.5-flash", "Alternatif")],
        "key_url": "https://aistudio.google.com/apikey"},
    "groq": {"label": "Groq", "pricing": "GRATUIT", "family": "openai", "base_url": "https://api.groq.com/openai/v1",
        "models": [("llama-3.3-70b-versatile", "Recommandé"), ("llama-3.1-8b-instant", "Rapide")],
        "key_url": "https://console.groq.com/keys"},
    "openrouter": {"label": "OpenRouter", "pricing": "GRATUIT", "family": "openai", "base_url": "https://openrouter.ai/api/v1",
        "models": [("openrouter/free", "Recommandé — gratuit"), ("meta-llama/llama-3.1-8b-instruct:free", "Alternatif gratuit")],
        "key_url": "https://openrouter.ai/keys"},
    "cerebras": {"label": "Cerebras", "pricing": "ESSAI", "family": "openai", "base_url": "https://api.cerebras.ai/v1",
        "models": [("llama3.1-8b", "Recommandé"), ("llama3.1-70b", "Alternatif")],
        "key_url": "https://cloud.cerebras.ai/"},
    "openai": {"label": "OpenAI", "pricing": "PAYANT", "family": "openai", "base_url": "https://api.openai.com/v1",
        "models": [("gpt-4o-mini", "Recommandé"), ("gpt-4o", "Alternatif")],
        "key_url": "https://platform.openai.com/api-keys"},
    "anthropic": {"label": "Anthropic Claude", "pricing": "PAYANT", "family": "anthropic",
        "models": [("claude-3-5-haiku-latest", "Recommandé"), ("claude-3-5-sonnet-latest", "Alternatif")],
        "key_url": "https://console.anthropic.com/settings/keys"},
    "deepseek": {"label": "DeepSeek", "pricing": "PAYANT", "family": "openai", "base_url": "https://api.deepseek.com/v1",
        "models": [("deepseek-chat", "Recommandé"), ("deepseek-reasoner", "Raisonnement approfondi")],
        "key_url": "https://platform.deepseek.com/api_keys"},
    "xai": {"label": "xAI Grok", "pricing": "PAYANT", "family": "openai", "base_url": "https://api.x.ai/v1",
        "models": [("grok-2-latest", "Recommandé")],
        "key_url": "https://console.x.ai/"},
}


class AIError(Exception):
    """Raised with a message safe to show the moderator (never a raw stack trace)."""


class AuthRequiredError(Exception):
    """Raised when an /api/ route needs a logged-in user and none is present."""


class PermissionDeniedError(Exception):
    """Raised when a logged-in user tries to reach a session/template they don't own."""


PUBLIC_API_EXACT = {"/api/auth/setup-status", "/api/auth/setup", "/api/auth/login", "/api/auth/logout", "/api/auth/me", "/api/participant"}


def is_public_api(path: str, method: str) -> bool:
    if path in PUBLIC_API_EXACT:
        return True
    if path.startswith("/api/sessions/") and method == "POST" and (path.endswith("/participants") or path.endswith("/responses") or path.endswith("/complete")):
        return True
    if path.startswith("/api/participants/") and method == "PUT":
        return True
    return False


def _http_json(url, headers, payload, timeout=25):
    req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"), headers=headers, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", "ignore")
        if e.code == 401 or e.code == 403:
            raise AIError("Clé API invalide ou non autorisée.")
        if e.code == 429:
            raise AIError("Quota atteint pour ce fournisseur. Réessayez plus tard.")
        if e.code == 404:
            raise AIError("Modèle indisponible chez ce fournisseur.")
        raise AIError(f"Le fournisseur IA a refusé la requête (code {e.code}).") from None
    except urllib.error.URLError:
        raise AIError("Connexion au fournisseur IA impossible.") from None
    except TimeoutError:
        raise AIError("Le fournisseur IA n'a pas répondu à temps.") from None


def _call_openai_compatible(base_url, api_key, model, system_prompt, user_prompt):
    data = _http_json(f"{base_url}/chat/completions", {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
        {"model": model, "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}], "temperature": 0.4})
    try:
        return data["choices"][0]["message"]["content"].strip()
    except (KeyError, IndexError, TypeError):
        raise AIError("Réponse du fournisseur IA illisible.") from None


def _call_gemini(api_key, model, system_prompt, user_prompt):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
    data = _http_json(url, {"Content-Type": "application/json"},
        {"systemInstruction": {"parts": [{"text": system_prompt}]}, "contents": [{"role": "user", "parts": [{"text": user_prompt}]}],
         "generationConfig": {"temperature": 0.4}})
    try:
        return data["candidates"][0]["content"]["parts"][0]["text"].strip()
    except (KeyError, IndexError, TypeError):
        if data.get("promptFeedback", {}).get("blockReason"):
            raise AIError("La demande a été bloquée par le fournisseur IA.") from None
        raise AIError("Réponse du fournisseur IA illisible.") from None


def _call_anthropic(api_key, model, system_prompt, user_prompt):
    data = _http_json("https://api.anthropic.com/v1/messages",
        {"Content-Type": "application/json", "x-api-key": api_key, "anthropic-version": "2023-06-01"},
        {"model": model, "max_tokens": 1200, "system": system_prompt, "messages": [{"role": "user", "content": user_prompt}]})
    try:
        return data["content"][0]["text"].strip()
    except (KeyError, IndexError, TypeError):
        raise AIError("Réponse du fournisseur IA illisible.") from None


def generate_ai_response(provider: str, model: str, system_prompt: str, user_prompt: str, api_key: str) -> str:
    cfg = AI_PROVIDERS.get(provider)
    if not cfg:
        raise AIError("Fournisseur IA inconnu.")
    if not api_key:
        raise AIError("Aucune clé API configurée pour ce fournisseur.")
    family = cfg["family"]
    if family == "openai":
        return _call_openai_compatible(cfg["base_url"], api_key, model, system_prompt, user_prompt)
    if family == "gemini":
        return _call_gemini(api_key, model, system_prompt, user_prompt)
    if family == "anthropic":
        return _call_anthropic(api_key, model, system_prompt, user_prompt)
    raise AIError("Fournisseur IA non implémenté.")


def get_ai_config(db):
    row = db.execute("SELECT enabled,provider,model,api_key FROM ai_config WHERE id=1").fetchone()
    if not row:
        return {"enabled": False, "provider": None, "model": None, "api_key": None}
    return {"enabled": bool(row["enabled"]), "provider": row["provider"], "model": row["model"], "api_key": row["api_key"]}


def require_ai(db):
    cfg = get_ai_config(db)
    if not cfg["enabled"]:
        raise AIError("Assistant IA désactivé.")
    if not cfg["provider"] or not cfg["api_key"]:
        raise AIError("Assistant IA non configuré (fournisseur ou clé manquants).")
    return cfg


def _n(v):
    return "—" if v is None else round(v, 1)


def _c(obj):
    return "non calculable (1 seul répondant)" if obj.get("consensusNote") == "single_respondent" else _n(obj.get("consensus"))


AI_SYSTEM_BASE = ("Tu assistes un modérateur d'atelier de diagnostic organisationnel EPC/SENEVAL. "
    "Tu interprètes des données déjà calculées ; tu ne recalcules jamais un score, tu n'inventes jamais un fait, "
    "une cause, une conséquence ou une recommandation absente des données fournies. "
    "Style professionnel, clair, factuel, sans jargon d'IA, sans formules comme « l'IA constate que ». Réponds en français.")


def ai_epc_context(db, sid):
    a = analysis(db, sid)
    g = a["global"]
    lines = [f"Atelier : {a['session']['name']}", f"Participants : {a['participantCount']} · questionnaires validés : {a['completedCount']}",
        f"Capacité globale : {_n(g['capacity'])}/100 · Consensus global : {_c(g)}/100 "
        f"(graduées : capacité {g['gradedCapacity']}, consensus {g['gradedConsensus']})", "", "Résultats par domaine :"]
    for d in a["domains"]:
        if d["capacity"] is None: continue
        lines.append(f"- {d['label']} : capacité {_n(d['capacity'])}, consensus {_c(d)} "
            f"(graduées {d['gradedCapacity']}/{d['gradedConsensus']}), {d['responses']} répondant(s)")
        for i in d["indicators"]:
            if i["capacity"] is None: continue
            lines.append(f"    · {i['label']} : capacité {_n(i['capacity'])}, consensus {_c(i)}")
    return "\n".join(lines)


def ai_priority_context(db, sid, pid):
    q = qualitative_data(db, sid)
    p = next((x for x in q["priorities"] if x["id"] == pid), None)
    if not p:
        raise AIError("Priorité introuvable.")
    an = next((x for x in q["analyses"] if x["priority_id"] == pid), None)
    entries = [e for e in q["entries"] if e["priority_id"] == pid]
    def block(kind, label):
        items = [e for e in entries if e["kind"] == kind]
        if not items: return f"{label} déjà saisi(e)s : aucun."
        return f"{label} déjà saisi(e)s :\n" + "\n".join(f"  - [{e['validation_status']}] {e['content']}" for e in items)
    lines = [f"Domaine : {p['domain_label']}", f"Indicateur : {p['indicator_code']} — {p['indicator_label']}",
        f"Description : {p['indicator_description']}", f"Constat déjà saisi : {an['problem'] if an and an['problem'] else 'aucun'}",
        block("cause", "Causes"), block("consequence", "Conséquences"), block("lever", "Leviers")]
    return p, an, entries, "\n".join(lines)


def log_ai_suggestion(db, sid, kind, target_id, provider, model):
    sug_id = str(uuid.uuid4())
    db.execute("INSERT INTO ai_suggestions VALUES (?,?,?,?,?,?,?,?,?)", (sug_id, sid, kind, target_id, provider, model, "proposed", now(), now()))
    db.commit()
    return sug_id


def _parse_json_list(text):
    t = text.strip()
    if t.startswith("```"):
        t = t.strip("`")
        if t.lower().startswith("json"): t = t[4:]
        t = t.strip()
    try:
        val = json.loads(t)
        return val if isinstance(val, list) else [val]
    except (json.JSONDecodeError, ValueError):
        m = re.search(r"\[.*\]", t, re.S)
        if m:
            try: return json.loads(m.group(0))
            except (json.JSONDecodeError, ValueError): pass
        raise AIError("Réponse du fournisseur IA illisible (format inattendu).") from None


AI_SECTION_LABELS = {
    "synthese_finale": "Synthèse finale", "resume_executif": "Résumé exécutif", "lecture_diagnostic": "Lecture du diagnostic",
    "synthese_domaines": "Synthèse par domaine", "synthese_priorites": "Synthèse des priorités",
    "synthese_recommandations": "Synthèse des recommandations", "synthese_formations": "Synthèse des besoins de formation",
    "synthese_plan": "Synthèse du plan d'action", "conclusion": "Conclusion générale proposée",
}
AI_REPORT_SECTIONS = {
    "synthese_finale": "Rédige une SYNTHÈSE FINALE structurée de l'atelier : participation, diagnostic (capacité/consensus/gradués), forces, fragilités et points de vigilance, priorités retenues, analyses (constats/causes/conséquences/leviers), recommandations retenues, besoins de formation — à partir uniquement des données fournies, sans jamais recalculer un score.",
    "resume_executif": "Rédige un RÉSUMÉ EXÉCUTIF de l'atelier : situation générale, principaux constats, points forts, points de vigilance, priorités retenues, principales orientations.",
    "lecture_diagnostic": "Rédige une LECTURE DU DIAGNOSTIC : tendances, écarts, convergences, divergences, domaines remarquables, à partir de la capacité et du consensus.",
    "synthese_domaines": "Rédige une SYNTHÈSE PAR DOMAINE : pour chaque domaine, résultat et interprétation prudente, en lien avec les analyses validées si disponibles.",
    "synthese_priorites": "Rédige une SYNTHÈSE DES PRIORITÉS : priorités retenues, constats, causes validées, leviers retenus.",
    "synthese_recommandations": "Rédige une SYNTHÈSE DES RECOMMANDATIONS retenues, regroupées en catégories cohérentes si pertinent.",
    "synthese_formations": "Rédige une SYNTHÈSE DES BESOINS DE FORMATION retenus.",
    "synthese_plan": "Rédige une SYNTHÈSE DU PLAN D'ACTION à partir des recommandations retenues.",
    "conclusion": "Propose une CONCLUSION GÉNÉRALE concise et institutionnelle. Précise qu'il s'agit d'une proposition, pas d'une décision validée.",
}


def ai_report_context(db, sid):
    r = report_data(db, sid)
    q, m = r["qualitative"], r["meta"]
    lines = [ai_epc_context(db, sid), "", f"Contexte de l'atelier : {m.get('context') or 'non renseigné'}", "", "Priorités et analyses validées :"]
    for p in q["priorities"]:
        an = next((x for x in q["analyses"] if x["priority_id"] == p["id"]), None)
        retained = [e for e in q["entries"] if e["priority_id"] == p["id"] and e["validation_status"] == "RETENU"]
        lines.append(f"- {p['domain_label']} — {p['indicator_label']} : constat="
            f"{an['problem'] if an and an['problem'] else 'non renseigné'} ; "
            f"causes={'; '.join(e['content'] for e in retained if e['kind']=='cause') or 'aucune'} ; "
            f"leviers={'; '.join(e['content'] for e in retained if e['kind']=='lever') or 'aucun'}")
    lines.append("\nRecommandations retenues :")
    kept_recs = [x for x in q["recommendations"] if x["status"] == "Retenue"]
    lines += ([f"- {x['title']} : {x['description']}" for x in kept_recs] or ["  aucune"])
    lines.append("\nThèmes de formation :")
    lines += ([f"- {t['title']} : {t['need_text'] or ''}" for t in q["trainingTopics"]] or ["  aucun"])
    return "\n".join(lines)


def connect(path: Path = DATABASE) -> sqlite3.Connection:
    path.parent.mkdir(parents=True, exist_ok=True)
    db = sqlite3.connect(path)
    db.row_factory = sqlite3.Row
    db.execute("PRAGMA foreign_keys = ON")
    return db


def init_schema(db: sqlite3.Connection) -> None:
    """Create/upgrade the schema only: cheap and safe to call on every request
    (called from Handler.db()). Never runs the heavy, potentially destructive
    reference-questionnaire migration, nor the retroactive template-forking
    pass — those only ever run once, from init_db(), at real process startup.
    """
    db.executescript("""
    CREATE TABLE IF NOT EXISTS templates (id TEXT PRIMARY KEY, name TEXT NOT NULL, version INTEGER NOT NULL, description TEXT, status TEXT NOT NULL, scale_json TEXT NOT NULL, scoring_json TEXT NOT NULL, consensus_json TEXT NOT NULL, grading_json TEXT NOT NULL, priority_json TEXT NOT NULL, created_at TEXT NOT NULL, updated_at TEXT NOT NULL, UNIQUE(name, version));
    CREATE TABLE IF NOT EXISTS domains (id TEXT PRIMARY KEY, template_id TEXT NOT NULL REFERENCES templates(id), code TEXT NOT NULL, label TEXT NOT NULL, description TEXT, display_order INTEGER NOT NULL, active INTEGER NOT NULL DEFAULT 1);
    CREATE TABLE IF NOT EXISTS indicators (id TEXT PRIMARY KEY, domain_id TEXT NOT NULL REFERENCES domains(id), code TEXT NOT NULL, label TEXT NOT NULL, description TEXT, response_type TEXT NOT NULL, required INTEGER NOT NULL DEFAULT 1, display_order INTEGER NOT NULL, active INTEGER NOT NULL DEFAULT 1, configuration_json TEXT NOT NULL DEFAULT '{}');
    CREATE TABLE IF NOT EXISTS sessions (id TEXT PRIMARY KEY, template_id TEXT NOT NULL REFERENCES templates(id), template_version INTEGER NOT NULL, name TEXT NOT NULL, organization TEXT, location TEXT, date TEXT, status TEXT NOT NULL, created_at TEXT NOT NULL, closed_at TEXT, description TEXT, expected_participants INTEGER);
    CREATE TABLE IF NOT EXISTS participants (id TEXT PRIMARY KEY, session_id TEXT NOT NULL REFERENCES sessions(id), anonymous_id TEXT NOT NULL, status TEXT NOT NULL, started_at TEXT NOT NULL, completed_at TEXT, display_name TEXT, UNIQUE(session_id, anonymous_id));
    CREATE TABLE IF NOT EXISTS responses (id TEXT PRIMARY KEY, session_id TEXT NOT NULL REFERENCES sessions(id), participant_id TEXT NOT NULL REFERENCES participants(id), indicator_id TEXT NOT NULL REFERENCES indicators(id), value_json TEXT NOT NULL, value_type TEXT NOT NULL, created_at TEXT NOT NULL, updated_at TEXT NOT NULL, UNIQUE(participant_id, indicator_id));
    CREATE TABLE IF NOT EXISTS priorities (id TEXT PRIMARY KEY, session_id TEXT NOT NULL REFERENCES sessions(id), domain_id TEXT NOT NULL REFERENCES domains(id), indicator_id TEXT NOT NULL REFERENCES indicators(id), votes INTEGER NOT NULL DEFAULT 0, selected_at TEXT NOT NULL, UNIQUE(session_id, indicator_id));
    CREATE TABLE IF NOT EXISTS analysis_notes (id TEXT PRIMARY KEY, session_id TEXT NOT NULL REFERENCES sessions(id), indicator_id TEXT REFERENCES indicators(id), kind TEXT NOT NULL, content TEXT NOT NULL, validation_status TEXT NOT NULL DEFAULT 'HYPOTHESE', created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
    CREATE TABLE IF NOT EXISTS recommendations (id TEXT PRIMARY KEY, session_id TEXT NOT NULL REFERENCES sessions(id), indicator_id TEXT REFERENCES indicators(id), title TEXT NOT NULL, description TEXT, lever TEXT, kind TEXT NOT NULL, owner TEXT, horizon TEXT, created_at TEXT NOT NULL);
    CREATE TABLE IF NOT EXISTS priority_analyses (id TEXT PRIMARY KEY, session_id TEXT NOT NULL REFERENCES sessions(id), priority_id TEXT NOT NULL REFERENCES priorities(id), problem TEXT NOT NULL DEFAULT '', created_at TEXT NOT NULL, updated_at TEXT NOT NULL, UNIQUE(session_id, priority_id));
    CREATE TABLE IF NOT EXISTS analysis_entries (id TEXT PRIMARY KEY, session_id TEXT NOT NULL REFERENCES sessions(id), priority_id TEXT NOT NULL REFERENCES priorities(id), parent_id TEXT REFERENCES analysis_entries(id), kind TEXT NOT NULL, content TEXT NOT NULL, item_type TEXT, comment TEXT, validation_status TEXT NOT NULL DEFAULT 'A_DISCUTER', created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
    CREATE TABLE IF NOT EXISTS workshop_recommendations (id TEXT PRIMARY KEY, session_id TEXT NOT NULL REFERENCES sessions(id), priority_id TEXT REFERENCES priorities(id), cause_id TEXT REFERENCES analysis_entries(id), lever_id TEXT REFERENCES analysis_entries(id), title TEXT NOT NULL, description TEXT NOT NULL, category TEXT NOT NULL DEFAULT 'Autre', priority_level TEXT NOT NULL DEFAULT 'Non définie', owner TEXT, horizon TEXT, comment TEXT, status TEXT NOT NULL DEFAULT 'Proposée', created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
    CREATE TABLE IF NOT EXISTS training_topics (id TEXT PRIMARY KEY, session_id TEXT NOT NULL REFERENCES sessions(id), priority_id TEXT REFERENCES priorities(id), recommendation_id TEXT REFERENCES workshop_recommendations(id), title TEXT NOT NULL, need_text TEXT, target_audience TEXT, priority_level TEXT NOT NULL DEFAULT 'Non définie', comment TEXT, created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
    CREATE TABLE IF NOT EXISTS session_report_meta (session_id TEXT PRIMARY KEY REFERENCES sessions(id), facilitator TEXT, audience TEXT, context TEXT, conclusion TEXT, updated_at TEXT NOT NULL);
    CREATE TABLE IF NOT EXISTS ai_config (id INTEGER PRIMARY KEY CHECK (id=1), enabled INTEGER NOT NULL DEFAULT 0, provider TEXT, model TEXT, api_key TEXT, updated_at TEXT NOT NULL);
    CREATE TABLE IF NOT EXISTS ai_suggestions (id TEXT PRIMARY KEY, session_id TEXT NOT NULL REFERENCES sessions(id), kind TEXT NOT NULL, target_id TEXT, provider TEXT NOT NULL, model TEXT NOT NULL, status TEXT NOT NULL DEFAULT 'proposed', created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
    CREATE TABLE IF NOT EXISTS report_ai_blocks (id TEXT PRIMARY KEY, session_id TEXT NOT NULL REFERENCES sessions(id), section_key TEXT NOT NULL, content TEXT NOT NULL, retained_at TEXT NOT NULL, UNIQUE(session_id, section_key));
    CREATE TABLE IF NOT EXISTS users (id TEXT PRIMARY KEY, email TEXT NOT NULL UNIQUE, password_hash TEXT NOT NULL, password_salt TEXT NOT NULL, role TEXT NOT NULL DEFAULT 'admin', display_name TEXT, created_at TEXT NOT NULL);
    CREATE TABLE IF NOT EXISTS auth_tokens (token_hash TEXT PRIMARY KEY, user_id TEXT NOT NULL REFERENCES users(id), created_at TEXT NOT NULL, expires_at TEXT NOT NULL);
    """)
    db.commit()
    existing_columns = {r["name"] for r in db.execute("PRAGMA table_info(participants)")}
    if "display_name" not in existing_columns:
        db.execute("ALTER TABLE participants ADD COLUMN display_name TEXT")
        db.commit()
    for col, decl in (("anonymous", "INTEGER"), ("participant_type", "TEXT"), ("profile", "TEXT"), ("sex", "TEXT"), ("age_range", "TEXT"), ("education_level", "TEXT"),
                      # Précision libre saisie quand education_level = "Autre" (mission :8810,
                      # demande Mouhamed BA) : "Autre" reste la seule catégorie statistique,
                      # ce texte n'est jamais utilisé pour créer de nouvelles catégories.
                      ("education_level_other", "TEXT")):
        if col not in existing_columns:
            db.execute(f"ALTER TABLE participants ADD COLUMN {col} {decl}")
    db.commit()
    session_columns = {r["name"] for r in db.execute("PRAGMA table_info(sessions)")}
    for col, decl in (("description", "TEXT"), ("expected_participants", "INTEGER"), ("owner_user_id", "TEXT")):
        if col not in session_columns:
            db.execute(f"ALTER TABLE sessions ADD COLUMN {col} {decl}")
    template_columns = {r["name"] for r in db.execute("PRAGMA table_info(templates)")}
    if "owner_user_id" not in template_columns:
        db.execute("ALTER TABLE templates ADD COLUMN owner_user_id TEXT")
    if "is_canonical" not in template_columns:
        db.execute("ALTER TABLE templates ADD COLUMN is_canonical INTEGER NOT NULL DEFAULT 0")
        db.commit()
        if db.execute("SELECT 1 FROM templates WHERE is_canonical=1").fetchone() is None:
            # One-time bootstrap for databases created before is_canonical existed: pick the
            # oldest ('EPC / SENEVAL', lowest version) row as canonical. clone_template always
            # assigns version = MAX+1, so the true, never-forked original is guaranteed to hold
            # the lowest version number for that name — the one deliberate, one-off use of name
            # matching, needed only to bootstrap the stable id-based flag itself.
            legacy = db.execute("SELECT id FROM templates WHERE name='EPC / SENEVAL' ORDER BY version ASC LIMIT 1").fetchone()
            if legacy:
                db.execute("UPDATE templates SET is_canonical=1 WHERE id=?", (legacy["id"],))
    db.commit()
    if db.execute("SELECT 1 FROM templates LIMIT 1").fetchone() is None:
        seed_epc(db)


def init_db(db: sqlite3.Connection) -> None:
    """Full startup sequence: schema + one-time/heavy migrations. Call this ONCE, from
    main(), before the HTTP server starts accepting requests — never from the per-request
    Handler.db() (use init_schema there instead)."""
    init_schema(db)
    ensure_private_templates_for_started_sessions(db)
    migrate_reference_questionnaire(db)
    migrate_ownership(db)
    ensure_epc_35_template(db)


def migrate_reference_questionnaire(db: sqlite3.Connection) -> None:
    """Keep the canonical EPC/SENEVAL reference template aligned with EPC_DOMAINS.

    Runs once at real process startup only (see init_db()), and only acts when the
    stored domains/indicators of the canonical row (identified by the stable
    is_canonical flag, never by name — a mission's private fork or a duplicate can
    share the exact same name without ever being mistaken for the reference) differ
    from EPC_DOMAINS (idempotent). A backup of the database is taken before any
    destructive change. By the time this runs, every session that had already started
    collecting has been forked onto its own private template (see
    ensure_private_templates_for_started_sessions, called first in init_db()), so this
    can only ever affect the canonical row itself — never a mission's historical data.
    """
    target_codes = [code for code, _, _ in EPC_DOMAINS]
    target_counts = {code: len(indicators) for code, _, indicators in EPC_DOMAINS}
    tpl = db.execute("SELECT id FROM templates WHERE is_canonical=1").fetchone()
    if not tpl:
        return
    tid = tpl["id"]
    domains = db.execute("SELECT id,code FROM domains WHERE template_id=? ORDER BY display_order", (tid,)).fetchall()
    up_to_date = [d["code"] for d in domains] == target_codes and all(
        db.execute("SELECT COUNT(*) FROM indicators WHERE domain_id=?", (d["id"],)).fetchone()[0] == target_counts[d["code"]]
        for d in domains
    )
    if up_to_date:
        return

    if DATABASE.exists():
        backup = DATABASE.with_name(f"{DATABASE.stem}.bak-{datetime.now().strftime('%Y%m%d-%H%M%S')}{DATABASE.suffix}")
        try:
            shutil.copy2(DATABASE, backup)
        except OSError:
            pass

    old_domain_ids = [d["id"] for d in domains]
    db.commit()
    db.execute("PRAGMA foreign_keys = OFF")
    try:
        if old_domain_ids:
            dph = ",".join("?" * len(old_domain_ids))
            indicator_ids = [r["id"] for r in db.execute(f"SELECT id FROM indicators WHERE domain_id IN ({dph})", old_domain_ids)]
            if indicator_ids:
                iph = ",".join("?" * len(indicator_ids))
                db.execute(f"DELETE FROM responses WHERE indicator_id IN ({iph})", indicator_ids)
                priority_ids = [r["id"] for r in db.execute(f"SELECT id FROM priorities WHERE indicator_id IN ({iph})", indicator_ids)]
                if priority_ids:
                    pph = ",".join("?" * len(priority_ids))
                    db.execute(f"UPDATE workshop_recommendations SET priority_id=NULL WHERE priority_id IN ({pph})", priority_ids)
                    db.execute(f"UPDATE training_topics SET priority_id=NULL WHERE priority_id IN ({pph})", priority_ids)
                    entry_ids = [r["id"] for r in db.execute(f"SELECT id FROM analysis_entries WHERE priority_id IN ({pph})", priority_ids)]
                    if entry_ids:
                        eph = ",".join("?" * len(entry_ids))
                        db.execute(f"UPDATE workshop_recommendations SET cause_id=NULL WHERE cause_id IN ({eph})", entry_ids)
                        db.execute(f"UPDATE workshop_recommendations SET lever_id=NULL WHERE lever_id IN ({eph})", entry_ids)
                        db.execute(f"DELETE FROM analysis_entries WHERE id IN ({eph})", entry_ids)
                    db.execute(f"DELETE FROM priority_analyses WHERE priority_id IN ({pph})", priority_ids)
                    db.execute(f"DELETE FROM priorities WHERE id IN ({pph})", priority_ids)
                db.execute(f"UPDATE analysis_notes SET indicator_id=NULL WHERE indicator_id IN ({iph})", indicator_ids)
                db.execute(f"UPDATE recommendations SET indicator_id=NULL WHERE indicator_id IN ({iph})", indicator_ids)
                db.execute(f"DELETE FROM indicators WHERE id IN ({iph})", indicator_ids)
            db.execute(f"DELETE FROM domains WHERE id IN ({dph})", old_domain_ids)
        for d_order, (code, label, indicators) in enumerate(EPC_DOMAINS, 1):
            did = str(uuid.uuid4())
            db.execute("INSERT INTO domains VALUES (?,?,?,?,?,?,?)", (did, tid, code, label, "", d_order, 1))
            for i_order, (reference, enonce) in enumerate(indicators, 1):
                db.execute("INSERT INTO indicators VALUES (?,?,?,?,?,?,?,?,?,?)", (str(uuid.uuid4()), did, reference, enonce, "", "numeric", 1, i_order, 1, "{}"))
        db.execute("UPDATE templates SET updated_at=? WHERE id=?", (now(), tid))
        db.commit()
    finally:
        db.execute("PRAGMA foreign_keys = ON")


def migrate_ownership(db: sqlite3.Connection) -> None:
    """Assign every ownerless session/template to the first account created.

    Runs on every startup (idempotent: a no-op once nothing is ownerless, or
    while no account exists yet). This is what lets the pilot who completes
    the first-run setup immediately keep the workshops/questionnaires that
    already existed before authentication was introduced.
    """
    user = db.execute("SELECT id FROM users ORDER BY created_at LIMIT 1").fetchone()
    if not user:
        return
    if db.execute("SELECT 1 FROM sessions WHERE owner_user_id IS NULL LIMIT 1").fetchone() is None \
            and db.execute("SELECT 1 FROM templates WHERE owner_user_id IS NULL AND name != 'EPC / SENEVAL' LIMIT 1").fetchone() is None:
        return
    if DATABASE.exists():
        backup = DATABASE.with_name(f"{DATABASE.stem}.bak-{datetime.now().strftime('%Y%m%d-%H%M%S')}{DATABASE.suffix}")
        try:
            shutil.copy2(DATABASE, backup)
        except OSError:
            pass
    db.execute("UPDATE sessions SET owner_user_id=? WHERE owner_user_id IS NULL", (user["id"],))
    # The reference EPC/SENEVAL template stays a common model (owner NULL); only
    # personal/custom questionnaires get attached to the first account.
    db.execute("UPDATE templates SET owner_user_id=? WHERE owner_user_id IS NULL AND name != 'EPC / SENEVAL'", (user["id"],))
    db.commit()


PBKDF2_ITERATIONS = 200_000
AUTH_TOKEN_TTL_DAYS = 14


def hash_password(password: str) -> tuple[str, str]:
    salt = secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), bytes.fromhex(salt), PBKDF2_ITERATIONS).hex()
    return digest, salt


def verify_password(password: str, password_hash: str, password_salt: str) -> bool:
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), bytes.fromhex(password_salt), PBKDF2_ITERATIONS).hex()
    return hmac.compare_digest(digest, password_hash)


def create_auth_token(db: sqlite3.Connection, user_id: str) -> str:
    token = secrets.token_urlsafe(32)
    token_hash = hashlib.sha256(token.encode("utf-8")).hexdigest()
    expires = (datetime.now(timezone.utc) + timedelta(days=AUTH_TOKEN_TTL_DAYS)).isoformat()
    db.execute("INSERT INTO auth_tokens VALUES (?,?,?,?)", (token_hash, user_id, now(), expires))
    db.commit()
    return token


def session_cookie_header(token: str | None = None, clear: bool = False) -> str:
    # No `Secure` flag: the VPS currently serves this app over plain HTTP (no TLS
    # termination in front of it), and a Secure cookie would simply never be sent
    # back by the browser, breaking login entirely. SameSite=Lax is the practical
    # CSRF mitigation available without adding a token scheme.
    if clear:
        return "epc_session=; Path=/; HttpOnly; SameSite=Lax; Max-Age=0"
    return f"epc_session={token}; Path=/; HttpOnly; SameSite=Lax; Max-Age={AUTH_TOKEN_TTL_DAYS*86400}"


def seed_epc(db: sqlite3.Connection) -> str:
    tid, stamp = str(uuid.uuid4()), now()
    scale = {"type": "numeric", "min": 1, "max": 5, "labels": {"1": "Totalement en désaccord", "2": "En désaccord", "3": "Neutre", "4": "D’accord", "5": "Totalement d’accord"}}
    scoring = {"capacity": "mean_divided_by_scale_max", "outputRange": [0, 100]}
    consensus = {"method": "standard_deviation", "normalization": "theoretical_range", "factor": 2}
    db.execute("INSERT INTO templates VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (tid, "EPC / SENEVAL", 1, "Configuration initiale issue du questionnaire actuel.", "active", json.dumps(scale), json.dumps(scoring), json.dumps(consensus), json.dumps(GRADING), json.dumps({"maxPerDomain": 3}), stamp, stamp, None, 1))
    for d_order, (code, label, indicators) in enumerate(EPC_DOMAINS, 1):
        did = str(uuid.uuid4())
        db.execute("INSERT INTO domains VALUES (?,?,?,?,?,?,?)", (did, tid, code, label, "", d_order, 1))
        for i_order, (reference, enonce) in enumerate(indicators, 1):
            db.execute("INSERT INTO indicators VALUES (?,?,?,?,?,?,?,?,?,?)", (str(uuid.uuid4()), did, reference, enonce, "", "numeric", 1, i_order, 1, "{}"))
    db.commit()
    return tid


# Nouveau questionnaire SenEval reduit, fourni par le pilote (mouhba@local) le 19/08/2026 :
# source de verite = le document local "6 Diagnostic de SenEval selon  EPC MB Envoye PAB.docx"
# (extrait mecaniquement via python-docx, aucune reformulation). 7 domaines x 5 indicateurs = 35.
# Ce DOCX ne fournit aucune reference courte separee par indicateur (contrairement a EPC_DOMAINS) :
# indicators.code est donc rempli avec le numero d'ordre global "1".."35", conformement a la
# consigne ("ordre 1 a 35"). Distinct et independant du referentiel canonique EPC_DOMAINS/70 :
# ensure_epc_35_template() ne touche jamais ce dernier ni ses forks.
EPC_DOMAINS_5 = [
    ('grh', '1 Gestion des ressources humaines', [
        'Le membre de l’association possède les compétences appropriées pour accomplir la mission',
        'Les pratiques de supervision renforcent les capacités du membre à atteindre les objectifs',
        "La formation offerte est directement liée aux priorités de l'organisation",
        "L'évaluation du membre favorise sa motivation et sa fidélisation",
        "L'allocation des tâches et des responsabilités favorise la fidélisation du membre",
    ]),
    ('grf', '2 Gestion des ressources financières', [
        'La stratégie de mobilisation des cotisations est efficace',
        "Les sources de financement de l’association sont diversifiées",
        "Le niveau d'appui des partenaires au développement est stable ou croissant",
        'Les ressources disponibles permettent de réaliser correctement les activités',
        'Les fonds mobilisés sont adéquatement alloués aux activités',
    ]),
    ('parteq', '3 Gestion de la participation équitable', [
        'Les parties prenantes participent à la conception des projets',
        "Les parties prenantes participent à la mise en œuvre des projets",
        'Les parties prenantes participent aux Groupes techniques de travail',
        'Les projets et activités sont suivis et évalués',
        "L'organisation entretient un dialogue régulier avec les décideurs et les institutions concernés",
    ]),
    ('dur', '4 Gestion de la durabilité des acquis', [
        'La durabilité institutionnelle est intégrée dès la conception des projets',
        'La durabilité institutionnelle est prise en compte pendant la mise en œuvre',
        "La durabilité institutionnelle est évaluée lors du suivi et de l'évaluation",
        "L'appui technique contribue effectivement à la durabilité des résultats",
        'La durabilité économique est intégrée dès la conception des projets',
    ]),
    ('partn', '5 Gestion du partenariat', [
        "L'organisation développe des partenariats productifs avec d'autres organisations",
        "Les partenariats font l'objet d'un suivi régulier",
        'Les partenariats reposent sur des mécanismes favorisant la confiance et la coopération',
        'Chaque partenaire contribue effectivement aux objectifs communs',
        'Les partenariats permettent de développer de nouveaux réseaux et relations stratégiques',
    ]),
    ('apporg', "6 Gestion de l'apprentissage organisationnel", [
        "Les projets font régulièrement l'objet d'un suivi-évaluation",
        "Les réunions favorisent l'apprentissage organisationnel",
        "L'organisation dispose des informations nécessaires pour répondre à ses priorités",
        'Les structures reconnaissent leur interdépendance dans la résolution des problèmes',
        "L'organisation encourage l'innovation et une prise de risque maîtrisée",
    ]),
    ('gouv', '7 Gestion stratégique et gouvernance', [
        'Les initiatives sont alignées sur les plans stratégiques et opérationnels',
        'Les progrès vers les objectifs stratégiques sont suivis régulièrement',
        "Les organes fonctionnent régulièrement en respectant  les objectifs de l'organisation",
        'La planification stratégique  est déclinée en plans d’action opérationnels',
        'Le Comité de coordination contribue efficacement à la mobilisation des ressources',
    ]),
]

EPC_35_TEMPLATE_NAME = "EPC / SENEVAL (7 domaines x 5 indicateurs)"


def ensure_epc_35_template(db: sqlite3.Connection) -> None:
    """Cree une fois, de maniere idempotente, le nouveau gabarit reduit EPC_DOMAINS_5 (7x35),
    fourni par le pilote le 19/08/2026 pour ses futures missions. Purement additif : un seul
    INSERT sur une ligne 'templates' entierement nouvelle (is_canonical=0), jamais un UPDATE/DELETE
    sur le referentiel canonique, sur EPC_DOMAINS, ni sur une mission existante. Reprend
    exactement la meme echelle/formule de capacite/consensus/bareme que seed_epc() : seul le nombre
    d'indicateurs par domaine change, aucune formule EPC n'est modifiee. Rattache au compte
    mouhba@local (recherche par email, jamais un id code en dur) pour qu'il apparaisse comme un
    gabarit lui appartenant, exactement comme s'il l'avait cree lui-meme depuis l'appli."""
    if db.execute("SELECT 1 FROM templates WHERE name=?", (EPC_35_TEMPLATE_NAME,)).fetchone():
        return
    owner = db.execute("SELECT id FROM users WHERE email=?", ("mouhba@local",)).fetchone()
    owner_user_id = owner["id"] if owner else None
    tid, stamp = str(uuid.uuid4()), now()
    scale = {"type": "numeric", "min": 1, "max": 5, "labels": {"1": "Totalement en désaccord", "2": "En désaccord", "3": "Neutre", "4": "D’accord", "5": "Totalement d’accord"}}
    scoring = {"capacity": "mean_divided_by_scale_max", "outputRange": [0, 100]}
    consensus = {"method": "standard_deviation", "normalization": "theoretical_range", "factor": 2}
    db.execute("INSERT INTO templates VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (tid, EPC_35_TEMPLATE_NAME, 1, "Questionnaire réduit fourni par le pilote (7 domaines x 5 indicateurs).", "active", json.dumps(scale), json.dumps(scoring), json.dumps(consensus), json.dumps(GRADING), json.dumps({"maxPerDomain": 3}), stamp, stamp, owner_user_id, 0))
    order_num = 1
    for d_order, (code, label, indicators) in enumerate(EPC_DOMAINS_5, 1):
        did = str(uuid.uuid4())
        db.execute("INSERT INTO domains VALUES (?,?,?,?,?,?,?)", (did, tid, code, label, "", d_order, 1))
        for i_order, enonce in enumerate(indicators, 1):
            db.execute("INSERT INTO indicators VALUES (?,?,?,?,?,?,?,?,?,?)", (str(uuid.uuid4()), did, str(order_num), enonce, "", "numeric", 1, i_order, 1, "{}"))
            order_num += 1
    db.commit()


def rows(db: sqlite3.Connection, sql: str, args=()):
    return [dict(r) for r in db.execute(sql, args).fetchall()]


def template_domain_indicator_counts(db: sqlite3.Connection, template_id: str):
    """Nombre de domaines/indicateurs ACTIFS d'un questionnaire, calculé depuis les
    données réelles (jamais codé en dur) — utilisé par /api/templates pour afficher
    par exemple "7 domaines x 10 indicateurs" (mission :8810, demande Mouhamed BA).
    Le canonique EPC/SENEVAL reste 7 domaines / 70 indicateurs, inchangé."""
    domain_count = db.execute("SELECT COUNT(*) FROM domains WHERE template_id=? AND active=1", (template_id,)).fetchone()[0]
    indicator_count = db.execute(
        "SELECT COUNT(*) FROM indicators i JOIN domains d ON d.id=i.domain_id WHERE d.template_id=? AND d.active=1 AND i.active=1",
        (template_id,),
    ).fetchone()[0]
    return domain_count, indicator_count


PARTICIPANT_PROFILE_FIELD_MAP = {"displayName": "display_name", "anonymous": "anonymous", "participantType": "participant_type", "profile": "profile", "sex": "sex", "ageRange": "age_range", "educationLevel": "education_level",
    # Précision libre saisie quand educationLevel = "Autre" (mission :8810) — jamais
    # utilisée comme catégorie statistique propre, seule "Autre" l'est.
    "educationLevelOther": "education_level_other"}


def indicator_code_conflicts(db: sqlite3.Connection, domain_id: str, code: str, exclude_id: str | None = None) -> bool:
    """Une "Référence" (indicators.code) dupliquée DANS LE MÊME DOMAINE crée une
    ambiguïté réelle (mission :8810, audit du champ Référence) : deux colonnes
    identiques dans l'export des réponses individuelles, et le même libellé sur les
    histogrammes quand la référence est courte/alphabétique (pdf_short_label). Ce
    n'est jamais un bug de calcul — chaque indicateur reste distinct par son id — mais
    une confusion évitable pour l'utilisateur. Même périmètre (par domaine, pas
    global) que le contrôle déjà appliqué à l'import XLSX (import_preview)."""
    row = db.execute(
        "SELECT 1 FROM indicators WHERE domain_id=? AND active=1 AND code=? AND id!=? LIMIT 1",
        (domain_id, code, exclude_id or ""),
    ).fetchone()
    return row is not None


# Mots grammaticaux (articles/prépositions/pronoms) + verbes/adverbes génériques des
# formulations EPC ("Nous offrons régulièrement...", "Il existe...") à retirer pour
# proposer une Référence courte (mission :8810, demande Mouhamed BA). Best-effort
# déterministe — jamais parfait, toujours corrigible par l'utilisateur avant
# enregistrement (voir propose_indicator_reference).
INDICATOR_REFERENCE_STOPWORDS = {
    "nous", "notre", "nos", "vous", "vos", "votre", "leur", "leurs", "on", "il", "ils", "elle", "elles",
    "le", "la", "les", "l", "de", "des", "du", "d", "et", "en", "au", "aux", "à", "a", "ce", "cet", "cette", "ces",
    "un", "une", "pour", "que", "qui", "dans", "avec", "sur", "par", "est", "sont", "ont", "avons", "avez", "être", "avoir",
    "tout", "tous", "toute", "toutes", "très", "plus", "moins", "bien", "mal", "comme",
    "régulièrement", "systématiquement", "activement", "efficacement", "généralement", "directement",
    "constamment", "correctement", "également", "notamment", "ainsi", "donc",
    "offrons", "offre", "offrent", "utilisons", "utilise", "utilisent", "procédons", "modifions", "modifie",
    "examinons", "examine", "engageons", "engage", "établissons", "établit", "disposons", "dispose",
    "reconnaissons", "reconnait", "impliquons", "implique", "faisons", "fait", "accordons", "accorde",
    "obtenons", "obtient", "participent", "participe", "contribuent", "contribue", "favorise", "favorisent",
    "encourage", "encouragent", "améliore", "améliorent", "reflète", "reflètent", "montre", "montrent",
    "permettent", "permet", "existe", "existent", "appropriée", "appropriées", "approprié", "appropriés",
}


def _significant_reference_words(label: str) -> list[str]:
    words = re.findall(r"[A-Za-zÀ-ÿŒœÆæ]+(?:-[A-Za-zÀ-ÿŒœÆæ]+)*", (label or "").strip())
    significant = [w for w in words if w.lower().strip("’'") not in INDICATOR_REFERENCE_STOPWORDS]
    return significant or words


def _reference_phrase(words: list[str]) -> str:
    if not words:
        return ""
    phrase = " ".join(words)
    return phrase[0].upper() + phrase[1:]


def generate_indicator_reference(label: str) -> str:
    """Propose une Référence courte (cible 3-4 mots) à partir de l'énoncé complet
    d'un indicateur, sans dépendance à un service IA (mission :8810). Retire les
    mots grammaticaux et les tournures génériques ("Nous offrons régulièrement..."),
    garde les mots significatifs dans leur ordre d'apparition, et capitalise
    uniquement la première lettre du résultat — jamais parfait, toujours modifiable
    par l'utilisateur avant enregistrement (voir INDICATOR_REFERENCE_STOPWORDS)."""
    return _reference_phrase(_significant_reference_words(label)[:4])


def propose_indicator_reference(db: sqlite3.Connection, domain_id: str, label: str, exclude_id: str | None = None) -> str:
    """Comme generate_indicator_reference, mais vérifie l'unicité dans le domaine
    (règle conservée de la mission précédente) : si la proposition à 4 mots existe
    déjà, tente une variante à 5 mots plutôt qu'un numéro incompréhensible ; si la
    collision persiste, retourne la proposition de base — l'enregistrement refusera
    alors le doublon et invitera l'utilisateur à la différencier lui-même."""
    words = _significant_reference_words(label)
    base = _reference_phrase(words[:4])
    if not base or not indicator_code_conflicts(db, domain_id, base, exclude_id=exclude_id):
        return base
    if len(words) > 4:
        extended = _reference_phrase(words[:5])
        if not indicator_code_conflicts(db, domain_id, extended, exclude_id=exclude_id):
            return extended
    return base


def create_session(db: sqlite3.Connection, owner_user_id: str, data: dict):
    """Le bouton "Nouveau diagnostic" ("Créer la mission") crée une MISSION/atelier
    qui pointe vers un questionnaire déjà existant (data["templateId"], le canonique
    EPC/SENEVAL par défaut) — il ne crée jamais de nouveau questionnaire. Vérifié ici
    explicitement (mission :8810, point B : le libellé "Créer le questionnaire" aurait
    été trompeur pour cette action, donc non appliqué). Retourne None si le
    questionnaire visé n'a aucun domaine actif avec question (rien n'est créé)."""
    template = template_payload(db, data["templateId"])
    if not template or not any(d["active"] and any(i["active"] for i in d["indicators"]) for d in template["domains"]):
        return None
    sid = str(uuid.uuid4())
    db.execute("INSERT INTO sessions VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (sid, template["id"], template["version"], data["name"], data.get("organization", ""), data.get("location", ""), data.get("date", ""),
         "open", now(), None, data.get("description", ""), int(data["expectedParticipants"]) if data.get("expectedParticipants") not in (None, "") else None, owner_user_id))
    db.commit()
    return sid


def update_participant_profile_fields(db: sqlite3.Connection, pid: str, data: dict) -> None:
    sets, vals = [], []
    for key, col in PARTICIPANT_PROFILE_FIELD_MAP.items():
        if key in data:
            v = data[key]
            v = (1 if v else 0) if key == "anonymous" else (v or None)
            sets.append(f"{col}=?"); vals.append(v)
    if sets:
        vals.append(pid)
        db.execute(f"UPDATE participants SET {','.join(sets)} WHERE id=?", vals)
        db.commit()


def session_label(db: sqlite3.Connection, sid: str) -> str:
    row = db.execute("SELECT name FROM sessions WHERE id=?", (sid,)).fetchone()
    return row["name"] if row else "atelier"


def template_payload(db, template_id: str):
    template = db.execute("SELECT * FROM templates WHERE id=?", (template_id,)).fetchone()
    if not template:
        return None
    out = dict(template)
    for key in ["scale_json", "scoring_json", "consensus_json", "grading_json", "priority_json"]:
        out[key[:-5]] = json.loads(out.pop(key))
    out["domains"] = rows(db, "SELECT * FROM domains WHERE template_id=? ORDER BY display_order", (template_id,))
    for domain in out["domains"]:
        domain["active"] = bool(domain["active"])
        domain["indicators"] = rows(db, "SELECT * FROM indicators WHERE domain_id=? ORDER BY display_order", (domain["id"],))
        for indicator in domain["indicators"]:
            indicator["required"] = bool(indicator["required"]); indicator["active"] = bool(indicator["active"]); indicator["configuration"] = json.loads(indicator.pop("configuration_json"))
    return out


def next_order(db, table, where_col, where_value):
    return (db.execute(f"SELECT COALESCE(MAX(display_order),0)+1 FROM {table} WHERE {where_col}=?", (where_value,)).fetchone()[0])


def clone_template_with_mapping(db, template_id, name=None, status="active"):
    """Same as clone_template, but also returns the old->new id maps for domains and
    indicators, so a caller can re-point a single session's own historical rows
    (responses/priorities/...) onto the fresh copy without touching any other session
    that still legitimately points at the original ids."""
    old = template_payload(db, template_id)
    if not old: raise ValueError("Configuration introuvable")
    tid, stamp = str(uuid.uuid4()), now()
    version = db.execute("SELECT COALESCE(MAX(version),0)+1 FROM templates WHERE name=?", (name or old["name"],)).fetchone()[0]
    db.execute("INSERT INTO templates VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (tid, name or old["name"], version, old["description"], status, json.dumps(old["scale"]), json.dumps(old["scoring"]), json.dumps(old["consensus"]), json.dumps(old["grading"]), json.dumps(old["priority"]), stamp, stamp, old.get("owner_user_id"), 0))
    domain_map, indicator_map = {}, {}
    for d in old["domains"]:
        new_did = str(uuid.uuid4()); domain_map[d["id"]] = new_did
        db.execute("INSERT INTO domains VALUES (?,?,?,?,?,?,?)",(new_did,tid,d["code"],d["label"],d["description"],d["display_order"],d["active"]))
        for i in d["indicators"]:
            new_iid = str(uuid.uuid4()); indicator_map[i["id"]] = new_iid
            db.execute("INSERT INTO indicators VALUES (?,?,?,?,?,?,?,?,?,?)",(new_iid,new_did,i["code"],i["label"],i["description"],i["response_type"],i["required"],i["display_order"],i["active"],json.dumps(i["configuration"])))
    db.commit()
    return tid, domain_map, indicator_map


def clone_template(db, template_id, name=None, status="active"):
    tid, _, _ = clone_template_with_mapping(db, template_id, name=name, status=status)
    return tid


def is_canonical_template(db, template_id) -> bool:
    row = db.execute("SELECT is_canonical FROM templates WHERE id=?", (template_id,)).fetchone()
    return bool(row and row["is_canonical"])


class StructuralEditForbiddenError(Exception):
    """Raised when an add/delete on domains/indicators targets the canonical EPC/SENEVAL
    reference template — never raised for a mission's private fork or a user's own custom
    questionnaire, which remain freely editable."""


CANONICAL_STRUCTURE_MESSAGES = {
    "add": "Ajout impossible : EPC / SENEVAL est le questionnaire de référence et doit garder exactement sa structure validée (7 domaines / 70 indicateurs, dans l'ordre validé). Dupliquez-le pour créer votre propre questionnaire modifiable.",
    "delete": "Suppression impossible : EPC / SENEVAL est le questionnaire de référence et doit garder exactement sa structure validée (7 domaines / 70 indicateurs, dans l'ordre validé). Dupliquez-le pour créer votre propre questionnaire modifiable.",
    "edit": "Modification impossible : EPC / SENEVAL est le questionnaire de référence et doit garder exactement son contenu validé. Dupliquez-le pour créer votre propre questionnaire modifiable.",
    "delete_template": "Suppression impossible. EPC/SENEVAL est le modèle de référence ; dupliquez-le pour le modifier.",
}


def guard_structural_edit(db, template_id, action="add") -> None:
    if is_canonical_template(db, template_id):
        raise StructuralEditForbiddenError(CANONICAL_STRUCTURE_MESSAGES[action])


DOMAIN_FROZEN_MESSAGE = "Ce questionnaire contient déjà des réponses dans {count} atelier(s) ({names}). Sa structure est désormais figée afin de préserver la cohérence des données : désactivez ce domaine pour préserver l'historique, ou dupliquez le questionnaire pour créer une nouvelle version modifiable."
INDICATOR_FROZEN_MESSAGE = "Ce questionnaire contient déjà {count} réponse(s) pour cette question. Sa structure est désormais figée afin de préserver la cohérence des données : désactivez cette question pour préserver l'historique, ou dupliquez le questionnaire pour créer une nouvelle version modifiable."


def ensure_private_template(db: sqlite3.Connection, session_id: str):
    """Fork this session's questionnaire onto a private, session-exclusive copy the
    moment it needs one — either because its very first participant is being created,
    or during the one-time retroactive pass for missions that already started collecting
    before this fix existed — and re-point this session's OWN historical rows
    (responses/priorities/analysis_notes/recommendations) onto the fresh copy's new ids,
    so its past answers keep resolving against the exact domains/indicators it actually
    used. No other session's rows are ever touched.

    "Still shared" vs "already private" is decided purely by a stable, race-free signal —
    never by name/text matching, and never by counting how many sessions currently share
    the template (a session that happens to be the only current user of a still-"active"
    pool template — canonical or a custom one — is NOT yet safe: another mission could
    start using that same template_id later and would then be free to edit it). The only
    templates ever marked status='session_locked' are ones already forked exclusively for
    one session by this very function, so that flag alone is a reliable, order-independent
    "already private" signal. Idempotent and cheap to call on every relevant request.
    """
    row = db.execute("SELECT template_id, owner_user_id FROM sessions WHERE id=?", (session_id,)).fetchone()
    if not row:
        return None
    tid = row["template_id"]
    current = db.execute("SELECT status FROM templates WHERE id=?", (tid,)).fetchone()
    if current and current["status"] == "session_locked":
        return tid
    new_tid, domain_map, indicator_map = clone_template_with_mapping(db, tid, status="session_locked")
    db.execute("UPDATE templates SET owner_user_id=? WHERE id=?", (row["owner_user_id"], new_tid))
    for old_iid, new_iid in indicator_map.items():
        db.execute("UPDATE responses SET indicator_id=? WHERE session_id=? AND indicator_id=?", (new_iid, session_id, old_iid))
        db.execute("UPDATE analysis_notes SET indicator_id=? WHERE session_id=? AND indicator_id=?", (new_iid, session_id, old_iid))
        db.execute("UPDATE recommendations SET indicator_id=? WHERE session_id=? AND indicator_id=?", (new_iid, session_id, old_iid))
        db.execute("UPDATE priorities SET indicator_id=? WHERE session_id=? AND indicator_id=?", (new_iid, session_id, old_iid))
    for old_did, new_did in domain_map.items():
        db.execute("UPDATE priorities SET domain_id=? WHERE session_id=? AND domain_id=?", (new_did, session_id, old_did))
    new_version = db.execute("SELECT version FROM templates WHERE id=?", (new_tid,)).fetchone()["version"]
    db.execute("UPDATE sessions SET template_id=?,template_version=? WHERE id=?", (new_tid, new_version, session_id))
    db.commit()
    return new_tid


def ensure_private_templates_for_started_sessions(db: sqlite3.Connection) -> None:
    """One-time (idempotent) retroactive pass: fork the questionnaire for every session
    that already has at least one participant, so missions whose collecte began before
    this fix was deployed — including any mission with real validated responses/
    priorities/analyses recorded — get protected the moment this fix ships. Safe to
    re-run: a no-op for sessions already on an exclusive template."""
    for s in rows(db, "SELECT DISTINCT session_id AS id FROM participants"):
        ensure_private_template(db, s["id"])


def delete_session(db: sqlite3.Connection, sid: str) -> bool:
    """Permanently delete a mission and all its data. If its questionnaire is a private fork
    (status='session_locked') that belongs exclusively to this mission — never the canonical
    template, never a template still referenced by another session — delete that fork too, so
    private forks don't outlive the one mission they were created for. Returns True if the fork
    was dropped, for tests/diagnostics."""
    session_row = db.execute("SELECT template_id FROM sessions WHERE id=?", (sid,)).fetchone()
    drop_template, tid = False, None
    if session_row:
        tid = session_row["template_id"]
        tpl = db.execute("SELECT status,is_canonical FROM templates WHERE id=?", (tid,)).fetchone()
        other_ref = db.execute("SELECT 1 FROM sessions WHERE template_id=? AND id!=? LIMIT 1", (tid, sid)).fetchone()
        drop_template = bool(tpl and not tpl["is_canonical"] and tpl["status"] == "session_locked" and not other_ref)
    for table in ("training_topics","workshop_recommendations","analysis_entries","priority_analyses","analysis_notes","recommendations","responses","priorities","participants","session_report_meta"):
        db.execute(f"DELETE FROM {table} WHERE session_id=?", (sid,))
    db.execute("DELETE FROM sessions WHERE id=?", (sid,))
    if drop_template:
        db.execute("DELETE FROM indicators WHERE domain_id IN (SELECT id FROM domains WHERE template_id=?)", (tid,))
        db.execute("DELETE FROM domains WHERE template_id=?", (tid,))
        db.execute("DELETE FROM templates WHERE id=?", (tid,))
    db.commit()
    return drop_template


def create_blank_template(db, data, owner_user_id=None):
    tid, stamp = str(uuid.uuid4()), now(); name=data.get("name", "Nouveau questionnaire").strip()
    if not name: raise ValueError("Le nom est obligatoire")
    version=db.execute("SELECT COALESCE(MAX(version),0)+1 FROM templates WHERE name=?",(name,)).fetchone()[0]
    scale={"type":"numeric","min":1,"max":5,"labels":{}}
    db.execute("INSERT INTO templates VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(tid,name,version,data.get("description", ""),"active",json.dumps(scale),json.dumps({"capacity":"mean_divided_by_scale_max","outputRange":[0,100]}),json.dumps({"method":"standard_deviation","normalization":"theoretical_range","factor":2}),json.dumps(GRADING),json.dumps({"maxPerDomain":3}),stamp,stamp,owner_user_id,0)); db.commit(); return tid


def matrix_xlsx(template):
    if not xlsxwriter: raise RuntimeError("Le générateur XLSX local n'est pas disponible")
    out=BytesIO(); wb=xlsxwriter.Workbook(out, {"in_memory": True}); head=wb.add_format({"bold":True,"bg_color":"#1F4E78","font_color":"#FFFFFF"}); wrap=wb.add_format({"text_wrap":True,"valign":"top"})
    guide=wb.add_worksheet("MODE D’EMPLOI"); guide.set_column(0,0,110); guide.write("A1","Cette matrice permet de préparer un questionnaire avant de l’importer dans l’outil.",head); guide.write_column("A3",["1. Dans la feuille PARAMETRES, remplacez la valeur d’exemple par le vrai nom de votre questionnaire.","2. Complétez la description (facultatif) et les libellés de l’échelle de notation (une seule fois, valables pour tout le questionnaire).","3. Dans la feuille QUESTIONNAIRE, saisissez une ligne par indicateur.","4. Répétez le nom du domaine pour les indicateurs appartenant au même domaine.","5. La numérotation sera générée automatiquement par l’outil.","6. Les lignes d’exemple (matrice PARAMETRES et QUESTIONNAIRE) peuvent être remplacées ou supprimées : elles servent uniquement de modèle, à l’image du questionnaire EPC/SENEVAL."],wrap)
    default_labels={"5":"Totalement d’accord","4":"D’accord","3":"Neutre","2":"Pas d’accord","1":"Totalement en désaccord"}
    smin,smax=int(template["scale"]["min"]),int(template["scale"]["max"])
    ps=wb.add_worksheet("PARAMETRES"); ps.write_row(0,0,["Nom du questionnaire (à remplacer par le vôtre)",template["name"]],head); ps.write_row(1,0,["Description",template["description"]],wrap); ps.write_row(3,0,["Note","Libellé (exemple EPC/SENEVAL, à adapter)"],head); labels=template["scale"].get("labels",{}); [ps.write_row(4+(smax-n),0,[n,labels.get(str(n)) or default_labels.get(str(n),'')]) for n in range(smax,smin-1,-1)]; ps.set_column(0,0,38); ps.set_column(1,1,55)
    ws=wb.add_worksheet("QUESTIONNAIRE"); ws.write_row(0,0,["Domaine","Référence","Indicateur qualitatif ou Capacité"],head); ws.freeze_panes(1,0); row=1
    if not template["domains"]: ws.write_row(row,0,["EXEMPLE — Gestion des Ressources Humaines","EXEMPLE — Formation au personnel","EXEMPLE — Nous offrons régulièrement la formation au personnel"],wrap); row+=1
    for d in template["domains"]:
        for i in d["indicators"]:
            ws.write_row(row,0,[d["label"],i["label"],i["description"]],wrap); row+=1
    ws.set_column(0,0,34); ws.set_column(1,1,38); ws.set_column(2,2,75); wb.close(); return out.getvalue()


def blank_matrix_xlsx():
    return matrix_xlsx({"name":"Exemple à remplacer : Diagnostic EPC / [nom de l’atelier]", "description":"", "version":1, "scale":{"type":"numeric","min":1,"max":5,"labels":{}}, "priority":{"maxPerDomain":3}, "domains":[]})

def report_rows(db, sid):
    a=analysis(db,sid); return a, [[d['label'],d['capacity'],_c(d),d['gradedCapacity'],d['gradedConsensus'],d['responses']] for d in a['domains']]

def individual_responses_rows(db, session_id: str):
    """Wide export shape: one row per validated (status='completed') participant of this
    session only, one column per active indicator. Anonymous participants never expose
    display_name, regardless of what was stored, so anonymat holds even if the client
    ever sends a stale name alongside anonymous=1."""
    session = db.execute("SELECT template_id FROM sessions WHERE id=?", (session_id,)).fetchone()
    if not session:
        return {"indicators": [], "rows": []}
    template = template_payload(db, session["template_id"])
    indicators = [i for d in template["domains"] for i in d["indicators"] if i["active"]]
    participants = rows(db, "SELECT * FROM participants WHERE session_id=? AND status='completed' ORDER BY started_at", (session_id,))
    out_rows = []
    for p in participants:
        anonymous = bool(p["anonymous"])
        responses = {r["indicator_id"]: json.loads(r["value_json"]) for r in db.execute("SELECT indicator_id,value_json FROM responses WHERE session_id=? AND participant_id=?", (session_id, p["id"]))}
        row = {"id": p["anonymous_id"], "name": "" if anonymous else (p["display_name"] or ""), "status": "Anonyme" if anonymous else "Nominatif",
               "participant_type": p["participant_type"] or "", "profile": p["profile"] or "", "sex": p["sex"] or "",
               "age_range": p["age_range"] or "", "education_level": p["education_level"] or ""}
        for ind in indicators:
            row[ind["id"]] = responses.get(ind["id"], "")
        out_rows.append(row)
    return {"indicators": indicators, "rows": out_rows}


INDIVIDUAL_RESPONSES_HEAD = ["Identifiant participant", "Nom", "Statut", "Type de participant", "Profil", "Sexe", "Tranche d'âge", "Niveau de scolarisation"]


def _individual_responses_table(db, session_id: str):
    data = individual_responses_rows(db, session_id)
    head = INDIVIDUAL_RESPONSES_HEAD + [i["code"] for i in data["indicators"]]
    body = [[row["id"], row["name"], row["status"], row["participant_type"], row["profile"], row["sex"], row["age_range"], row["education_level"]] + [row[i["id"]] for i in data["indicators"]] for row in data["rows"]]
    return head, body


def individual_responses_xlsx(db, session_id: str):
    head, body = _individual_responses_table(db, session_id)
    out = BytesIO(); wb = xlsxwriter.Workbook(out, {"in_memory": True}); h = wb.add_format({"bold": True, "bg_color": "#1F4E78", "font_color": "#FFFFFF"})
    s = wb.add_worksheet("Réponses individuelles"); s.write_row(0, 0, head, h)
    for n, row in enumerate(body): s.write_row(n + 1, 0, row)
    s.set_column(0, len(head) - 1, 20)
    wb.close(); return out.getvalue()


def individual_responses_csv(db, session_id: str):
    head, body = _individual_responses_table(db, session_id)
    buf = StringIO(); writer = csv.writer(buf); writer.writerow(head)
    for row in body: writer.writerow(row)
    return buf.getvalue().encode()


PARTICIPANT_FILTER_LABELS = {"participant_type": "Type de participant", "profile": "Profil", "sex": "Sexe", "age_range": "Tranche d'âge", "education_level": "Niveau de scolarisation"}


def filtered_analysis_rows(db, session_id: str, participant_filter):
    """Header block (mission/filtres actifs/N) + domain table + indicator table for
    the filtered-analysis export — distinct from individual_responses_rows (raw
    per-participant answers)."""
    a = analysis(db, session_id, participant_filter)
    filt_desc = "; ".join(f"{PARTICIPANT_FILTER_LABELS[col]}={val}" for col, val in (participant_filter or {}).items()) or "Tous les participants"
    header = [["Mission", a["session"]["name"]], ["Filtres actifs", filt_desc], ["N (validés)", a["completedCount"]]]
    domain_rows = [[d["label"], d["capacity"], _c(d), d["gradedCapacity"], d["gradedConsensus"], d["responses"]] for d in a["domains"]]
    indicator_rows = [[d["label"], i["label"], i["capacity"], _c(i), i["responses"], i["missing"]] for d in a["domains"] for i in d["indicators"]]
    return header, domain_rows, indicator_rows


def filtered_analysis_xlsx(db, session_id: str, participant_filter):
    header, domain_rows, indicator_rows = filtered_analysis_rows(db, session_id, participant_filter)
    out = BytesIO(); wb = xlsxwriter.Workbook(out, {"in_memory": True}); h = wb.add_format({"bold": True, "bg_color": "#1F4E78", "font_color": "#FFFFFF"})
    s = wb.add_worksheet("Filtre"); [s.write_row(n, 0, row) for n, row in enumerate(header)]; s.set_column(0, 1, 30)
    d = wb.add_worksheet("Domaines"); d.write_row(0, 0, ["Domaine", "Capacité", "Consensus", "Cap. graduée", "Cons. gradué", "Réponses"], h); [d.write_row(n + 1, 0, row) for n, row in enumerate(domain_rows)]; d.set_column(0, 5, 22)
    i = wb.add_worksheet("Indicateurs"); i.write_row(0, 0, ["Domaine", "Référence", "Capacité", "Consensus", "Réponses", "Manquants"], h); [i.write_row(n + 1, 0, row) for n, row in enumerate(indicator_rows)]; i.set_column(0, 5, 24)
    wb.close(); return out.getvalue()


def filtered_analysis_csv(db, session_id: str, participant_filter):
    header, domain_rows, indicator_rows = filtered_analysis_rows(db, session_id, participant_filter)
    buf = StringIO(); writer = csv.writer(buf)
    for row in header: writer.writerow(row)
    writer.writerow([]); writer.writerow(["Domaine", "Capacité", "Consensus", "Cap. graduée", "Cons. gradué", "Réponses"])
    for row in domain_rows: writer.writerow(row)
    writer.writerow([]); writer.writerow(["Domaine", "Référence", "Capacité", "Consensus", "Réponses", "Manquants"])
    for row in indicator_rows: writer.writerow(row)
    return buf.getvalue().encode()


def participant_profile_breakdown(db, session_id: str):
    """Count of validated participants per value, for each of the 5 fixed profile
    fields — blank/NULL excluded (never invented as a fake 'non renseigné' count)."""
    out = []
    for col, label in PARTICIPANT_FILTER_LABELS.items():
        for r in rows(db, f"SELECT {col} AS v, COUNT(*) AS n FROM participants WHERE session_id=? AND status='completed' AND {col} IS NOT NULL AND {col}!='' GROUP BY {col} ORDER BY n DESC", (session_id,)):
            out.append([label, r["v"], r["n"]])
    return out


def findings_rows(findings):
    """Flatten objective_findings()'s dict into plain rows for XLSX/DOCX/report display."""
    rows_out = []
    for d in findings["forces"]["domains"]: rows_out.append(["Force", "Domaine", d["label"], d["capacity"], d["consensus"]])
    for i in findings["forces"]["indicators"]: rows_out.append(["Force", "Indicateur", f"{i['domain']} — {i['label']}", i["capacity"], i["consensus"]])
    for d in findings["fragilites"]["domains"]: rows_out.append(["Fragilité", "Domaine", d["label"], d["capacity"], d["consensus"]])
    for i in findings["fragilites"]["indicators"]: rows_out.append(["Fragilité", "Indicateur", f"{i['domain']} — {i['label']}", i["capacity"], i["consensus"]])
    for v in findings["vigilance"]:
        if v["reason"] == "ecart_sous_populations":
            rows_out.append(["Vigilance", "Écart entre sous-populations", v["label"], v.get("gap"), ""])
        else:
            rows_out.append(["Vigilance", "Domaine", v["label"], v.get("capacity"), v.get("consensus")])
    return rows_out


def report_xlsx(db,sid):
    a,rs=report_rows(db,sid); q=qualitative_data(db,sid); meta=report_data(db,sid)["meta"]; template=template_payload(db,a['session']['template_id']); out=BytesIO(); wb=xlsxwriter.Workbook(out,{"in_memory":True}); h=wb.add_format({"bold":True,"bg_color":"#1F4E78","font_color":"#FFFFFF"})
    analyses={x['priority_id']:x for x in q['analyses']}; priority_rows=[[p['id'],p['domain_label'],p['indicator_code'],p['indicator_label'],analyses.get(p['id'],{}).get('problem','')] for p in q['priorities']]
    sheets=[("Synthèse",["Atelier","Organisation","Lieu","Date","Animateur","Public","Contexte","Conclusion","Capacité","Consensus"],[[a['session']['name'],a['session']['organization'],a['session']['location'],a['session']['date'],meta['facilitator'],meta['audience'],meta['context'],meta['conclusion'],a['global']['capacity'],_c(a['global'])]]),("Profil_participants",["Champ","Valeur","N"],participant_profile_breakdown(db,sid)),("Domaines",["Domaine","Capacité","Consensus","Cap. graduée","Cons. gradué","Réponses"],rs),("Indicateurs",["Domaine","Référence","Capacité","Consensus","Réponses","Manquants"],[[d['label'],i['label'],i['capacity'],_c(i),i['responses'],i['missing']] for d in a['domains'] for i in d['indicators']]),("Constats",["Type","Niveau","Libellé","Capacité","Consensus"],findings_rows(a['findings'])),("Priorités",["ID priorité","Domaine","Référence","Indicateur","Constat"],priority_rows),("Analyses",["ID","Priorité","Constat"],[[x['id'],x['priority_id'],x['problem']] for x in q['analyses']]),("Causes",["ID","Priorité","Parent","Cause","Type","Statut"],[[x['id'],x['priority_id'],x['parent_id'],x['content'],x['item_type'],x['validation_status']] for x in q['entries'] if x['kind']=='cause']),("Conséquences",["ID","Priorité","Conséquence","Statut"],[[x['id'],x['priority_id'],x['content'],x['validation_status']] for x in q['entries'] if x['kind']=='consequence']),("Leviers",["ID","Priorité","Levier","Commentaire","Statut"],[[x['id'],x['priority_id'],x['content'],x['comment'],x['validation_status']] for x in q['entries'] if x['kind']=='lever']),("Recommandations",["ID","Priorité","Cause","Levier","Titre","Description","Catégorie","Niveau","Responsable","Échéance","Statut"],[[x['id'],x['priority_id'],x['cause_id'],x['lever_id'],x['title'],x['description'],x['category'],x['priority_level'],x['owner'],x['horizon'],x['status']] for x in q['recommendations']]),("Formations",["ID","Priorité","Recommandation","Intitulé","Besoin","Public","Niveau","Commentaire"],[[x['id'],x['priority_id'],x['recommendation_id'],x['title'],x['need_text'],x['target_audience'],x['priority_level'],x['comment']] for x in q['trainingTopics']]),("Plan_action",["N°","Action / recommandation","Origine","Responsable","Échéance","Priorité","Statut"],[[n+1,x['title'],x['priority_id'] or '—',x['owner'] or '—',x['horizon'] or '—',x['priority_level'],x['status']] for n,x in enumerate(q['recommendations']) if x['status']=='Retenue']),("Questionnaire",["Domaine","Référence","Indicateur","Échelle"],[[d['label'],i['label'],i['description'],f"{template['scale']['min']}–{template['scale']['max']}"] for d in template['domains'] for i in d['indicators'] if i['active']])]
    for name,head,data in sheets:
        s=wb.add_worksheet(name);s.write_row(0,0,head,h);[s.write_row(n+1,0,row) for n,row in enumerate(data)];s.set_column(0,len(head)-1,24)
    ai_blocks=rows(db,"SELECT section_key,content FROM report_ai_blocks WHERE session_id=?",(sid,))
    if ai_blocks:
        wrap=wb.add_format({"text_wrap":True,"valign":"top"})
        ai_sheet=wb.add_worksheet("Synthèse_IA"); ai_sheet.write_row(0,0,["Section","Contenu proposé par l'assistant IA, retenu par le modérateur"],h)
        for n,b in enumerate(ai_blocks): ai_sheet.write_row(n+1,0,[AI_SECTION_LABELS.get(b['section_key'],b['section_key']),b['content']],wrap)
        ai_sheet.set_column(0,0,28); ai_sheet.set_column(1,1,110)
    domains=[d for d in a['domains'] if d.get('capacity') is not None]
    if PILImage is not None and domains:
        gs=wb.add_worksheet("Graphiques"); gs.set_column(0,0,4); row=1
        gs.write(0,0,"Vue synthétique",h)
        row=xlsx_add_image(gs,row,0,pil_chart_grid([
            docx_radar_chart(domains),
            docx_bars_chart(domains,"standard","Notes standardisées par domaine",with_mean=True),
            docx_bars_chart(domains,"graded","Notes graduées par domaine",with_mean=True),
            docx_grid_chart(domains),
        ]))
        if len(domains)>1:
            row+=1; gs.write(row,0,"Analyse comparative — cohorte des domaines",h); row+=1
            row=xlsx_add_image(gs,row,0,docx_cohort_chart(domains))
        for i in range(0,len(domains),4):
            group=domains[i:i+4]
            tiles=[]
            for d in group:
                inds=[ind for ind in d["indicators"] if ind.get("capacity") is not None]
                if inds:
                    caption=f"Capacité {pdf_fmt(d['capacity'])} · Consensus {pdf_c(d)} · {d['responses']} répondant(s)"
                    tiles.append((docx_bars_chart(inds,"standard",d["label"]),caption))
            if tiles:
                row+=1; gs.write(row,0,"Détail par domaine",h); row+=1
                row=xlsx_add_image(gs,row,0,pil_chart_grid(tiles))
        gs.activate()
    wb.close();return out.getvalue()

DOCX_FONT_CACHE = {}


def docx_font(size, bold=False):
    key = (size, bold)
    if key in DOCX_FONT_CACHE:
        return DOCX_FONT_CACHE[key]
    names = ["arialbd.ttf", "Arial Bold.ttf", "DejaVuSans-Bold.ttf", "LiberationSans-Bold.ttf"] if bold else \
            ["arial.ttf", "Arial.ttf", "DejaVuSans.ttf", "LiberationSans-Regular.ttf"]
    dirs = ["C:/Windows/Fonts", "/usr/share/fonts/truetype/dejavu", "/usr/share/fonts/truetype/liberation",
            "/System/Library/Fonts/Supplemental", "/System/Library/Fonts"]
    font = None
    for d in dirs:
        for n in names:
            try:
                font = ImageFont.truetype(f"{d}/{n}", size); break
            except Exception:
                continue
        if font:
            break
    if font is None:
        try:
            font = ImageFont.load_default(size=size)
        except TypeError:
            font = ImageFont.load_default()
    DOCX_FONT_CACHE[key] = font
    return font


def docx_text_w(draw, text, font):
    b = draw.textbbox((0, 0), text, font=font)
    return b[2] - b[0]


def docx_rotated_text(img, text, x, y, angle, font, fill, anchor_end=False):
    draw0 = ImageDraw.Draw(img)
    b = draw0.textbbox((0, 0), text, font=font)
    w, h = b[2] - b[0], b[3] - b[1]
    pad = 4
    txt_img = PILImage.new("RGBA", (w + pad * 2, h + pad * 2), (255, 255, 255, 0))
    ImageDraw.Draw(txt_img).text((pad, pad), text, font=font, fill=fill)
    rot = txt_img.rotate(angle, expand=1, resample=PILImage.BICUBIC)
    px = x - (rot.width if anchor_end else 0)
    img.paste(rot, (int(px), int(y)), rot)


def docx_dashed_line(draw, p1, p2, color, dash=8, gap=5, width=2):
    x1, y1 = p1; x2, y2 = p2
    length = math.hypot(x2 - x1, y2 - y1)
    if length == 0:
        return
    ux, uy = (x2 - x1) / length, (y2 - y1) / length
    d = 0
    while d < length:
        e = min(d + dash, length)
        draw.line([x1 + ux * d, y1 + uy * d, x1 + ux * e, y1 + uy * e], fill=color, width=width)
        d += dash + gap


def docx_bars_chart(items, mode, title, with_mean=False):
    """Vertical grouped column chart, raster twin of static/app.js bars() / pdf_bars_chart()."""
    grad = mode == "graded"
    data = [d for d in items if d.get("capacity") is not None]
    if with_mean and data:
        mc = sum(d["capacity"] for d in data) / len(data)
        cons_data = [d for d in data if d.get("consensus") is not None]
        ms = sum(d["consensus"] for d in cons_data) / len(cons_data) if cons_data else None
        gc = sum(d.get("gradedCapacity") or 0 for d in data) / len(data)
        gs = sum(d.get("gradedConsensus") or 0 for d in data) / len(data)
        data = data + [{"label": "Moyen", "code": "Moyen", "capacity": mc, "consensus": ms, "gradedCapacity": gc, "gradedConsensus": gs}]
    w = max(900, 90 + len(data) * 140) if data else 900
    h = 560
    img = PILImage.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    f_title, f_axis, f_legend, f_val, f_cat, f_cat_b, f_xtitle = docx_font(24, True), docx_font(15), docx_font(15), docx_font(14), docx_font(14), docx_font(14, True), docx_font(17, True)
    draw.text((w / 2, 30), title, font=f_title, fill="black", anchor="mm")
    if not data:
        draw.text((30, 66), "Pas encore de données disponibles.", font=f_axis, fill="black")
        return img
    left, right = 100, w - 30
    top, bottom = 72, h - 120
    plot_w, plot_h = right - left, bottom - top
    axis_min, axis_max = 20, 100
    scaleY = lambda v: bottom - (v - axis_min) / (axis_max - axis_min) * plot_h
    draw.rectangle([left, top, right, bottom], fill="#e9e9e9", outline="black")
    for v in (20, 40, 60, 80, 100):
        yy = scaleY(v)
        draw.line([left, yy, right, yy], fill="#bbbbbb")
        draw.text((left - 12, yy), str(v), font=f_axis, fill="black", anchor="rm")
    cap_label = "Capacité graduée" if grad else "Capacité standardisée"
    cons_label = "Consensus gradué" if grad else "Consensus standardisé"
    draw.rectangle([left, 40, left + 16, 56], fill="#9999FF")
    draw.text((left + 24, 48), cap_label, font=f_legend, fill="black", anchor="lm")
    draw.rectangle([left + 240, 40, left + 256, 56], fill="#993366")
    draw.text((left + 264, 48), cons_label, font=f_legend, fill="black", anchor="lm")
    n = len(data); step = plot_w / n
    for i, d in enumerate(data):
        x = left + i * step
        cap = d.get("gradedCapacity") if grad else d.get("capacity")
        cons = d.get("gradedConsensus") if grad else d.get("consensus")
        cv, sv = max(axis_min, cap or 0), max(axis_min, cons or 0)
        bw = step * 0.32
        y_cap, y_cons = scaleY(cv), scaleY(sv)
        draw.rectangle([x + step * .12, y_cap, x + step * .12 + bw, bottom], fill="#9999FF", outline="black")
        draw.rectangle([x + step * .12 + bw + 4, y_cons, x + step * .12 + bw + 4 + bw, bottom], fill="#993366", outline="black")
        cap_txt = "—" if cap is None else (str(round(cap)) if grad else f"{cap:.1f}")
        cons_txt = "—" if cons is None else (str(round(cons)) if grad else f"{cons:.1f}")
        draw.text((x + step * .12 + bw / 2, y_cap - 12), cap_txt, font=f_val, fill="black", anchor="mm")
        draw.text((x + step * .12 + bw + 4 + bw / 2, y_cons - 12), cons_txt, font=f_val, fill="black", anchor="mm")
        docx_rotated_text(img, pdf_short_label(d), x + step * .12 + bw + 2, bottom + 8, -40, f_cat_b if d.get("label") == "Moyen" else f_cat, "black")
    draw.text((left + plot_w / 2, h - 34), "Domaines de compétence", font=f_xtitle, fill="black", anchor="mm")
    return img


def docx_grid_chart(items):
    """Quadrant scatter with numbered markers, raster twin of graduatedGrid() / pdf_grid_chart()."""
    data = [d for d in items if d.get("capacity") is not None and d.get("consensus") is not None]
    w, h = 700, 620
    img = PILImage.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    f_title, f_axis, f_axis_title, f_quad, f_num, f_legend = docx_font(22, True), docx_font(14), docx_font(16, True), docx_font(13), docx_font(15, True), docx_font(13)
    draw.text((w / 2, 26), "Positionnement des domaines", font=f_title, fill="black", anchor="mm")
    if not data:
        draw.text((30, 60), "Pas encore de données disponibles pour la grille graduée.", font=f_axis, fill="black")
        return img
    left, right = 100, w - 40
    top, bottom = 60, h - 150
    plot_w, plot_h = right - left, bottom - top
    sx = lambda v: left + v / 100 * plot_w
    sy = lambda v: bottom - v / 100 * plot_h
    draw.rectangle([left, top, right, bottom], fill="#e9e9e9", outline="black")
    for v in (0, 20, 40, 60, 80, 100):
        draw.line([sx(v), top, sx(v), bottom], fill="#bbbbbb")
        draw.line([left, sy(v), right, sy(v)], fill="#bbbbbb")
        draw.text((sx(v), bottom + 16), str(v), font=f_axis, fill="black", anchor="mm")
        draw.text((left - 12, sy(v)), str(v), font=f_axis, fill="black", anchor="rm")
    draw.line([sx(50), top, sx(50), bottom], fill="#333333", width=2)
    draw.line([left, sy(50), right, sy(50)], fill="#333333", width=2)
    draw.text((left + plot_w / 2, bottom + 38), "Capacité", font=f_axis_title, fill="black", anchor="mm")
    docx_rotated_text(img, "Consensus", 22, top + plot_h / 2 + 45, 90, f_axis_title, "black")
    draw.text((left + 6, top + 8), "Faible capacité / consensus élevé", font=f_quad, fill="black")
    draw.text((right - 6 - docx_text_w(draw, "Capacité élevée / consensus élevé", f_quad), top + 8), "Capacité élevée / consensus élevé", font=f_quad, fill="black")
    draw.text((left + 6, bottom - 20), "Faible capacité / consensus faible", font=f_quad, fill="black")
    draw.text((right - 6 - docx_text_w(draw, "Capacité élevée / consensus faible", f_quad), bottom - 20), "Capacité élevée / consensus faible", font=f_quad, fill="black")
    avg_cap = sum(d["capacity"] for d in data) / len(data)
    avg_cons = sum(d["consensus"] for d in data) / len(data)
    order = sorted(range(len(data)), key=lambda i: data[i]["capacity"] + data[i]["consensus"])
    rank = {idx: r for r, idx in enumerate(order)}
    dirs = [(16, -14), (16, 18), (-16, -14), (-16, 18)]
    for i, d in enumerate(data):
        cx, cy = sx(d["capacity"]), sy(d["consensus"])
        r = 7
        draw.polygon([(cx, cy - r), (cx + r, cy), (cx, cy + r), (cx - r, cy)], fill="#000080")
        dx, dy = dirs[rank[i] % 4]
        if cx + dx > right - 16: dx = -abs(dx)
        elif cx + dx < left + 16: dx = abs(dx)
        if cy + dy < top + 12: dy = abs(dy)
        elif cy + dy > bottom - 12: dy = -abs(dy)
        draw.text((cx + dx, cy + dy), str(i + 1), font=f_num, fill="#000080", anchor="lm" if dx > 0 else "rm")
    ax, ay = sx(avg_cap), sy(avg_cons)
    label = "Moyenne générale"
    label_w = docx_text_w(draw, label, f_num)
    best = None
    for ddx, ddy, anchor in ((20, -16, "l"), (20, 22, "l"), (-20, -16, "r"), (-20, 22, "r")):
        lx, ly = ax + ddx, ay + ddy
        box_left = lx if anchor == "l" else lx - label_w
        box_right = box_left + label_w
        box_top, box_bottom = ly - 10, ly + 10
        in_bounds = box_left >= left + 2 and box_right <= right - 2 and box_top >= top + 2 and box_bottom <= bottom - 2
        min_dist = min(
            math.hypot(sx(d["capacity"]) - max(box_left, min(sx(d["capacity"]), box_right)),
                       sy(d["consensus"]) - max(box_top, min(sy(d["consensus"]), box_bottom)))
            for d in data
        )
        score = (in_bounds, min_dist)
        if best is None or score > best[0]:
            best = (score, lx, ly, anchor)
    _, lx, ly, anchor = best
    r = 7
    draw.polygon([(ax, ay - r), (ax + r, ay), (ax, ay + r), (ax - r, ay)], fill="white", outline="black", width=2)
    draw.text((lx, ly), label, font=f_num, fill="black", anchor="lm" if anchor == "l" else "rm")
    legend = " · ".join(f"{i + 1} {d['label']}" for i, d in enumerate(data))
    words, lines, cur = legend.split(" "), [], ""
    for word in words:
        trial = (cur + " " + word).strip()
        if not cur or docx_text_w(draw, trial, f_legend) <= w - 60:
            cur = trial
        else:
            lines.append(cur); cur = word
    if cur:
        lines.append(cur)
    ly2 = bottom + 74
    for line in lines:
        draw.text((30, ly2), line, font=f_legend, fill="#444444"); ly2 += 20
    return img


def docx_cohort_chart(items):
    """Standardised scores (bars) and graded scores (lines) across domains + the
    cohort mean as a 7th column, raster twin of the redesigned cohortChart()."""
    data = [d for d in items if d.get("capacity") is not None]
    w = max(900, 90 + (len(data) + 1) * 140) if data else 900
    h = 560
    img = PILImage.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    f_title, f_axis, f_legend, f_val, f_cat, f_cat_b, f_xtitle = docx_font(20, True), docx_font(15), docx_font(14), docx_font(14), docx_font(14), docx_font(14, True), docx_font(17, True)
    draw.text((w / 2, 24), "Notes standardisées et graduées — cohorte des domaines", font=f_title, fill="black", anchor="mm")
    if not data:
        draw.text((30, 60), "Pas encore de données disponibles pour l’analyse de cohorte.", font=f_axis, fill="black")
        return img
    mc = sum(d["capacity"] for d in data) / len(data)
    cons_data = [d for d in data if d.get("consensus") is not None]
    ms = sum(d["consensus"] for d in cons_data) / len(cons_data) if cons_data else None
    # Mirrors static/app.js cohort(): missing graded values count as 0, divided by
    # the full domain count (not just the domains that have a graded value).
    gc = sum(d.get("gradedCapacity") or 0 for d in data) / len(data)
    gs = sum(d.get("gradedConsensus") or 0 for d in data) / len(data)
    all_data = data + [{"label": "Moyen", "code": "Moyen", "capacity": mc, "consensus": ms, "gradedCapacity": gc, "gradedConsensus": gs}]
    left, right = 100, w - 30
    top, bottom = 108, h - 120
    plot_w, plot_h = right - left, bottom - top
    axis_min, axis_max = 20, 100
    scaleY = lambda v: bottom - (v - axis_min) / (axis_max - axis_min) * plot_h
    draw.rectangle([left, top, right, bottom], fill="#e9e9e9", outline="black")
    for v in (20, 40, 60, 80, 100):
        yy = scaleY(v)
        draw.line([left, yy, right, yy], fill="#bbbbbb")
        draw.text((left - 12, yy), str(v), font=f_axis, fill="black", anchor="rm")
    draw.rectangle([left, 56, left + 16, 72], fill="#9999FF")
    draw.text((left + 24, 64), "Capacité standardisée", font=f_legend, fill="black", anchor="lm")
    draw.rectangle([left + 240, 56, left + 256, 72], fill="#993366")
    draw.text((left + 264, 64), "Consensus standardisé", font=f_legend, fill="black", anchor="lm")
    draw.line([left, 86, left + 32, 86], fill="#800000", width=3)
    draw.text((left + 38, 86), "Capacité graduée", font=f_legend, fill="black", anchor="lm")
    draw.line([left + 220, 86, left + 252, 86], fill="#0000ff", width=3)
    draw.text((left + 258, 86), "Consensus gradué", font=f_legend, fill="black", anchor="lm")
    n = len(all_data); step = plot_w / n; bw = step * .32
    cap_pts, cons_pts = [], []
    for i, d in enumerate(all_data):
        x = left + i * step
        c_, s_ = d["capacity"], d["consensus"]
        cv, sv = max(axis_min, c_ or 0), max(axis_min, s_ or 0)
        y_cap, y_cons = scaleY(cv), scaleY(sv)
        draw.rectangle([x + step * .12, y_cap, x + step * .12 + bw, bottom], fill="#9999FF", outline="black")
        draw.rectangle([x + step * .12 + bw + 4, y_cons, x + step * .12 + bw + 4 + bw, bottom], fill="#993366", outline="black")
        lx = x + step * .12 + bw + 2
        docx_rotated_text(img, pdf_short_label(d), lx, bottom + 8, -40, f_cat_b if d["label"] == "Moyen" else f_cat, "black")
        if d.get("gradedCapacity") is not None:
            cap_pts.append((lx, scaleY(d["gradedCapacity"])))
        if d.get("gradedConsensus") is not None:
            cons_pts.append((lx, scaleY(d["gradedConsensus"])))
    if len(cap_pts) > 1:
        draw.line(cap_pts, fill="#800000", width=3)
    for px, py in cap_pts:
        draw.polygon([(px, py - 6), (px + 6, py + 5), (px - 6, py + 5)], fill="#800000")
    if len(cons_pts) > 1:
        draw.line(cons_pts, fill="#0000ff", width=3)
    for px, py in cons_pts:
        draw.line([px - 5, py - 5, px + 5, py + 5], fill="#0000ff", width=2)
        draw.line([px - 5, py + 5, px + 5, py - 5], fill="#0000ff", width=2)
    draw.text((left + plot_w / 2, h - 34), "Domaines", font=f_xtitle, fill="black", anchor="mm")
    return img


def docx_radar_chart(items):
    """Complementary synthesis radar, raster twin of radar() / pdf_radar_chart()."""
    data = [d for d in items if d.get("capacity") is not None]
    w, h = 560, 580
    img = PILImage.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    f_legend, f_label = docx_font(15), docx_font(13)
    if len(data) < 3:
        draw.text((30, 30), "Radar non disponible : au moins 3 domaines avec des données sont nécessaires.", font=f_legend, fill="black")
        return img
    n = len(data)
    cx, cy, r = w / 2, h / 2 + 16, 190
    draw.rectangle([20, 20, 34, 34], fill="#176b4b")
    draw.text((42, 27), "Capacité", font=f_legend, fill="black", anchor="lm")
    draw.rectangle([170, 20, 184, 34], fill="#536271")
    draw.text((192, 27), "Consensus", font=f_legend, fill="black", anchor="lm")
    pts_cap, pts_cons = [], []
    for i, d in enumerate(data):
        ang = -math.pi / 2 + i * 2 * math.pi / n
        ex, ey = cx + math.cos(ang) * r, cy + math.sin(ang) * r
        draw.line([cx, cy, ex, ey], fill="#aaaabb")
        lx, ly = cx + math.cos(ang) * (r + 20), cy + math.sin(ang) * (r + 20)
        draw.text((lx, ly), pdf_short_label(d)[:9], font=f_label, fill="black", anchor="mm")
        cv, sv = (d.get("capacity") or 0) / 100, (d.get("consensus") or 0) / 100
        pts_cap.append((cx + math.cos(ang) * r * cv, cy + math.sin(ang) * r * cv))
        pts_cons.append((cx + math.cos(ang) * r * sv, cy + math.sin(ang) * r * sv))
    draw.polygon(pts_cap, fill="#cfe3da", outline="#176b4b")
    draw.polygon(pts_cons, outline="#536271", width=2)
    return img


def docx_add_chart(doc, img):
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(6.3))


DOCX_NAVY = RGBColor(0x1B, 0x3A, 0x5C) if RGBColor else None
DOCX_MUTED = RGBColor(0x64, 0x74, 0x8B) if RGBColor else None


def docx_shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def docx_style_heading(heading, color=None):
    for run in heading.runs:
        run.font.color.rgb = color or DOCX_NAVY
    return heading


def docx_note(doc, text, size=9, color=None, italic=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.color.rgb = color or DOCX_MUTED
    return p


def docx_style_table(table, header_bg="1B3A5C"):
    table.style = "Table Grid"
    for cell in table.rows[0].cells:
        docx_shade_cell(cell, header_bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(table.rows[1:]):
        if i % 2 == 1:
            for cell in row.cells:
                docx_shade_cell(cell, "F2F5F8")
    return table


def docx_metrics_row(doc, metrics):
    """metrics: list of (label, value) tuples, rendered as a row of KPI cards like the web app."""
    t = doc.add_table(rows=2, cols=len(metrics))
    t.autofit = True
    for col, (label, value) in enumerate(metrics):
        vcell = t.cell(0, col)
        vcell.text = str(value)
        vp = vcell.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.runs[0]; vr.font.bold = True; vr.font.size = Pt(18); vr.font.color.rgb = DOCX_NAVY
        lcell = t.cell(1, col)
        lcell.text = label
        lp = lcell.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.runs[0]; lr.font.size = Pt(9); lr.font.color.rgb = DOCX_MUTED
        docx_shade_cell(vcell, "EAF1F9"); docx_shade_cell(lcell, "EAF1F9")
    return t


def pdf_level(v):
    """Mirrors static/app.js level(): qualitative reading of a capacity score."""
    if v is None: return "—"
    if v < 20: return ""
    if v <= 39: return "Loin en dessous de la moyenne"
    if v <= 59: return "En dessous de la moyenne"
    if v <= 70: return "Moyen"
    if v <= 80: return "Au-dessus de la moyenne"
    return "Bien au-dessus de la moyenne"


def pil_chart_grid(images, cols=2, pad=24, bg="white"):
    """Compose several chart images (as returned by the docx_*_chart helpers) into one
    dashboard image, so a report page shows several charts side by side instead of
    one chart per page. Each entry is either a PIL image, or an (image, caption) tuple
    whose caption is drawn above the chart (used for per-domain stat lines)."""
    items = [x if isinstance(x, tuple) else (x, None) for x in images if x is not None]
    if not items:
        return None
    cell_w = max(img.width for img, _ in items)
    cell_h = max(img.height for img, _ in items)
    caption_h = 34 if any(cap for _, cap in items) else 0
    rows = math.ceil(len(items) / cols)
    grid_w = cols * cell_w + pad * (cols + 1)
    grid_h = rows * (cell_h + caption_h) + pad * (rows + 1)
    canvas_img = PILImage.new("RGB", (grid_w, grid_h), bg)
    draw = ImageDraw.Draw(canvas_img)
    font = docx_font(20, True)
    for idx, (img, caption) in enumerate(items):
        r, col = divmod(idx, cols)
        x0 = pad + col * (cell_w + pad)
        y0 = pad + r * (cell_h + caption_h + pad)
        if caption:
            draw.text((x0 + cell_w / 2, y0 + caption_h / 2), caption, font=font, fill="black", anchor="mm")
        x = x0 + (cell_w - img.width) // 2
        y = y0 + caption_h + (cell_h - img.height) // 2
        canvas_img.paste(img, (x, y))
    return canvas_img


def xlsx_add_image(ws, row, col, img, scale=0.55):
    """Embed a PIL image in an xlsxwriter sheet at (row, col); returns the next free row."""
    if img is None:
        return row
    buf = BytesIO(); img.save(buf, format="PNG"); buf.seek(0)
    ws.insert_image(row, col, "chart.png", {"image_data": buf, "x_scale": scale, "y_scale": scale})
    return row + math.ceil(img.height * scale / 15) + 2


def report_docx(db, sid):
    """Mirrors the web app's 'Diagnostic terminé' page (the one 'Imprimer / Télécharger en
    PDF' prints) section for section: title/meta, KPI cards, Vue synthétique, Synthèse par
    domaine, Priorités retenues — same order, same headings, same table columns."""
    a = analysis(db, sid)
    if not a:
        raise ValueError("Session introuvable")
    session, g = a["session"], a["global"]
    domains = [d for d in a["domains"] if d.get("capacity") is not None]
    priorities = qualitative_data(db, sid)["priorities"]
    doc = Document()

    docx_style_heading(doc.add_heading("Diagnostic terminé", level=0))
    progress_target = session["expected_participants"] or a["participantCount"]
    docx_note(doc, f"{session['name']} · {a['participantCount']} commencés · {a['completedCount']} validés · progression {round(a['completedCount']/progress_target*100) if progress_target else 0}%", size=11, italic=False)
    doc.add_paragraph()
    docx_metrics_row(doc, [
        ("Capacité", pdf_fmt(g["capacity"])),
        ("Consensus", pdf_c(g)),
        ("Capacité graduée", pdf_fmt(g["gradedCapacity"])),
        ("Consensus gradué", pdf_fmt(g["gradedConsensus"])),
    ])

    profile_breakdown = participant_profile_breakdown(db, sid)
    if profile_breakdown:
        docx_style_heading(doc.add_heading("Profil des participants", level=2))
        pt = doc.add_table(rows=1, cols=3)
        for c, x in zip(pt.rows[0].cells, ["Champ", "Valeur", "N"]):
            c.text = x
        for field, value, n in profile_breakdown:
            row = pt.add_row().cells
            for c, v in zip(row, [field, value, n]):
                c.text = str(v)
        docx_style_table(pt)

    def add_findings_section():
        rows_ = findings_rows(a["findings"])
        if not rows_: return
        docx_style_heading(doc.add_heading("Forces / Fragilités / Points de vigilance", level=2))
        docx_note(doc, "Constats chiffrés issus des données ci-dessus — jamais présentés comme des causes.", italic=True)
        ft = doc.add_table(rows=1, cols=5)
        for c, x in zip(ft.rows[0].cells, ["Type", "Niveau", "Libellé", "Capacité", "Consensus"]):
            c.text = x
        for row_data in rows_:
            row = ft.add_row().cells
            for c, v in zip(row, row_data):
                c.text = "" if v is None else str(v)
        docx_style_table(ft)

    if not domains or PILImage is None:
        docx_style_heading(doc.add_heading("Synthèse par domaine", level=2))
        t = doc.add_table(rows=1, cols=6)
        for c, x in zip(t.rows[0].cells, ["Domaine", "Capacité", "Consensus", "Graduées", "Niveau", "Réponses"]):
            c.text = x
        for d in domains:
            row = t.add_row().cells
            graded = f"{d['gradedCapacity'] if d['gradedCapacity'] is not None else '—'} / {d['gradedConsensus'] if d['gradedConsensus'] is not None else '—'}"
            for c, v in zip(row, [d["label"], pdf_fmt(d["capacity"]), pdf_c(d), graded, pdf_level(d["capacity"]), d["responses"]]):
                c.text = str(v)
        docx_style_table(t)
        if not domains:
            doc.add_paragraph("Pas encore de données suffisantes pour la restitution graphique EPC.")
        elif PILImage is None:
            doc.add_paragraph("Graphiques indisponibles : le paquet optionnel Pillow n'est pas installé (voir README.md).")
        add_findings_section()
        docx_style_heading(doc.add_heading("Priorités retenues", level=2))
        doc.add_paragraph(f"{len(priorities)} priorité(s) sélectionnée(s)." if priorities else "Aucune priorité sélectionnée.")
        out = BytesIO(); doc.save(out); return out.getvalue()

    doc.add_page_break()
    docx_style_heading(doc.add_heading("Vue synthétique", level=2))
    docx_note(doc, "Radar global : vue complémentaire, ne remplace pas les graphiques EPC ci-dessous.")
    docx_add_chart(doc, pil_chart_grid([
        docx_radar_chart(domains),
        docx_bars_chart(domains, "standard", "Notes standardisées par domaine", with_mean=True),
        docx_bars_chart(domains, "graded", "Notes graduées par domaine", with_mean=True),
        docx_grid_chart(domains),
    ]))
    if len(domains) > 1:
        docx_add_chart(doc, docx_cohort_chart(domains))
    docx_note(doc, "20–39 : Loin en dessous · 40–59 : En dessous · 60–70 : Moyen · 71–80 : Au-dessus · 81–100 : Bien au-dessus", size=9, color=RGBColor(0x18, 0x6E, 0x42))

    doc.add_page_break()
    docx_style_heading(doc.add_heading("Synthèse par domaine", level=2))
    t = doc.add_table(rows=1, cols=6)
    for c, x in zip(t.rows[0].cells, ["Domaine", "Capacité", "Consensus", "Graduées", "Niveau", "Réponses"]):
        c.text = x
    for d in domains:
        row = t.add_row().cells
        graded = f"{d['gradedCapacity'] if d['gradedCapacity'] is not None else '—'} / {d['gradedConsensus'] if d['gradedConsensus'] is not None else '—'}"
        for c, v in zip(row, [d["label"], pdf_fmt(d["capacity"]), pdf_c(d), graded, pdf_level(d["capacity"]), d["responses"]]):
            c.text = str(v)
    docx_style_table(t)

    add_findings_section()

    docx_style_heading(doc.add_heading("Priorités retenues", level=2))
    doc.add_paragraph(f"{len(priorities)} priorité(s) sélectionnée(s)." if priorities else "Aucune priorité sélectionnée.")

    ai_blocks = rows(db, "SELECT section_key,content FROM report_ai_blocks WHERE session_id=?", (sid,))
    if ai_blocks:
        doc.add_page_break()
        docx_style_heading(doc.add_heading("Synthèse assistée par IA", level=2))
        docx_note(doc, "Propositions rédigées avec l'aide de l'assistant IA et explicitement retenues par le modérateur. "
            "Les données, scores et graphiques EPC ci-dessus restent la source primaire du diagnostic.", italic=False)
        for key in AI_SECTION_LABELS:
            block = next((b for b in ai_blocks if b["section_key"] == key), None)
            if not block: continue
            docx_style_heading(doc.add_heading(AI_SECTION_LABELS[key], level=3))
            doc.add_paragraph(block["content"])

    out = BytesIO(); doc.save(out); return out.getvalue()

PDF_STOPWORDS = {"de", "des", "du", "la", "le", "les", "et", "en", "au", "aux", "à", "a", "l"}


def pdf_fmt(v):
    return "—" if v is None else f"{v:.1f}"


def pdf_c(obj):
    return "Non calculable" if obj.get("consensusNote") == "single_respondent" else pdf_fmt(obj.get("consensus"))


def pdf_short_label(item):
    """Mirrors static/app.js shortLabel(): configured code if it reads as a short
    alphabetic label, otherwise an abbreviation generated from the item's own name."""
    code = (item.get("code") or "").strip()
    if re.fullmatch(r"[A-Za-zÀ-ÿ]{2,10}", code):
        return code.upper()
    if code:
        # Référence courte d'un indicateur (ex. "Formation au personnel") : utilisée
        # telle quelle en abscisse (mission :8810, demande Mouhamed BA), au lieu de
        # l'abréviation calculée ci-dessous à partir de l'énoncé complet.
        return code
    label = (item.get("label") or "").strip()
    words = [w for w in label.split() if w and w.lower().replace("’", "").replace("'", "") not in PDF_STOPWORDS]
    if not words:
        words = label.split()
    if len(words) >= 2:
        return "".join(w[0].upper() for w in words[:4])
    return label[:4].upper()



def read_xlsx(raw):
    ns={"x":"http://schemas.openxmlformats.org/spreadsheetml/2006/main","r":"http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
    try: z=zipfile.ZipFile(BytesIO(raw)); wb=ET.fromstring(z.read("xl/workbook.xml")); rels=ET.fromstring(z.read("xl/_rels/workbook.xml.rels"))
    except Exception as e: raise ValueError("Le fichier n'est pas un XLSX valide") from e
    targets={r.attrib["Id"]:r.attrib["Target"].lstrip("/") for r in rels}; shared=[]
    if "xl/sharedStrings.xml" in z.namelist(): shared=["".join(x.itertext()) for x in ET.fromstring(z.read("xl/sharedStrings.xml")).findall("x:si",ns)]
    sheets={}
    for sh in wb.findall("x:sheets/x:sheet",ns):
        rid=sh.attrib["{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"]; target=targets[rid]; target="xl/"+target if not target.startswith("xl/") else target
        grid=[]
        for row in ET.fromstring(z.read(target)).findall(".//x:sheetData/x:row",ns):
            vals=[]
            for c in row.findall("x:c",ns):
                v=c.find("x:v",ns); text="" if v is None else v.text or ""
                if c.attrib.get("t")=="s": text=shared[int(text)]
                elif c.attrib.get("t")=="inlineStr": text="".join(c.itertext())
                vals.append(text)
            grid.append(vals)
        sheets[sh.attrib["name"].upper()]=grid
    return sheets


def import_preview(raw):
    sheets=read_xlsx(raw); errors=[]
    if "QUESTIONNAIRE" in sheets:
        q=sheets["QUESTIONNAIRE"]; p=sheets.get("PARAMETRES",[])
        if not q or q[0][:3] != ["Domaine","Référence","Indicateur qualitatif ou Capacité"]: raise ValueError("Colonnes QUESTIONNAIRE invalides")
        values={r[0]:r[1] if len(r)>1 else "" for r in p[:2] if r}; labels={}
        for r in p[4:]:
            if len(r)>1 and r[0]: labels[str(r[0])]=r[1]
        level_nums=sorted(int(k) for k in labels if str(k).lstrip('-').isdigit())
        smin,smax=(level_nums[0],level_nums[-1]) if level_nums else (1,5)
        grouped={}; order=[]
        for r in q[1:]:
            r=(r+["", "", ""])[:3]
            if not any(r) or r[0].startswith("EXEMPLE"): continue
            if not all(r): errors.append("Chaque ligne doit contenir Domaine, Référence et Indicateur qualitatif ou Capacité"); continue
            if r[0] not in grouped: grouped[r[0]]=[]; order.append(r[0])
            grouped[r[0]].append({"code":"","label":r[1],"description":r[2],"response_type":"numeric","required":True,"active":True,"display_order":len(grouped[r[0]])})
        if not values.get("Nom du questionnaire","").strip(): errors.append("Nom du questionnaire obligatoire")
        return {"errors":errors,"template":{"name":values.get("Nom du questionnaire",""),"description":values.get("Description",""),"scale":{"type":"numeric","min":smin,"max":smax,"labels":labels},"priority":{"maxPerDomain":3},"domains":[{"label":d,"display_order":n+1,"indicators":grouped[d]} for n,d in enumerate(order)]},"rows":sum(len(v) for v in grouped.values())}
    if "INDICATEURS" not in sheets or "PARAMETRES" not in sheets: raise ValueError("Les feuilles INDICATEURS et PARAMETRES sont obligatoires")
    rows_i=sheets["INDICATEURS"]; header_index=next((n for n,r in enumerate(rows_i) if r[:len(MATRIX_COLUMNS)]==MATRIX_COLUMNS),None)
    if header_index is None: errors.append("Colonnes INDICATEURS invalides ou dans un ordre incorrect"); header_index=len(rows_i)
    params={r[0]:r[1] if len(r)>1 else "" for r in sheets["PARAMETRES"] if r and r[0] in PARAMETERS}
    for p in PARAMETERS:
        if p not in params: errors.append(f"Paramètre manquant : {p}")
    domains={}
    for n,r in enumerate(rows_i[header_index+1:],header_index+2):
        r=(r+[""]*9)[:9]; dom,do,code,label,desc,io,typ,required,active=r
        if not any(r): continue
        if code == "EXEMPLE-01": continue
        if not dom or not code or not label: errors.append(f"Ligne {n}: Domaine, Code indicateur et Indicateur sont obligatoires"); continue
        try: do=int(float(do)); io=int(float(io))
        except: errors.append(f"Ligne {n}: les ordres doivent être des nombres entiers"); continue
        if typ not in ("numeric","text","boolean"): errors.append(f"Ligne {n}: Type réponse invalide ({typ})"); continue
        def flag(v): return str(v).strip().lower() in ("oui","true","1","yes")
        key=(dom,do); domains.setdefault(key,[]).append({"code":code,"label":label,"description":desc,"response_type":typ,"required":flag(required),"active":flag(active),"display_order":io})
    for key, inds in domains.items():
        if len({i["code"] for i in inds}) != len(inds): errors.append(f"Domaine {key[0]}: codes indicateurs dupliqués")
        if len({i["display_order"] for i in inds}) != len(inds): errors.append(f"Domaine {key[0]}: ordres indicateurs dupliqués")
    try: lo=float(params.get("Valeur minimum",1)); hi=float(params.get("Valeur maximum",5)); assert lo<=hi
    except: errors.append("Bornes d'échelle invalides")
    if not params.get("Nom questionnaire","").strip(): errors.append("Nom questionnaire obligatoire")
    return {"errors":errors,"template":{"name":params.get("Nom questionnaire",""),"description":params.get("Description",""),"scale":{"type":params.get("Type d'échelle","numeric"),"min":lo if not errors else 1,"max":hi if not errors else 5,"labels":json.loads(params.get("Libellés des valeurs") or "{}")},"priority":{"maxPerDomain":int(float(params.get("Nombre de priorités par domaine",3) or 3))},"domains":[{"label":k[0],"display_order":k[1],"indicators":sorted(v,key=lambda x:x["display_order"])} for k,v in sorted(domains.items(),key=lambda x:x[0][1])]},"rows":sum(len(x) for x in domains.values())}


def save_import(db, data, owner_user_id=None):
    if data["errors"]: raise ValueError("La matrice comporte des erreurs")
    tid=create_blank_template(db,data["template"],owner_user_id); t=data["template"]; db.execute("UPDATE templates SET scale_json=?,priority_json=? WHERE id=?",(json.dumps(t["scale"]),json.dumps(t["priority"]),tid))
    for d in t["domains"]:
        did=str(uuid.uuid4()); code="domain-"+uuid.uuid4().hex[:8]; db.execute("INSERT INTO domains VALUES (?,?,?,?,?,?,?)",(did,tid,code,d["label"],"",d["display_order"],1))
        for i in d["indicators"]: db.execute("INSERT INTO indicators VALUES (?,?,?,?,?,?,?,?,?,?)",(str(uuid.uuid4()),did,i["code"],i["label"],i["description"],i["response_type"],int(i["required"]),i["display_order"],int(i["active"]),"{}"))
    db.commit(); return tid


def grade(value, norm):
    """Classify value into a graduated band from norm (low, high, result) tuples.

    Matches the reference KOICA calculation tool: the continuous standardized
    score is rounded to the nearest whole number first, then looked up against
    the bands' own authored integer bounds (which are contiguous once rounded,
    e.g. ...59-63, 64-67, 68-71... — no gap-filling needed post-rounding).
    """
    if value is None:
        return None
    v = round(max(0.0, min(100.0, float(value))))
    for low, high, result in norm:
        if low <= v <= high:
            return result
    return None


# Query-param key (camelCase, matches PARTICIPANT_PROFILE_FIELDS in static/app.js) -> participants column.
PARTICIPANT_FILTER_FIELDS = {"participantType": "participant_type", "profile": "profile", "sex": "sex", "ageRange": "age_range", "educationLevel": "education_level"}

FORCE_THRESHOLD = 71  # matches the "Au-dessus de la moyenne" band floor already used by level() (static/app.js)
FRAGILE_THRESHOLD = 60  # matches the "Moyen" band floor
VIGILANCE_GAP_THRESHOLD = 15  # capacity-point gap between sub-populations flagged as a point de vigilance


def _participant_filter_sql(participant_filter):
    """Build an ' AND p.<col>=? AND ...' fragment (+ matching params) from a
    {db_column: value} dict. Empty/None input returns ("", []) so callers get
    byte-identical SQL to the unfiltered path — this is what guarantees "Tous
    les participants" stays strictly unchanged."""
    if not participant_filter:
        return "", []
    frag = "".join(f" AND p.{col}=?" for col in participant_filter)
    return frag, list(participant_filter.values())


def objective_findings(result, comparison=None):
    """Deterministic, non-AI findings from an already-computed analysis() result:
    forces (highest capacity), fragilités (lowest capacity), points de vigilance
    (capacity/consensus disagree, or — when `comparison` categories are given — a
    large capacity gap between sub-populations on the same domain). Every item
    carries the numbers that justify it; this is a reading of the data, never a
    claimed cause."""
    domains = [d for d in result["domains"] if d["capacity"] is not None]
    indicators = [dict(i, domain=d["label"]) for d in result["domains"] for i in d["indicators"] if i["capacity"] is not None]

    forces_domains = sorted([d for d in domains if d["capacity"] >= FORCE_THRESHOLD], key=lambda d: -d["capacity"])[:3]
    forces_indicators = sorted([i for i in indicators if i["capacity"] >= FORCE_THRESHOLD], key=lambda i: -i["capacity"])[:5]
    fragile_domains = sorted([d for d in domains if d["capacity"] < FRAGILE_THRESHOLD], key=lambda d: d["capacity"])[:3]
    fragile_indicators = sorted([i for i in indicators if i["capacity"] < FRAGILE_THRESHOLD], key=lambda i: i["capacity"])[:5]

    vigilance = []
    for d in domains:
        if d["consensus"] is None:
            continue
        if d["capacity"] >= FORCE_THRESHOLD and d["consensus"] < FRAGILE_THRESHOLD:
            vigilance.append({"level": "domain", "id": d["id"], "label": d["label"], "capacity": d["capacity"], "consensus": d["consensus"], "reason": "capacite_elevee_consensus_faible"})
        elif d["capacity"] < FRAGILE_THRESHOLD and d["consensus"] >= FORCE_THRESHOLD:
            vigilance.append({"level": "domain", "id": d["id"], "label": d["label"], "capacity": d["capacity"], "consensus": d["consensus"], "reason": "capacite_faible_consensus_eleve"})

    if comparison:
        by_domain = {}
        for cat in comparison:
            for d in cat.get("domains", []):
                if d.get("capacity") is None:
                    continue
                by_domain.setdefault(d["id"], {"label": d["label"], "values": []})["values"].append({"category": cat["value"], "capacity": d["capacity"]})
        for did, info in by_domain.items():
            caps = [v["capacity"] for v in info["values"]]
            if len(caps) < 2:
                continue
            gap = max(caps) - min(caps)
            if gap >= VIGILANCE_GAP_THRESHOLD:
                vigilance.append({"level": "domain", "id": did, "label": info["label"], "reason": "ecart_sous_populations", "gap": gap, "values": info["values"]})

    return {"forces": {"domains": forces_domains, "indicators": forces_indicators}, "fragilites": {"domains": fragile_domains, "indicators": fragile_indicators}, "vigilance": vigilance}


def analysis(db, session_id: str, participant_filter=None):
    """Capacité/consensus are computed only from participants with status='completed'
    (a questionnaire opened but abandoned mid-way must never silently shift the
    published score) — participantCount below still counts every participant row
    (started or completed) so "commencés" stays visible separately from
    "validés"/completedCount.

    participant_filter: optional {participants_column: value} dict (e.g.
    {"sex": "Homme"}) to scope the computation to a sub-population — always
    recomputed from the matching individual completed responses, never a
    filter applied after aggregation. None/empty means every participant, and
    produces byte-identical SQL to the pre-filter version of this function."""
    session = db.execute("SELECT * FROM sessions WHERE id=?", (session_id,)).fetchone()
    if not session:
        return None
    template = template_payload(db, session["template_id"])
    scale, rules, consensus, norm = template["scale"], template["scoring"], template["consensus"], template["grading"]
    low, high, amplitude = float(scale["min"]), float(scale["max"]), float(scale["max"] - scale["min"])
    output_max = float(rules.get("outputRange", [0, 100])[1])
    filt_sql, filt_params = _participant_filter_sql(participant_filter)
    all_values, output_domains, all_participant_ids = [], [], set()
    for domain in template["domains"]:
        indicators = [i for i in domain["indicators"] if i["active"]]
        output_indicators, participant_means = [], {}
        for indicator in indicators:
            response_rows = rows(db, "SELECT r.participant_id,r.value_json FROM responses r JOIN participants p ON p.id=r.participant_id WHERE r.session_id=? AND r.indicator_id=? AND p.status='completed'" + filt_sql, (session_id, indicator["id"], *filt_params))
            values = [float(json.loads(r["value_json"])) for r in response_rows if isinstance(json.loads(r["value_json"]), (int, float))]
            for r in response_rows:
                value = json.loads(r["value_json"])
                if isinstance(value, (int, float)):
                    participant_means.setdefault(r["participant_id"], []).append(float(value))
                    all_participant_ids.add(r["participant_id"])
            mean = sum(values) / len(values) if values else None
            sd = (sum((v - mean) ** 2 for v in values) / (len(values) - 1)) ** .5 if len(values) > 1 else None
            capacity = mean / high * output_max if mean is not None else None
            cons = max(0, output_max - consensus["factor"] * (sd / amplitude * output_max)) if sd is not None else None
            missing_count = len(rows(db, "SELECT id FROM participants p WHERE p.session_id=?" + filt_sql, (session_id, *filt_params)))
            output_indicators.append({"id": indicator["id"], "code": indicator["code"], "label": indicator["label"], "responses": len(values), "missing": max(0, missing_count - len(values)), "mean": mean, "capacity": capacity, "dispersion": sd, "consensus": cons, "consensusNote": "single_respondent" if len(values) == 1 else None, "gradedCapacity": grade(capacity, norm) if capacity is not None else None, "gradedConsensus": grade(cons, norm) if cons is not None else None, "distribution": {str(k): values.count(k) for k in range(int(low), int(high) + 1)}})
            all_values.extend(values)
        person_scores = [sum(values) / len(values) for values in participant_means.values() if values]
        dmean = sum(person_scores) / len(person_scores) if person_scores else None
        dsd = (sum((v - dmean) ** 2 for v in person_scores) / (len(person_scores) - 1)) ** .5 if len(person_scores) > 1 else None
        dcap = dmean / high * output_max if dmean is not None else None
        dcons = max(0, output_max - consensus["factor"] * (dsd / amplitude * output_max)) if dsd is not None else None
        output_domains.append({"id": domain["id"], "code": domain["code"], "label": domain["label"], "responses": len(person_scores), "capacity": dcap, "dispersion": dsd, "consensus": dcons, "consensusNote": "single_respondent" if len(person_scores) == 1 else None, "gradedCapacity": grade(dcap, norm) if dcap is not None else None, "gradedConsensus": grade(dcons, norm) if dcons is not None else None, "indicators": output_indicators})
    # Mirrors the reference KOICA tool's "Moyenne" row: global capacity/consensus
    # (standardized and graduated alike) are unweighted averages across domains,
    # not a response-weighted pool of every individual answer. The graduated
    # global score in particular averages the domains' own graduated scores —
    # it is not grade() applied to the averaged standardized score.
    domain_caps=[d["capacity"] for d in output_domains if d["capacity"] is not None]; domain_cons=[d["consensus"] for d in output_domains if d["consensus"] is not None]
    global_capacity=sum(domain_caps)/len(domain_caps) if domain_caps else None
    gc=sum(domain_cons)/len(domain_cons) if domain_cons else None
    graded_caps=[d["gradedCapacity"] for d in output_domains if d["gradedCapacity"] is not None]; graded_cons=[d["gradedConsensus"] for d in output_domains if d["gradedConsensus"] is not None]
    global_graded_capacity=sum(graded_caps)/len(graded_caps) if graded_caps else None
    global_graded_consensus=sum(graded_cons)/len(graded_cons) if graded_cons else None
    participants=len(rows(db,"SELECT id FROM participants p WHERE p.session_id=?"+filt_sql,(session_id,*filt_params))); completed=len(rows(db,"SELECT id FROM participants p WHERE p.session_id=? AND p.status='completed'"+filt_sql,(session_id,*filt_params)))
    result = {"session": dict(session), "participantCount": participants, "completedCount":completed, "domains": output_domains, "global": {"responses": len(all_values), "capacity": global_capacity, "consensus":gc, "consensusNote": "single_respondent" if len(all_participant_ids) == 1 else None, "gradedCapacity": global_graded_capacity, "gradedConsensus": global_graded_consensus}}
    result["findings"] = objective_findings(result)
    # Only computed for the unfiltered ("Tous") call — a sub-population's own profile
    # breakdown would trivially collapse to ~100% one category, which isn't useful,
    # and this avoids extra queries on every filtered/compare_population() call.
    result["participantProfile"] = participant_profile_breakdown(db, session_id) if not participant_filter else None
    return result


def compare_population(db, session_id: str, axis_col: str):
    """One analysis() call per distinct value of `axis_col` actually present
    among this session's validated participants — categories with 0 respondents
    are simply absent (never shown at N=0), per the consignes. Each call reuses
    analysis()'s own math untouched, so this can never drift from the single-
    population numbers."""
    values = [r["v"] for r in rows(db, f"SELECT DISTINCT {axis_col} AS v FROM participants WHERE session_id=? AND status='completed' AND {axis_col} IS NOT NULL AND {axis_col}!=''", (session_id,))]
    categories = []
    for value in values:
        a = analysis(db, session_id, {axis_col: value})
        categories.append({"value": value, "N": a["completedCount"], "capacity": a["global"]["capacity"], "consensus": a["global"]["consensus"], "consensusNote": a["global"]["consensusNote"], "gradedCapacity": a["global"]["gradedCapacity"], "gradedConsensus": a["global"]["gradedConsensus"], "domains": [{"id": d["id"], "label": d["label"], "capacity": d["capacity"], "consensus": d["consensus"]} for d in a["domains"]]})
    return categories


def qualitative_data(db, session_id: str):
    """Qualitative workshop data kept separate from EPC numerical results."""
    priorities = rows(db, """SELECT p.*,d.label AS domain_label,i.code AS indicator_code,
        i.label AS indicator_label,i.description AS indicator_description
        FROM priorities p JOIN domains d ON d.id=p.domain_id JOIN indicators i ON i.id=p.indicator_id
        WHERE p.session_id=? ORDER BY d.display_order,i.display_order""", (session_id,))
    return {
        "priorities": priorities,
        "analyses": rows(db, "SELECT * FROM priority_analyses WHERE session_id=?", (session_id,)),
        "entries": rows(db, "SELECT * FROM analysis_entries WHERE session_id=? ORDER BY created_at", (session_id,)),
        "recommendations": rows(db, "SELECT * FROM workshop_recommendations WHERE session_id=? ORDER BY created_at", (session_id,)),
        "trainingTopics": rows(db, "SELECT * FROM training_topics WHERE session_id=? ORDER BY created_at", (session_id,)),
    }

def report_data(db, session_id: str):
    a=analysis(db,session_id)
    if not a: return None
    meta=db.execute("SELECT * FROM session_report_meta WHERE session_id=?",(session_id,)).fetchone()
    return {"analysis":a,"template":template_payload(db,a["session"]["template_id"]),"qualitative":qualitative_data(db,session_id),"meta":dict(meta) if meta else {"facilitator":"","audience":"","context":"","conclusion":""}}


class Handler(SimpleHTTPRequestHandler):
    def db(self):
        db = connect(); init_schema(db); return db

    def json(self, code, payload, cookie=None):
        raw = json.dumps(payload, ensure_ascii=False).encode()
        self.send_response(code); self.send_header("Content-Type", "application/json; charset=utf-8"); self.send_header("Content-Length", str(len(raw)))
        if cookie is not None: self.send_header("Set-Cookie", cookie)
        self.end_headers(); self.wfile.write(raw)

    def body(self):
        size = int(self.headers.get("Content-Length", 0)); return json.loads(self.rfile.read(size) or b"{}")

    def raw_body(self): return self.rfile.read(int(self.headers.get("Content-Length", 0)))

    def current_user(self, db):
        raw = self.headers.get("Cookie")
        if not raw: return None
        jar = SimpleCookie(); jar.load(raw)
        morsel = jar.get("epc_session")
        if not morsel: return None
        token_hash = hashlib.sha256(morsel.value.encode("utf-8")).hexdigest()
        row = db.execute("SELECT u.* FROM auth_tokens t JOIN users u ON u.id=t.user_id WHERE t.token_hash=? AND t.expires_at>?", (token_hash, now())).fetchone()
        return dict(row) if row else None

    def check_ownership(self, path, db, user):
        parts = path.split("/")
        if path.startswith("/api/sessions/") and len(parts) > 3 and parts[3]:
            row = db.execute("SELECT owner_user_id FROM sessions WHERE id=?", (parts[3],)).fetchone()
            if row and user["role"] != "admin" and row["owner_user_id"] not in (None, user["id"]):
                raise PermissionDeniedError()
        elif path.startswith("/api/templates/") and len(parts) > 3 and parts[3] not in ("matrix.xlsx", "import"):
            row = db.execute("SELECT owner_user_id FROM templates WHERE id=?", (parts[3],)).fetchone()
            if row and row["owner_user_id"] is not None and user["role"] != "admin" and row["owner_user_id"] != user["id"]:
                raise PermissionDeniedError()

    def require_auth(self, path, db):
        """Call first inside each verb handler's try block. Returns the current
        user (or None for the small public whitelist) and enforces per-row
        ownership for /api/sessions|templates/<id>... routes."""
        user = self.current_user(db)
        if path.startswith("/api/") and not is_public_api(path, self.command):
            if user is None: raise AuthRequiredError()
            self.check_ownership(path, db, user)
        return user

    def do_GET(self):
        path, query = urlparse(self.path).path, parse_qs(urlparse(self.path).query)
        db = self.db()
        try:
            user = self.require_auth(path, db)
            if path == "/api/auth/setup-status":
                return self.json(200, {"needsSetup": db.execute("SELECT 1 FROM users LIMIT 1").fetchone() is None})
            if path == "/api/auth/me":
                return self.json(200, {"user": {"id": user["id"], "email": user["email"], "role": user["role"], "displayName": user["display_name"]} if user else None})
            if path == "/api/templates":
                if user["role"] == "admin": tpl_rows = rows(db, "SELECT id,name,version,description,status,created_at,updated_at,is_canonical FROM templates WHERE status='active' ORDER BY is_canonical DESC,name,version DESC")
                else: tpl_rows = rows(db, "SELECT id,name,version,description,status,created_at,updated_at,is_canonical FROM templates WHERE status='active' AND (owner_user_id IS NULL OR owner_user_id=?) ORDER BY is_canonical DESC,name,version DESC", (user["id"],))
                # domain_count/indicator_count : nombre de domaines/indicateurs ACTIFS,
                # calculé depuis les données réelles pour que l'accueil et "Voir tous
                # les modèles" affichent la structure de chaque questionnaire (ex :
                # "7 domaines x 10 indicateurs") sans rien coder en dur (mission :8810).
                for t in tpl_rows:
                    t["domain_count"], t["indicator_count"] = template_domain_indicator_counts(db, t["id"])
                return self.json(200, tpl_rows)
            if path == "/api/templates/matrix.xlsx":
                data=blank_matrix_xlsx(); name=export_filename("matrice-questionnaire-vierge", ext="xlsx"); self.send_response(200); self.send_header("Content-Type","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"); self.send_header("Content-Disposition",f"attachment; filename={name}"); self.send_header("Content-Length",str(len(data))); self.end_headers(); self.wfile.write(data); return
            if path.startswith("/api/templates/") and path.endswith("/matrix.xlsx"):
                template=template_payload(db,path.split("/")[3]); data=matrix_xlsx(template); name=export_filename(template["name"],"matrice-questionnaire", ext="xlsx"); self.send_response(200); self.send_header("Content-Type","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"); self.send_header("Content-Disposition",f"attachment; filename={name}"); self.send_header("Content-Length",str(len(data))); self.end_headers(); self.wfile.write(data); return
            if path.startswith("/api/templates/"): return self.json(200, template_payload(db, path.rsplit("/", 1)[1]) or {"error": "Configuration introuvable"})
            if path == "/api/sessions":
                if user["role"] == "admin": return self.json(200, rows(db, "SELECT * FROM sessions ORDER BY created_at DESC"))
                return self.json(200, rows(db, "SELECT * FROM sessions WHERE owner_user_id=? ORDER BY created_at DESC", (user["id"],)))
            if path.startswith("/api/sessions/") and path.endswith("/analysis"):
                pf={col: query[qp][0] for qp,col in PARTICIPANT_FILTER_FIELDS.items() if query.get(qp) and query[qp][0]}
                result = analysis(db, path.split("/")[3], pf or None); return self.json(200, result or {"error": "Session introuvable"})
            if path.startswith("/api/sessions/") and path.endswith("/compare"):
                sid=path.split("/")[3]; axis=query.get("axis",[""])[0]
                if axis not in PARTICIPANT_FILTER_FIELDS: return self.json(400,{"error":"Axe de comparaison invalide."})
                categories=compare_population(db,sid,PARTICIPANT_FILTER_FIELDS[axis])
                return self.json(200,{"axis":axis,"categories":categories,"findings":objective_findings(analysis(db,sid),comparison=categories)})
            if path.startswith("/api/sessions/") and path.endswith("/workshop-data"):
                sid=path.split("/")[3]; return self.json(200,{"priorities":rows(db,"SELECT * FROM priorities WHERE session_id=?",(sid,)),"notes":rows(db,"SELECT * FROM analysis_notes WHERE session_id=?",(sid,)),"recommendations":rows(db,"SELECT * FROM recommendations WHERE session_id=?",(sid,))})
            if path.startswith("/api/sessions/") and path.endswith("/qualitative-data"):
                return self.json(200, qualitative_data(db, path.split("/")[3]))
            if path.startswith("/api/sessions/") and path.endswith("/report-data"):
                return self.json(200, report_data(db, path.split("/")[3]) or {"error":"Session introuvable"})
            if path.startswith("/api/sessions/") and path.endswith("/export.json"):
                sid = path.split("/")[3]; return self.json(200, {"report":report_data(db,sid), "responses": rows(db, "SELECT * FROM responses WHERE session_id=?", (sid,)), "analysisNotes": rows(db, "SELECT * FROM analysis_notes WHERE session_id=?", (sid,))})
            if path.startswith("/api/sessions/") and path.endswith("/report.xlsx"):
                sid=path.split("/")[3];data=report_xlsx(db,sid);mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet';name=export_filename(session_label(db,sid),"diagnostic",ext="xlsx")
            elif path.startswith("/api/sessions/") and path.endswith("/report.docx"):
                sid=path.split("/")[3];data=report_docx(db,sid);mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document';name=export_filename(session_label(db,sid),"rapport",ext="docx")
            elif path.startswith("/api/sessions/") and path.endswith("/individual-responses.xlsx"):
                sid=path.split("/")[3];data=individual_responses_xlsx(db,sid);mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet';name=export_filename(session_label(db,sid),"reponses-individuelles",ext="xlsx")
            elif path.startswith("/api/sessions/") and path.endswith("/filtered-analysis.xlsx"):
                sid=path.split("/")[3];pf={col: query[qp][0] for qp,col in PARTICIPANT_FILTER_FIELDS.items() if query.get(qp) and query[qp][0]};data=filtered_analysis_xlsx(db,sid,pf or None);mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet';name=export_filename(session_label(db,sid),"analyse-filtree",ext="xlsx")
            else: data=None
            if data is not None:
                self.send_response(200);self.send_header('Content-Type',mime);self.send_header('Content-Disposition',f'attachment; filename={name}');self.send_header('Content-Length',str(len(data)));self.send_header('Cache-Control','no-store');self.end_headers();self.wfile.write(data);return
            if path.startswith("/api/sessions/") and path.endswith("/responses.csv"):
                sid = path.split("/")[3]; buf = StringIO(); writer = csv.writer(buf); writer.writerow(["participant", "nom / organisation", "indicator", "value", "updated_at"])
                for r in db.execute("SELECT p.anonymous_id,p.display_name,i.code,r.value_json,r.updated_at FROM responses r JOIN participants p ON p.id=r.participant_id JOIN indicators i ON i.id=r.indicator_id WHERE r.session_id=?", (sid,)): writer.writerow(r)
                data = buf.getvalue().encode(); name=export_filename(session_label(db,sid),"reponses",ext="csv"); self.send_response(200); self.send_header("Content-Type", "text/csv; charset=utf-8"); self.send_header("Content-Disposition", f"attachment; filename={name}"); self.end_headers(); self.wfile.write(data); return
            if path.startswith("/api/sessions/") and path.endswith("/individual-responses.csv"):
                sid = path.split("/")[3]; data = individual_responses_csv(db, sid); name=export_filename(session_label(db,sid),"reponses-individuelles",ext="csv"); self.send_response(200); self.send_header("Content-Type", "text/csv; charset=utf-8"); self.send_header("Content-Disposition", f"attachment; filename={name}"); self.end_headers(); self.wfile.write(data); return
            if path.startswith("/api/sessions/") and path.endswith("/filtered-analysis.csv"):
                sid = path.split("/")[3]; pf={col: query[qp][0] for qp,col in PARTICIPANT_FILTER_FIELDS.items() if query.get(qp) and query[qp][0]}; data = filtered_analysis_csv(db, sid, pf or None); name=export_filename(session_label(db,sid),"analyse-filtree",ext="csv"); self.send_response(200); self.send_header("Content-Type", "text/csv; charset=utf-8"); self.send_header("Content-Disposition", f"attachment; filename={name}"); self.end_headers(); self.wfile.write(data); return
            if path.startswith("/api/sessions/") and path.endswith("/ai/report-blocks"):
                sid = path.split("/")[3]
                return self.json(200, rows(db, "SELECT section_key,content,retained_at FROM report_ai_blocks WHERE session_id=?", (sid,)))
            if path == "/api/ai/config":
                cfg = get_ai_config(db)
                return self.json(200, {"enabled": cfg["enabled"], "provider": cfg["provider"], "model": cfg["model"], "keyConfigured": bool(cfg["api_key"]),
                    "providers": {k: {"label": v["label"], "pricing": v["pricing"], "models": v["models"], "keyUrl": v["key_url"]} for k, v in AI_PROVIDERS.items()}})
            if path == "/api/participant":
                sid, pid = query.get("session", [None])[0], query.get("participant", [None])[0]
                participant = db.execute("SELECT * FROM participants WHERE id=? AND session_id=?", (pid, sid)).fetchone()
                session = db.execute("SELECT * FROM sessions WHERE id=?", (sid,)).fetchone()
                return self.json(200, {"session": dict(session) if session else None, "participant": dict(participant) if participant else None, "template": template_payload(db, session["template_id"]) if session else None, "responses": {r["indicator_id"]: json.loads(r["value_json"]) for r in db.execute("SELECT * FROM responses WHERE participant_id=?", (pid,))}})
            return self.serve_static(path)
        except AuthRequiredError: return self.json(401, {"error": "Connexion requise."})
        except PermissionDeniedError: return self.json(403, {"error": "Accès refusé : cette ressource ne vous appartient pas."})
        except Exception: return self.json(500, {"error": "Erreur interne inattendue : le fichier n'a pas pu être généré. Réessayez ou contactez le support."})
        finally: db.close()

    def do_POST(self):
        path = urlparse(self.path).path
        if path == "/api/templates/import/preview":
            guard_db = self.db()
            try:
                self.require_auth(path, guard_db)
            except AuthRequiredError:
                return self.json(401, {"error": "Connexion requise."})
            finally:
                guard_db.close()
            raw=self.raw_body(); marker=b"\r\n\r\n"; start=raw.find(marker)+len(marker); end=raw.rfind(b"\r\n--"); uploaded=raw[start:end]
            try:
                preview=import_preview(uploaded); token=str(uuid.uuid4()); IMPORTS[token]=preview; return self.json(200,{"token":token,**preview})
            except ValueError as e: return self.json(400,{"error":str(e)})
        data, db = self.body(), self.db()
        try:
            user = self.require_auth(path, db)
            if path == "/api/auth/setup":
                if db.execute("SELECT 1 FROM users LIMIT 1").fetchone() is not None:
                    return self.json(409, {"error": "Un compte existe déjà : utilisez la connexion."})
                email = (data.get("email") or "").strip().lower()
                password = data.get("password") or ""
                if not email or "@" not in email: return self.json(400, {"error": "Adresse email invalide."})
                if len(password) < 8: return self.json(400, {"error": "Le mot de passe doit contenir au moins 8 caractères."})
                uid = str(uuid.uuid4()); digest, salt = hash_password(password)
                db.execute("INSERT INTO users VALUES (?,?,?,?,?,?,?)", (uid, email, digest, salt, "admin", data.get("displayName") or "", now()))
                db.commit()
                migrate_ownership(db)
                token = create_auth_token(db, uid)
                return self.json(201, {"user": {"id": uid, "email": email, "role": "admin"}}, cookie=session_cookie_header(token))
            if path == "/api/auth/login":
                email = (data.get("email") or "").strip().lower()
                row = db.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
                if not row or not verify_password(data.get("password") or "", row["password_hash"], row["password_salt"]):
                    return self.json(401, {"error": "Identifiants incorrects."})
                token = create_auth_token(db, row["id"])
                return self.json(200, {"user": {"id": row["id"], "email": row["email"], "role": row["role"], "displayName": row["display_name"]}}, cookie=session_cookie_header(token))
            if path == "/api/auth/logout":
                raw = self.headers.get("Cookie")
                if raw:
                    jar = SimpleCookie(); jar.load(raw); morsel = jar.get("epc_session")
                    if morsel: db.execute("DELETE FROM auth_tokens WHERE token_hash=?", (hashlib.sha256(morsel.value.encode("utf-8")).hexdigest(),)); db.commit()
                return self.json(200, {"ok": True}, cookie=session_cookie_header(clear=True))
            if path == "/api/ai/test":
                try:
                    cfg = require_ai(db)
                    t0 = time.time()
                    sample = generate_ai_response(cfg["provider"], cfg["model"], "Réponds uniquement par le mot OK.", "Confirme que la connexion fonctionne.", cfg["api_key"])
                    return self.json(200, {"ok": True, "provider": AI_PROVIDERS[cfg["provider"]]["label"], "model": cfg["model"], "latencyMs": int((time.time()-t0)*1000)})
                except AIError as e: return self.json(200, {"ok": False, "reason": str(e)})

            if path.startswith("/api/sessions/") and path.endswith("/ai/diagnostic"):
                sid = path.split("/")[3]
                try:
                    cfg = require_ai(db)
                    context = ai_epc_context(db, sid)
                    system = AI_SYSTEM_BASE + (" Analyse conjointement capacité et consensus pour les domaines remarquables "
                        "(faible capacité + consensus élevé = constat partagé par le groupe ; faible capacité + consensus faible = perceptions divergentes ; "
                        "capacité élevée + consensus élevé = force reconnue ; capacité élevée + consensus faible = expérience hétérogène). "
                        "Structure ta réponse avec exactement ces titres, en majuscules : POINTS SAILLANTS / POINTS DE VIGILANCE / "
                        "CONTRASTES CAPACITÉ / CONSENSUS / QUESTIONS À APPROFONDIR AVEC LE GROUPE. Reste concis et exploitable en atelier.")
                    text = generate_ai_response(cfg["provider"], cfg["model"], system, context, cfg["api_key"])
                    sug_id = log_ai_suggestion(db, sid, "diagnostic", None, cfg["provider"], cfg["model"])
                    return self.json(200, {"id": sug_id, "text": text})
                except AIError as e: return self.json(409, {"error": str(e)})

            if path.startswith("/api/sessions/") and "/ai/priority/" in path and path.endswith("/prepare"):
                parts = path.split("/"); sid, pid = parts[3], parts[6]
                try:
                    cfg = require_ai(db)
                    _, _, _, context = ai_priority_context(db, sid, pid)
                    system = AI_SYSTEM_BASE + (" Pour cette priorité, propose uniquement : 1) un constat reformulé à partir des seules données fournies, "
                        "2) des questions à poser au groupe, 3) les points nécessitant clarification. Ne propose ni cause ni recommandation ici.")
                    text = generate_ai_response(cfg["provider"], cfg["model"], system, context, cfg["api_key"])
                    sug_id = log_ai_suggestion(db, sid, "priority_prepare", pid, cfg["provider"], cfg["model"])
                    return self.json(200, {"id": sug_id, "text": text})
                except AIError as e: return self.json(409, {"error": str(e)})

            if path.startswith("/api/sessions/") and "/ai/priority/" in path and path.endswith("/entries"):
                parts = path.split("/"); sid, pid = parts[3], parts[6]
                kind = data.get("kind")
                if kind not in ("cause", "consequence", "lever"): return self.json(400, {"error": "Type d'hypothèse invalide."})
                try:
                    cfg = require_ai(db)
                    _, _, _, context = ai_priority_context(db, sid, pid)
                    kind_fr = {"cause": "des hypothèses de CAUSES", "consequence": "des CONSÉQUENCES possibles", "lever": "des LEVIERS d'action possibles"}[kind]
                    system = AI_SYSTEM_BASE + (f" Propose 3 à 5 {kind_fr} pour cette priorité, formulées comme des hypothèses à discuter — "
                        "jamais présentées comme établies (n'écris jamais « cause identifiée »). "
                        "Réponds uniquement par une liste, une hypothèse par ligne, sans numérotation ni tiret, sans autre texte.")
                    text = generate_ai_response(cfg["provider"], cfg["model"], system, context, cfg["api_key"])
                    items = [l.strip("-•* ").strip() for l in text.split("\n") if l.strip()]
                    sug_id = log_ai_suggestion(db, sid, f"priority_{kind}", pid, cfg["provider"], cfg["model"])
                    return self.json(200, {"id": sug_id, "items": items})
                except AIError as e: return self.json(409, {"error": str(e)})

            if path.startswith("/api/sessions/") and path.endswith("/ai/recommendations"):
                sid = path.split("/")[3]
                try:
                    cfg = require_ai(db)
                    q = qualitative_data(db, sid)
                    retained = [e for e in q["entries"] if e["validation_status"] == "RETENU"]
                    chain_lines = []
                    for p in q["priorities"]:
                        an = next((x for x in q["analyses"] if x["priority_id"] == p["id"]), None)
                        causes = [e["content"] for e in retained if e["priority_id"] == p["id"] and e["kind"] == "cause"]
                        levers = [e["content"] for e in retained if e["priority_id"] == p["id"] and e["kind"] == "lever"]
                        if not causes and not levers: continue
                        chain_lines.append(f"Priorité (id={p['id']}) : {p['domain_label']} — {p['indicator_label']}\n"
                            f"Constat : {an['problem'] if an and an['problem'] else 'non renseigné'}\n"
                            f"Causes validées : {'; '.join(causes) or 'aucune'}\nLeviers validés : {'; '.join(levers) or 'aucun'}")
                    if not chain_lines:
                        return self.json(200, {"id": None, "items": [], "note": "Aucune cause ni levier validé par le groupe pour l'instant : rien à proposer."})
                    context = "\n\n".join(chain_lines)
                    system = AI_SYSTEM_BASE + (" Pour chaque priorité listée, propose une ou deux recommandations d'action fondées UNIQUEMENT "
                        "sur les causes/leviers validés fournis — ne produis pas de liste générique de bonnes pratiques sans lien avec ces données. "
                        "Réponds en JSON strict, uniquement une liste d'objets avec ces clés : priorityId, title, description, category, "
                        "priorityLevel (Haute, Moyenne ou Basse), owner, horizon. Aucun texte hors JSON.")
                    text = generate_ai_response(cfg["provider"], cfg["model"], system, context, cfg["api_key"])
                    items = _parse_json_list(text)
                    sug_id = log_ai_suggestion(db, sid, "recommendation", None, cfg["provider"], cfg["model"])
                    return self.json(200, {"id": sug_id, "items": items})
                except AIError as e: return self.json(409, {"error": str(e)})

            if path.startswith("/api/sessions/") and path.endswith("/ai/training"):
                sid = path.split("/")[3]
                try:
                    cfg = require_ai(db)
                    q = qualitative_data(db, sid)
                    recs = q["recommendations"]
                    if not recs:
                        return self.json(200, {"id": None, "items": [], "note": "Aucune recommandation disponible : rien à proposer."})
                    context = "\n".join(f"- {r['title']} : {r['description']} (catégorie {r['category']})" for r in recs)
                    system = AI_SYSTEM_BASE + (" À partir de ces recommandations, identifie les besoins de formation qu'elles font apparaître. "
                        "Réponds en JSON strict, uniquement une liste d'objets avec ces clés : title, targetAudience, needText, "
                        "priorityLevel (Haute, Moyenne ou Basse). Aucun texte hors JSON.")
                    text = generate_ai_response(cfg["provider"], cfg["model"], system, context, cfg["api_key"])
                    items = _parse_json_list(text)
                    sug_id = log_ai_suggestion(db, sid, "training", None, cfg["provider"], cfg["model"])
                    return self.json(200, {"id": sug_id, "items": items})
                except AIError as e: return self.json(409, {"error": str(e)})

            if path.startswith("/api/sessions/") and path.endswith("/ai/plan"):
                sid = path.split("/")[3]
                try:
                    cfg = require_ai(db)
                    q = qualitative_data(db, sid)
                    retained = [r for r in q["recommendations"] if r["status"] == "Retenue"]
                    if not retained:
                        return self.json(200, {"id": None, "items": [], "note": "Aucune recommandation retenue pour l'instant : rien à structurer."})
                    context = "\n".join(f"- {r['title']} : {r['description']} (responsable : {r['owner'] or 'non renseigné'}, "
                        f"échéance : {r['horizon'] or 'non renseignée'})" for r in retained)
                    system = AI_SYSTEM_BASE + (" Structure ces recommandations retenues en plan d'action. N'invente aucun engagement organisationnel : "
                        "si un responsable ou une échéance ne sont pas fournis dans les données, indique-le explicitement plutôt que d'en inventer un. "
                        "Réponds en JSON strict, uniquement une liste d'objets avec ces clés : action, owner, horizon, expectedResult, "
                        "indicator, dependencies. Aucun texte hors JSON.")
                    text = generate_ai_response(cfg["provider"], cfg["model"], system, context, cfg["api_key"])
                    items = _parse_json_list(text)
                    sug_id = log_ai_suggestion(db, sid, "plan", None, cfg["provider"], cfg["model"])
                    return self.json(200, {"id": sug_id, "items": items})
                except AIError as e: return self.json(409, {"error": str(e)})

            if path.startswith("/api/sessions/") and path.endswith("/ai/report/section"):
                sid = path.split("/")[3]
                section = data.get("section")
                if section not in AI_REPORT_SECTIONS: return self.json(400, {"error": "Section de rapport invalide."})
                try:
                    cfg = require_ai(db)
                    context = ai_report_context(db, sid)
                    system = AI_SYSTEM_BASE + " " + AI_REPORT_SECTIONS[section] + (" N'invente aucune étape non réalisée : si les données "
                        "nécessaires sont absentes, indique-le sobrement plutôt que d'inventer un contenu.")
                    text = generate_ai_response(cfg["provider"], cfg["model"], system, context, cfg["api_key"])
                    sug_id = log_ai_suggestion(db, sid, f"report_{section}", None, cfg["provider"], cfg["model"])
                    return self.json(200, {"id": sug_id, "section": section, "text": text})
                except AIError as e: return self.json(409, {"error": str(e)})

            if path.startswith("/api/sessions/") and path.endswith("/ai/report/full"):
                sid = path.split("/")[3]
                try:
                    cfg = require_ai(db)
                    context = ai_report_context(db, sid)
                    results = {}
                    for section, instruction in AI_REPORT_SECTIONS.items():
                        system = AI_SYSTEM_BASE + " " + instruction + (" N'invente aucune étape non réalisée : si les données "
                            "nécessaires sont absentes, indique-le sobrement plutôt que d'inventer un contenu.")
                        results[section] = generate_ai_response(cfg["provider"], cfg["model"], system, context, cfg["api_key"])
                    sug_id = log_ai_suggestion(db, sid, "report_full", None, cfg["provider"], cfg["model"])
                    return self.json(200, {"id": sug_id, "sections": results})
                except AIError as e: return self.json(409, {"error": str(e)})

            if path == "/api/templates": return self.json(201,{"id":create_blank_template(db,data,user["id"])})
            if path == "/api/templates/import/confirm":
                preview=IMPORTS.pop(data.get("token"),None)
                if not preview: return self.json(400,{"error":"Aperçu d'import introuvable ou expiré"})
                return self.json(201,{"id":save_import(db,preview,user["id"])})
            if path.startswith("/api/templates/") and path.endswith("/clone"):
                return self.json(201,{"id":clone_template(db,path.split("/")[3],data.get("name"))})
            if path.startswith("/api/templates/") and path.endswith("/domains"):
                tid=path.split("/")[3]; guard_structural_edit(db,tid,"add")
                did=str(uuid.uuid4()); code=data.get("code") or "domain-"+uuid.uuid4().hex[:8]; db.execute("INSERT INTO domains VALUES (?,?,?,?,?,?,?)",(did,tid,code,data["label"],data.get("description",""),int(data.get("displayOrder") or next_order(db,"domains","template_id",tid)),int(data.get("active",True)))); db.commit(); return self.json(201,{"id":did})
            if path.startswith("/api/domains/") and path.endswith("/indicators"):
                did=path.split("/")[3]
                owner_domain=db.execute("SELECT template_id FROM domains WHERE id=?",(did,)).fetchone()
                if owner_domain: guard_structural_edit(db,owner_domain["template_id"],"add")
                code=(data.get("code") or "").strip()
                if code and indicator_code_conflicts(db,did,code): return self.json(400,{"error":"Cette référence est déjà utilisée par une autre question de ce domaine."})
                iid=str(uuid.uuid4()); db.execute("INSERT INTO indicators VALUES (?,?,?,?,?,?,?,?,?,?)",(iid,did,code or "indicator-"+uuid.uuid4().hex[:8],data["label"],data.get("description",""),data.get("responseType","numeric"),int(data.get("required",True)),int(data.get("displayOrder") or next_order(db,"indicators","domain_id",did)),int(data.get("active",True)),json.dumps(data.get("configuration",{})))); db.commit(); return self.json(201,{"id":iid})
            if path == "/api/sessions":
                sid = create_session(db, user["id"], data)
                if sid is None: return self.json(400,{"error":"Impossible de créer une session : le questionnaire ne contient aucun domaine avec question."})
                return self.json(201, {"id": sid})
            if path.endswith("/participants"):
                sid = path.split("/")[3]; session = db.execute("SELECT status FROM sessions WHERE id=?", (sid,)).fetchone()
                if not session or session["status"] != "open": return self.json(409, {"error": "Collecte fermée"})
                ensure_private_template(db, sid)
                pid, label = str(uuid.uuid4()), data.get("anonymousId") or f"P-{uuid.uuid4().hex[:6]}"; db.execute("INSERT INTO participants (id,session_id,anonymous_id,status,started_at,completed_at,display_name) VALUES (?,?,?,?,?,?,?)", (pid, sid, label, "in_progress", now(), None, data.get("displayName") or None)); db.commit(); return self.json(201, {"id": pid, "anonymousId": label})
            if path.endswith("/responses"):
                sid = path.split("/")[3]; stamp = now(); db.execute("INSERT INTO responses VALUES (?,?,?,?,?,?,?,?) ON CONFLICT(participant_id,indicator_id) DO UPDATE SET value_json=excluded.value_json,value_type=excluded.value_type,updated_at=excluded.updated_at", (str(uuid.uuid4()), sid, data["participantId"], data["indicatorId"], json.dumps(data["value"]), data.get("valueType", "numeric"), stamp, stamp)); db.commit(); return self.json(200, {"ok": True})
            if path.endswith("/complete"):
                pid = data["participantId"]; db.execute("UPDATE participants SET status='completed',completed_at=? WHERE id=?", (now(), pid)); db.commit(); return self.json(200, {"ok": True})
            if path.endswith("/status"):
                sid = path.split("/")[3]; status = data["status"]; db.execute("UPDATE sessions SET status=?,closed_at=? WHERE id=?", (status, now() if status == "closed" else None, sid)); db.commit(); return self.json(200, {"ok": True})
            if path.endswith("/priorities"):
                sid = path.split("/")[3]; db.execute("INSERT INTO priorities VALUES (?,?,?,?,?,?) ON CONFLICT(session_id,indicator_id) DO UPDATE SET votes=excluded.votes", (str(uuid.uuid4()), sid, data["domainId"], data["indicatorId"], int(data.get("votes", 0)), now())); db.commit(); return self.json(200, {"ok": True})
            if path.endswith("/priority-analyses"):
                sid=path.split("/")[3]; stamp=now(); priority_id=data["priorityId"]
                db.execute("INSERT INTO priority_analyses VALUES (?,?,?,?,?,?) ON CONFLICT(session_id,priority_id) DO UPDATE SET problem=excluded.problem,updated_at=excluded.updated_at",(str(uuid.uuid4()),sid,priority_id,data.get("problem",""),stamp,stamp)); db.commit(); return self.json(201,{"ok":True})
            if path.endswith("/analysis-entries"):
                sid=path.split("/")[3]; stamp=now(); eid=str(uuid.uuid4())
                db.execute("INSERT INTO analysis_entries VALUES (?,?,?,?,?,?,?,?,?,?,?)",(eid,sid,data["priorityId"],data.get("parentId") or None,data["kind"],data["content"],data.get("itemType") or None,data.get("comment") or None,data.get("validationStatus","A_DISCUTER"),stamp,stamp)); db.commit(); return self.json(201,{"id":eid})
            if path.endswith("/recommendations-v2"):
                sid=path.split("/")[3]; stamp=now(); rid=str(uuid.uuid4())
                db.execute("INSERT INTO workshop_recommendations VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(rid,sid,data.get("priorityId") or None,data.get("causeId") or None,data.get("leverId") or None,data["title"],data["description"],data.get("category","Autre"),data.get("priorityLevel","Non définie"),data.get("owner") or None,data.get("horizon") or None,data.get("comment") or None,data.get("status","Proposée"),stamp,stamp)); db.commit(); return self.json(201,{"id":rid})
            if path.endswith("/training-topics"):
                sid=path.split("/")[3]; stamp=now(); tid=str(uuid.uuid4())
                db.execute("INSERT INTO training_topics VALUES (?,?,?,?,?,?,?,?,?,?,?)",(tid,sid,data.get("priorityId") or None,data.get("recommendationId") or None,data["title"],data.get("needText") or None,data.get("targetAudience") or None,data.get("priorityLevel","Non définie"),data.get("comment") or None,stamp,stamp)); db.commit(); return self.json(201,{"id":tid})
            if path.endswith("/report-meta"):
                sid=path.split("/")[3]; db.execute("INSERT INTO session_report_meta VALUES (?,?,?,?,?,?) ON CONFLICT(session_id) DO UPDATE SET facilitator=excluded.facilitator,audience=excluded.audience,context=excluded.context,conclusion=excluded.conclusion,updated_at=excluded.updated_at",(sid,data.get("facilitator",""),data.get("audience",""),data.get("context",""),data.get("conclusion",""),now())); db.commit(); return self.json(200,{"ok":True})
            if path.endswith("/analysis-notes"):
                sid = path.split("/")[3]; stamp=now(); db.execute("INSERT INTO analysis_notes VALUES (?,?,?,?,?,?,?,?,?)", (str(uuid.uuid4()), sid, data.get("indicatorId"), data["kind"], data["content"], data.get("validationStatus", "HYPOTHESE"), stamp, stamp)); db.commit(); return self.json(201, {"ok": True})
            if path.endswith("/recommendations"):
                sid=path.split("/")[3]; db.execute("INSERT INTO recommendations VALUES (?,?,?,?,?,?,?,?,?,?)",(str(uuid.uuid4()),sid,data.get("indicatorId"),data["title"],data.get("description",""),data.get("lever",""),data.get("kind","action"),data.get("owner",""),data.get("horizon",""),now())); db.commit(); return self.json(201,{"ok":True})
            return self.json(404, {"error": "Route inconnue"})
        except (KeyError, ValueError) as e: return self.json(400, {"error": f"Requête invalide : champ manquant ou incorrect ({e})."})
        except sqlite3.IntegrityError: return self.json(409, {"error": "Action impossible : cette donnée est encore utilisée ailleurs."})
        except AuthRequiredError: return self.json(401, {"error": "Connexion requise."})
        except PermissionDeniedError: return self.json(403, {"error": "Accès refusé : cette ressource ne vous appartient pas."})
        except StructuralEditForbiddenError as e: return self.json(409, {"error": str(e)})
        except Exception: return self.json(500, {"error": "Erreur interne inattendue. Aucune donnée n'a été modifiée."})
        finally: db.close()

    def do_PUT(self):
        path, data, db = urlparse(self.path).path, self.body(), self.db()
        try:
            user = self.require_auth(path, db)
            if path == "/api/ai/config":
                if data.get("provider") and data["provider"] not in AI_PROVIDERS: return self.json(400, {"error": "Fournisseur IA inconnu."})
                cur = get_ai_config(db)
                api_key = data["apiKey"] if data.get("apiKey") else cur["api_key"]
                db.execute("INSERT INTO ai_config (id,enabled,provider,model,api_key,updated_at) VALUES (1,?,?,?,?,?) ON CONFLICT(id) DO UPDATE SET enabled=excluded.enabled,provider=excluded.provider,model=excluded.model,api_key=excluded.api_key,updated_at=excluded.updated_at",
                    (int(bool(data.get("enabled"))), data.get("provider"), data.get("model"), api_key, now()))
                db.commit(); return self.json(200, {"ok": True})
            if path.startswith("/api/sessions/") and path.endswith("/ai/report-block"):
                sid = path.split("/")[3]; section = data.get("sectionKey"); content = data.get("content", "")
                if section not in AI_REPORT_SECTIONS: return self.json(400, {"error": "Section de rapport invalide."})
                db.execute("INSERT INTO report_ai_blocks (id,session_id,section_key,content,retained_at) VALUES (?,?,?,?,?) "
                    "ON CONFLICT(session_id,section_key) DO UPDATE SET content=excluded.content,retained_at=excluded.retained_at",
                    (str(uuid.uuid4()), sid, section, content, now())); db.commit(); return self.json(200, {"ok": True})
            if path.startswith("/api/sessions/"):
                sid=path.split("/")[3]; expected=int(data["expectedParticipants"]) if data.get("expectedParticipants") not in (None,"") else None
                if data.get("templateId"):
                    tpl=db.execute("SELECT version FROM templates WHERE id=?",(data["templateId"],)).fetchone()
                    if not tpl: return self.json(404,{"error":"Questionnaire introuvable"})
                    db.execute("UPDATE sessions SET name=?,organization=?,location=?,date=?,description=?,expected_participants=?,template_id=?,template_version=? WHERE id=?",(data["name"],data.get("organization",''),data.get("location",''),data.get("date",''),data.get("description",''),expected,data["templateId"],tpl["version"],sid))
                else:
                    db.execute("UPDATE sessions SET name=?,organization=?,location=?,date=?,description=?,expected_participants=? WHERE id=?",(data["name"],data.get("organization",''),data.get("location",''),data.get("date",''),data.get("description",''),expected,sid))
                db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/participants/"):
                update_participant_profile_fields(db, path.split("/")[3], data)
                return self.json(200,{"ok":True})
            if path.startswith("/api/priority-analyses/"):
                db.execute("UPDATE priority_analyses SET problem=?,updated_at=? WHERE id=?",(data.get("problem",""),now(),path.split("/")[3])); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/ai-suggestions/"):
                status = data.get("status")
                if status not in ("proposed", "modified", "retained", "rejected"): return self.json(400, {"error": "Statut invalide."})
                db.execute("UPDATE ai_suggestions SET status=?,updated_at=? WHERE id=?", (status, now(), path.split("/")[3])); db.commit(); return self.json(200, {"ok": True})
            if path.startswith("/api/analysis-entries/"):
                db.execute("UPDATE analysis_entries SET parent_id=?,content=?,item_type=?,comment=?,validation_status=?,updated_at=? WHERE id=?",(data.get("parentId") or None,data["content"],data.get("itemType") or None,data.get("comment") or None,data.get("validationStatus","A_DISCUTER"),now(),path.split("/")[3])); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/recommendations-v2/"):
                db.execute("UPDATE workshop_recommendations SET priority_id=?,cause_id=?,lever_id=?,title=?,description=?,category=?,priority_level=?,owner=?,horizon=?,comment=?,status=?,updated_at=? WHERE id=?",(data.get("priorityId") or None,data.get("causeId") or None,data.get("leverId") or None,data["title"],data["description"],data.get("category","Autre"),data.get("priorityLevel","Non définie"),data.get("owner") or None,data.get("horizon") or None,data.get("comment") or None,data.get("status","Proposée"),now(),path.split("/")[3])); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/training-topics/"):
                db.execute("UPDATE training_topics SET priority_id=?,recommendation_id=?,title=?,need_text=?,target_audience=?,priority_level=?,comment=?,updated_at=? WHERE id=?",(data.get("priorityId") or None,data.get("recommendationId") or None,data["title"],data.get("needText") or None,data.get("targetAudience") or None,data.get("priorityLevel","Non définie"),data.get("comment") or None,now(),path.split("/")[3])); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/templates/"):
                tid=path.split("/")[3]; guard_structural_edit(db,tid,"edit")
                tpl_status=db.execute("SELECT status FROM templates WHERE id=?",(tid,)).fetchone()["status"]
                used=db.execute("SELECT 1 FROM sessions WHERE template_id=? LIMIT 1",(tid,)).fetchone()
                if used and tpl_status != "session_locked":
                    tid=clone_template(db,tid); data["versionCreated"]=True
                old=template_payload(db,tid); scale=data.get("scale",old["scale"]); db.execute("UPDATE templates SET name=?,description=?,scale_json=?,priority_json=?,updated_at=? WHERE id=?",(data.get("name",old["name"]),data.get("description",old["description"]),json.dumps(scale),json.dumps(data.get("priority",old["priority"])),now(),tid)); db.commit(); return self.json(200,{"id":tid,"versionCreated":data.get("versionCreated",False)})
            if path.startswith("/api/domains/"):
                did=path.split("/")[3]
                owner_domain=db.execute("SELECT template_id FROM domains WHERE id=?",(did,)).fetchone()
                if owner_domain: guard_structural_edit(db,owner_domain["template_id"],"edit")
                db.execute("UPDATE domains SET label=?,description=?,display_order=?,active=? WHERE id=?",(data["label"],data.get("description",""),int(data.get("displayOrder",1)),int(data.get("active",True)),did)); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/indicators/"):
                iid=path.split("/")[3]
                if not (data.get("code") or "").strip(): return self.json(400,{"error":"La référence est obligatoire."})
                if not (data.get("label") or "").strip(): return self.json(400,{"error":"La question est obligatoire."})
                owner_domain=db.execute("SELECT d.template_id FROM indicators i JOIN domains d ON d.id=i.domain_id WHERE i.id=?",(iid,)).fetchone()
                if owner_domain: guard_structural_edit(db,owner_domain["template_id"],"edit")
                if indicator_code_conflicts(db,data["domainId"],data["code"].strip(),exclude_id=iid): return self.json(400,{"error":"Cette référence est déjà utilisée par une autre question de ce domaine."})
                db.execute("UPDATE indicators SET domain_id=?,code=?,label=?,description=?,response_type=?,required=?,display_order=?,active=?,configuration_json=? WHERE id=?",(data["domainId"],data["code"],data["label"],data.get("description",""),data.get("responseType","numeric"),int(data.get("required",True)),int(data.get("displayOrder",1)),int(data.get("active",True)),json.dumps(data.get("configuration",{})),iid)); db.commit(); return self.json(200,{"ok":True})
            return self.json(404,{"error":"Route inconnue"})
        except (KeyError, ValueError) as e: return self.json(400, {"error": f"Requête invalide : champ manquant ou incorrect ({e})."})
        except sqlite3.IntegrityError: return self.json(409, {"error": "Action impossible : cette donnée est encore utilisée ailleurs."})
        except AuthRequiredError: return self.json(401, {"error": "Connexion requise."})
        except PermissionDeniedError: return self.json(403, {"error": "Accès refusé : cette ressource ne vous appartient pas."})
        except StructuralEditForbiddenError as e: return self.json(409, {"error": str(e)})
        except Exception: return self.json(500, {"error": "Erreur interne inattendue. Aucune donnée n'a été modifiée."})
        finally: db.close()

    def do_DELETE(self):
        path, db=urlparse(self.path).path,self.db()
        try:
            user = self.require_auth(path, db)
            if path.startswith("/api/sessions/") and "/ai/report-block/" in path:
                sid=path.split("/")[3]; section=path.rsplit("/",1)[1]
                db.execute("DELETE FROM report_ai_blocks WHERE session_id=? AND section_key=?",(sid,section)); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/analysis-entries/"):
                eid=path.split("/")[3]; dependent=db.execute("SELECT COUNT(*) FROM workshop_recommendations WHERE cause_id=? OR lever_id=?",(eid,eid)).fetchone()[0]
                if dependent and parse_qs(urlparse(self.path).query).get("force",["0"])[0] != "1": return self.json(409,{"error":f"Cette entrée est utilisée par {dependent} recommandation(s). Confirmez la suppression.","dependencies":dependent})
                db.execute("UPDATE workshop_recommendations SET cause_id=NULL WHERE cause_id=?",(eid,)); db.execute("UPDATE workshop_recommendations SET lever_id=NULL WHERE lever_id=?",(eid,)); db.execute("UPDATE analysis_entries SET parent_id=NULL WHERE parent_id=?",(eid,)); db.execute("DELETE FROM analysis_entries WHERE id=?",(eid,)); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/recommendations-v2/"):
                rid=path.split("/")[3]; dependent=db.execute("SELECT COUNT(*) FROM training_topics WHERE recommendation_id=?",(rid,)).fetchone()[0]
                if dependent and parse_qs(urlparse(self.path).query).get("force",["0"])[0] != "1": return self.json(409,{"error":f"Cette recommandation est liée à {dependent} thème(s) de formation. Confirmez la suppression.","dependencies":dependent})
                db.execute("UPDATE training_topics SET recommendation_id=NULL WHERE recommendation_id=?",(rid,)); db.execute("DELETE FROM workshop_recommendations WHERE id=?",(rid,)); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/training-topics/"):
                db.execute("DELETE FROM training_topics WHERE id=?",(path.split("/")[3],)); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/templates/"):
                tid=path.split("/")[3]
                guard_structural_edit(db,tid,"delete_template")
                if db.execute("SELECT 1 FROM sessions WHERE template_id=? LIMIT 1",(tid,)).fetchone():
                    if parse_qs(urlparse(self.path).query).get("force",["0"])[0] == "1": db.execute("UPDATE templates SET status='archived',updated_at=? WHERE id=?",(now(),tid)); db.commit(); return self.json(200,{"ok":True,"archived":True})
                    return self.json(409,{"error":"Suppression impossible. Ce questionnaire est utilisé par une ou plusieurs sessions d’atelier. Vous pouvez le conserver, créer une nouvelle version, ou confirmer son retrait de la liste des modèles."})
                db.execute("DELETE FROM indicators WHERE domain_id IN (SELECT id FROM domains WHERE template_id=?)",(tid,)); db.execute("DELETE FROM domains WHERE template_id=?",(tid,)); db.execute("DELETE FROM templates WHERE id=?",(tid,)); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/domains/"):
                did=path.split("/")[3]
                owner_domain=db.execute("SELECT template_id FROM domains WHERE id=?",(did,)).fetchone()
                if owner_domain: guard_structural_edit(db,owner_domain["template_id"],"delete")
                affected=rows(db,"SELECT DISTINCT s.id,s.name FROM sessions s JOIN responses r ON r.session_id=s.id JOIN indicators i ON i.id=r.indicator_id WHERE i.domain_id=?",(did,))
                if affected:
                    names=", ".join(a["name"] for a in affected)
                    return self.json(409,{"error":DOMAIN_FROZEN_MESSAGE.format(count=len(affected),names=names),"sessions":affected})
                db.execute("DELETE FROM indicators WHERE domain_id=?",(did,)); db.execute("DELETE FROM domains WHERE id=?",(did,)); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/indicators/"):
                iid=path.split("/")[3]
                owner_domain=db.execute("SELECT d.template_id FROM indicators i JOIN domains d ON d.id=i.domain_id WHERE i.id=?",(iid,)).fetchone()
                if owner_domain: guard_structural_edit(db,owner_domain["template_id"],"delete")
                used=db.execute("SELECT COUNT(*) FROM responses WHERE indicator_id=?",(iid,)).fetchone()[0]
                if used:
                    return self.json(409,{"error":INDICATOR_FROZEN_MESSAGE.format(count=used),"dependencies":used})
                db.execute("DELETE FROM indicators WHERE id=?",(iid,)); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/sessions/") and "/priorities/" in path:
                parts=path.split("/"); db.execute("DELETE FROM priorities WHERE session_id=? AND indicator_id=?",(parts[3],parts[5])); db.commit(); return self.json(200,{"ok":True})
            if path.startswith("/api/sessions/") and len(path.rstrip("/").split("/")) == 4:
                sid=path.rstrip("/").split("/")[3]
                delete_session(db,sid); return self.json(200,{"ok":True})
            return self.json(404,{"error":"Route inconnue"})
        except (KeyError, ValueError) as e: return self.json(400, {"error": f"Requête invalide : champ manquant ou incorrect ({e})."})
        except sqlite3.IntegrityError: return self.json(409, {"error": "Action impossible : cette donnée est encore utilisée ailleurs."})
        except AuthRequiredError: return self.json(401, {"error": "Connexion requise."})
        except PermissionDeniedError: return self.json(403, {"error": "Accès refusé : cette ressource ne vous appartient pas."})
        except StructuralEditForbiddenError as e: return self.json(409, {"error": str(e)})
        except Exception: return self.json(500, {"error": "Erreur interne inattendue. Aucune donnée n'a été modifiée."})
        finally: db.close()

    def serve_static(self, path):
        requested = "index.html" if path in ("/", "") else path.lstrip("/")
        target = (STATIC / requested).resolve()
        if not str(target).startswith(str(STATIC.resolve())) or not target.is_file(): self.send_error(404); return
        self.send_response(200); self.send_header("Content-Type", "text/html; charset=utf-8" if target.suffix == ".html" else "application/javascript; charset=utf-8" if target.suffix == ".js" else "text/css; charset=utf-8"); self.send_header("Cache-Control","no-store, max-age=0"); self.end_headers(); self.wfile.write(target.read_bytes())


def main():
    db=connect(); init_db(db); db.close()
    host=os.environ.get("HOST","127.0.0.1")
    port=int(os.environ.get("PORT","8000"))
    server=ThreadingHTTPServer((host, port), Handler)
    print(f"EPC Workshop Engine: http://{host}:{port}")
    server.serve_forever()

if __name__ == "__main__": main()
