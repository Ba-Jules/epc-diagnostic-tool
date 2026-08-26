"""Manifeste de restitution (lot 6, cf. AUDIT_MODULARISATION_8800.md) : quelles
sections de rapport, quels graphiques et quels prompts IA sont disponibles pour
un modele de questionnaire donne.

Un seul modele existe aujourd'hui (epc_seneval) - son manifeste liste
exactement ce qui existe deja, sans rien changer au rendu ("conserver
integralement le rapport EPC actuel", note de rollback de l'audit : "renderer
EPC historique selectionne pour model_key=epc_seneval"). Un futur second
modele apporterait son propre manifeste plutot que de brancher des
conditions supplementaires dans le code de rendu - c'est le seul point que
ce lot ajoute : un point d'indirection, pas une reecriture des generateurs
XLSX/DOCX (docx_bars_chart, docx_radar_chart, etc. restent inchanges, ils ne
savent toujours dessiner que ce que l'EPC dessine deja).
"""
from __future__ import annotations

import csv
import json
import math
import re
from io import BytesIO, StringIO

from .db import MODEL_KEY_EPC_SENEVAL, rows, template_payload
from .profile import get_participant_profile_values, participant_profile_breakdown
from .scoring import analysis, filtered_analysis


def findings_rows(findings: dict) -> list[list]:
    """Flattens objective_findings() (epc/scoring.py) into plain rows for
    XLSX/DOCX/report display - mission de parite :8810->:8820, cf.
    consignes_claude.txt."""
    out: list[list] = []
    for d in findings["forces"]["domains"]: out.append(["Force", "Domaine", d["label"], d["capacity"], d["consensus"]])
    for i in findings["forces"]["indicators"]: out.append(["Force", "Indicateur", f"{i['domain']} — {i['label']}", i["capacity"], i["consensus"]])
    for d in findings["fragilites"]["domains"]: out.append(["Fragilité", "Domaine", d["label"], d["capacity"], d["consensus"]])
    for i in findings["fragilites"]["indicators"]: out.append(["Fragilité", "Indicateur", f"{i['domain']} — {i['label']}", i["capacity"], i["consensus"]])
    for v in findings["vigilance"]:
        if v["reason"] == "ecart_sous_populations":
            out.append(["Vigilance", "Écart entre sous-populations", v["label"], v.get("gap"), ""])
        else:
            out.append(["Vigilance", "Domaine", v["label"], v.get("capacity"), v.get("consensus")])
    return out

# Ordre et cles des feuilles du classeur XLSX (report_xlsx) - identique a
# l'ordre historique des onglets, jamais reordonne ni omis pour epc_seneval.
REPORT_SECTIONS_EPC_SENEVAL = [
    "synthese", "profil_participants", "domaines", "indicateurs", "constats",
    "priorites", "analyses", "causes", "consequences", "leviers",
    "recommandations", "formations", "plan_action", "questionnaire",
]

AI_SYSTEM_PROMPT_EPC_SENEVAL = (
    "Tu assistes un modérateur d'atelier de diagnostic organisationnel EPC/SENEVAL. "
    "Tu interprètes des données déjà calculées ; tu ne recalcules jamais un score, tu n'inventes jamais un fait, "
    "une cause, une conséquence ou une recommandation absente des données fournies. "
    "Style professionnel, clair, factuel, sans jargon d'IA, sans formules comme « l'IA constate que ». Réponds en français."
)

# section_key -> (libelle affiche, instruction donnee a l'IA) ; l'ordre est
# l'ordre d'affichage dans le rapport DOCX/XLSX et de generation pour
# "generer tout le rapport".
AI_REPORT_SECTIONS_EPC_SENEVAL = {
    "resume_executif": ("Résumé exécutif", "Rédige un RÉSUMÉ EXÉCUTIF de l'atelier : situation générale, principaux constats, points forts, points de vigilance, priorités retenues, principales orientations."),
    "lecture_diagnostic": ("Lecture du diagnostic", "Rédige une LECTURE DU DIAGNOSTIC : tendances, écarts, convergences, divergences, domaines remarquables, à partir de la capacité et du consensus."),
    "synthese_domaines": ("Synthèse par domaine", "Rédige une SYNTHÈSE PAR DOMAINE : pour chaque domaine, résultat et interprétation prudente, en lien avec les analyses validées si disponibles."),
    "synthese_priorites": ("Synthèse des priorités", "Rédige une SYNTHÈSE DES PRIORITÉS : priorités retenues, constats, causes validées, leviers retenus."),
    "synthese_recommandations": ("Synthèse des recommandations", "Rédige une SYNTHÈSE DES RECOMMANDATIONS retenues, regroupées en catégories cohérentes si pertinent."),
    "synthese_formations": ("Synthèse des besoins de formation", "Rédige une SYNTHÈSE DES BESOINS DE FORMATION retenus."),
    "synthese_plan": ("Synthèse du plan d'action", "Rédige une SYNTHÈSE DU PLAN D'ACTION à partir des recommandations retenues."),
    "conclusion": ("Conclusion générale proposée", "Propose une CONCLUSION GÉNÉRALE concise et institutionnelle. Précise qu'il s'agit d'une proposition, pas d'une décision validée."),
}

RESTITUTION_MANIFESTS = {
    MODEL_KEY_EPC_SENEVAL: {
        "modelKey": MODEL_KEY_EPC_SENEVAL,
        "modelName": "EPC / SENEVAL",
        "reportSections": REPORT_SECTIONS_EPC_SENEVAL,
        "aiSystemPrompt": AI_SYSTEM_PROMPT_EPC_SENEVAL,
        "aiReportSections": {k: v[1] for k, v in AI_REPORT_SECTIONS_EPC_SENEVAL.items()},
        "aiSectionLabels": {k: v[0] for k, v in AI_REPORT_SECTIONS_EPC_SENEVAL.items()},
    },
}


def resolve_model_key(template: dict) -> str:
    """A template without its own model_key (a custom/blank questionnaire,
    cf. lot 2) still runs on the single EPC scoring engine that exists today
    (lot 3 deliberately deferred a multi-strategy registry) - so it gets the
    same restitution manifest as the canonical EPC/SENEVAL model."""
    return template.get("model_key") or MODEL_KEY_EPC_SENEVAL


def restitution_manifest(template: dict) -> dict:
    return RESTITUTION_MANIFESTS.get(resolve_model_key(template), RESTITUTION_MANIFESTS[MODEL_KEY_EPC_SENEVAL])


def session_restitution_manifest(db, session_id: str) -> dict:
    """Resolves the manifest for the model session_id's questionnaire belongs
    to - the one lookup point every route needing model-driven prompts,
    report sections or charts should call, instead of hardcoding a model."""
    session = db.execute("SELECT template_id FROM sessions WHERE id=?", (session_id,)).fetchone()
    if not session:
        return RESTITUTION_MANIFESTS[MODEL_KEY_EPC_SENEVAL]
    return restitution_manifest(template_payload(db, session["template_id"]))


# --- Extraction mecanique du lot 6b (AUDIT_MODULARISATION_8800.md) : generateurs
# XLSX/DOCX et donnees qualitatives, deplaces depuis app.py a l'identique (verifie
# par diff avant suppression, comme le lot 1). Aucune formule ni mise en page ne
# change ici, seul l'emplacement du code bouge.

try:  # Used only to generate the downloadable Excel template; the app stays local.
    import xlsxwriter
except ImportError:
    xlsxwriter = None
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None
    Inches = None
try:  # Used to draw the report chart images (Word and Excel both embed the same PNGs); reports degrade to text/tables without it.
    from PIL import Image as PILImage, ImageDraw, ImageFont
except ImportError:
    PILImage = None


def _n(v):
    return "—" if v is None else round(v, 1)


def _c(obj):
    return "non calculable (1 seul répondant)" if obj.get("consensusNote") == "single_respondent" else _n(obj.get("consensus"))


def report_rows(db, sid):
    a=analysis(db,sid)
    if not a:
        raise ValueError("Session introuvable")
    return a, [[d['label'],d['capacity'],_c(d),d['gradedCapacity'],d['gradedConsensus'],d['responses']] for d in a['domains']]

# Export "reponses individuelles" (mission de parite :8810->:8820, cf.
# consignes_claude.txt) : format large, une ligne par participant valide,
# une colonne par champ de profil ACTIF de la session (jamais les 5 champs
# fixes de :8810 - lus dynamiquement depuis profile_fields) + une colonne par
# indicateur actif. Cloisonne a UNE session (donc a un groupe) : jamais de
# donnee d'une autre session/campagne.
def individual_responses_rows(db, session_id: str) -> dict:
    session = db.execute("SELECT template_id, profile_schema_id FROM sessions WHERE id=?", (session_id,)).fetchone()
    if not session:
        return {"indicators": [], "profileFields": [], "rows": []}
    template = template_payload(db, session["template_id"])
    indicators = [i for d in template["domains"] for i in d["indicators"] if i["active"]]
    profile_fields = rows(db, "SELECT * FROM profile_fields WHERE schema_id=? AND active=1 ORDER BY display_order", (session["profile_schema_id"],)) if session["profile_schema_id"] else []
    participants = rows(db, "SELECT * FROM participants WHERE session_id=? AND status='completed' ORDER BY started_at", (session_id,))
    out_rows = []
    for p in participants:
        nominative = bool((p["display_name"] or "").strip())
        responses = {r["indicator_id"]: json.loads(r["value_json"]) for r in db.execute("SELECT indicator_id,value_json FROM responses WHERE session_id=? AND participant_id=?", (session_id, p["id"]))}
        profile_values = get_participant_profile_values(db, p["id"])
        row = {"id": p["anonymous_id"], "name": p["display_name"] if nominative else "", "status": "Nominatif" if nominative else "Anonyme"}
        for field in profile_fields:
            value = profile_values.get(field["field_key"])
            row[field["id"]] = ", ".join(str(v) for v in value) if isinstance(value, list) else ("" if value is None else value)
        for ind in indicators:
            row[ind["id"]] = responses.get(ind["id"], "")
        out_rows.append(row)
    return {"indicators": indicators, "profileFields": profile_fields, "rows": out_rows}


INDIVIDUAL_RESPONSES_HEAD = ["Identifiant participant", "Nom", "Statut"]


def _individual_responses_table(db, session_id: str):
    data = individual_responses_rows(db, session_id)
    head = INDIVIDUAL_RESPONSES_HEAD + [f["label"] for f in data["profileFields"]] + [i["code"] for i in data["indicators"]]
    body = [[row["id"], row["name"], row["status"]] + [row[f["id"]] for f in data["profileFields"]] + [row[i["id"]] for i in data["indicators"]] for row in data["rows"]]
    return head, body


def individual_responses_xlsx(db, session_id: str):
    head, body = _individual_responses_table(db, session_id)
    out = BytesIO(); wb = xlsxwriter.Workbook(out, {"in_memory": True}); h = wb.add_format({"bold": True, "bg_color": "#1F4E78", "font_color": "#FFFFFF"})
    s = wb.add_worksheet("Réponses individuelles"); s.write_row(0, 0, head, h)
    for n, row in enumerate(body): s.write_row(n + 1, 0, row)
    s.set_column(0, len(head) - 1, 20)
    wb.close(); return out.getvalue()


def individual_responses_csv(db, session_id: str):
    head, body = _individual_responses_table(db, session_id)
    buf = StringIO(); writer = csv.writer(buf); writer.writerow(head)
    for row in body: writer.writerow(row)
    return buf.getvalue().encode()


# Export "resultats filtres" (mission de parite :8810->:8820, cf.
# consignes_claude.txt) : reutilise filtered_analysis() (epc/scoring.py) -
# combine autant de dimensions que fournies (AND entre dimensions, OR entre
# valeurs d'une meme dimension, jamais un nom de champ code en dur) -,
# l'export correspond donc exactement aux resultats affiches par l'ecran de
# filtrage, y compris la suppression des petits effectifs (MIN_COHORT_N).
def filtered_analysis_rows(db, session_id: str, filters: dict):
    a = filtered_analysis(db, session_id, filters)
    if a is None:
        return None
    session_row = db.execute("SELECT name, campaign_id, group_code FROM sessions WHERE id=?", (session_id,)).fetchone()
    filt_desc = "; ".join(f"{f['fieldLabel']}={','.join(str(v) for v in f['values'])}" for f in a["filters"]["applied"]) or "Tous les participants"
    header = [
        ["Mission", session_row["name"]],
        ["Campagne", session_row["campaign_id"] or "—"],
        ["Groupe", session_row["group_code"] or "—"],
        ["Filtres actifs", filt_desc],
        ["N (validés)", a["completedCount"]],
    ]
    domain_rows = [[d["label"], d["capacity"], _c(d), d["gradedCapacity"], d["gradedConsensus"], d["responses"]] for d in a["domains"]]
    indicator_rows = [[d["label"], i["label"], i["capacity"], _c(i), i["responses"], i["missing"]] for d in a["domains"] for i in d["indicators"]]
    return header, domain_rows, indicator_rows


def filtered_analysis_xlsx(db, session_id: str, filters: dict):
    result = filtered_analysis_rows(db, session_id, filters)
    if result is None:
        raise ValueError("Session introuvable")
    header, domain_rows, indicator_rows = result
    out = BytesIO(); wb = xlsxwriter.Workbook(out, {"in_memory": True}); h = wb.add_format({"bold": True, "bg_color": "#1F4E78", "font_color": "#FFFFFF"})
    s = wb.add_worksheet("Filtre"); [s.write_row(n, 0, row) for n, row in enumerate(header)]; s.set_column(0, 1, 30)
    d = wb.add_worksheet("Domaines"); d.write_row(0, 0, ["Domaine", "Capacité", "Consensus", "Cap. graduée", "Cons. gradué", "Réponses"], h); [d.write_row(n + 1, 0, row) for n, row in enumerate(domain_rows)]; d.set_column(0, 5, 22)
    i = wb.add_worksheet("Indicateurs"); i.write_row(0, 0, ["Domaine", "Référence", "Capacité", "Consensus", "Réponses", "Manquants"], h); [i.write_row(n + 1, 0, row) for n, row in enumerate(indicator_rows)]; i.set_column(0, 5, 24)
    wb.close(); return out.getvalue()


def filtered_analysis_csv(db, session_id: str, filters: dict):
    result = filtered_analysis_rows(db, session_id, filters)
    if result is None:
        raise ValueError("Session introuvable")
    header, domain_rows, indicator_rows = result
    buf = StringIO(); writer = csv.writer(buf)
    for row in header: writer.writerow(row)
    writer.writerow([]); writer.writerow(["Domaine", "Capacité", "Consensus", "Cap. graduée", "Cons. gradué", "Réponses"])
    for row in domain_rows: writer.writerow(row)
    writer.writerow([]); writer.writerow(["Domaine", "Référence", "Capacité", "Consensus", "Réponses", "Manquants"])
    for row in indicator_rows: writer.writerow(row)
    return buf.getvalue().encode()


def report_xlsx(db,sid):
    a,rs=report_rows(db,sid); q=qualitative_data(db,sid); meta=report_data(db,sid)["meta"]; template=template_payload(db,a['session']['template_id']); manifest=restitution_manifest(template); out=BytesIO(); wb=xlsxwriter.Workbook(out,{"in_memory":True}); h=wb.add_format({"bold":True,"bg_color":"#1F4E78","font_color":"#FFFFFF"})
    analyses={x['priority_id']:x for x in q['analyses']}; priority_rows=[[p['id'],p['domain_label'],p['indicator_code'],p['indicator_label'],analyses.get(p['id'],{}).get('problem','')] for p in q['priorities']]
    sheets=[("synthese","Synthèse",["Atelier","Organisation","Lieu","Date","Animateur","Public","Contexte","Conclusion","Capacité","Consensus"],[[a['session']['name'],a['session']['organization'],a['session']['location'],a['session']['date'],meta['facilitator'],meta['audience'],meta['context'],meta['conclusion'],a['global']['capacity'],_c(a['global'])]]),("profil_participants","Profil_participants",["Champ","Valeur","N"],participant_profile_breakdown(db,sid)),("domaines","Domaines",["Domaine","Capacité","Consensus","Cap. graduée","Cons. gradué","Réponses"],rs),("indicateurs","Indicateurs",["Domaine","Référence","Capacité","Consensus","Réponses","Manquants"],[[d['label'],i['label'],i['capacity'],_c(i),i['responses'],i['missing']] for d in a['domains'] for i in d['indicators']]),("constats","Constats",["Type","Niveau","Libellé","Capacité","Consensus"],findings_rows(a['findings'])),("priorites","Priorités",["ID priorité","Domaine","Référence","Indicateur","Constat"],priority_rows),("analyses","Analyses",["ID","Priorité","Constat"],[[x['id'],x['priority_id'],x['problem']] for x in q['analyses']]),("causes","Causes",["ID","Priorité","Parent","Cause","Type","Statut"],[[x['id'],x['priority_id'],x['parent_id'],x['content'],x['item_type'],x['validation_status']] for x in q['entries'] if x['kind']=='cause']),("consequences","Conséquences",["ID","Priorité","Conséquence","Statut"],[[x['id'],x['priority_id'],x['content'],x['validation_status']] for x in q['entries'] if x['kind']=='consequence']),("leviers","Leviers",["ID","Priorité","Levier","Commentaire","Statut"],[[x['id'],x['priority_id'],x['content'],x['comment'],x['validation_status']] for x in q['entries'] if x['kind']=='lever']),("recommandations","Recommandations",["ID","Priorité","Cause","Levier","Titre","Description","Catégorie","Niveau","Responsable","Échéance","Statut"],[[x['id'],x['priority_id'],x['cause_id'],x['lever_id'],x['title'],x['description'],x['category'],x['priority_level'],x['owner'],x['horizon'],x['status']] for x in q['recommendations']]),("formations","Formations",["ID","Priorité","Recommandation","Intitulé","Besoin","Public","Niveau","Commentaire"],[[x['id'],x['priority_id'],x['recommendation_id'],x['title'],x['need_text'],x['target_audience'],x['priority_level'],x['comment']] for x in q['trainingTopics']]),("plan_action","Plan_action",["N°","Action / recommandation","Origine","Responsable","Échéance","Priorité","Statut"],[[n+1,x['title'],x['priority_id'] or '—',x['owner'] or '—',x['horizon'] or '—',x['priority_level'],x['status']] for n,x in enumerate(q['recommendations']) if x['status']=='Retenue']),("questionnaire","Questionnaire",["Domaine","Référence","Indicateur","Échelle"],[[d['label'],i['label'],i['description'],f"{template['scale']['min']}–{template['scale']['max']}"] for d in template['domains'] for i in d['indicators'] if i['active']])]
    for key,name,head,data in sheets:
        if key not in manifest["reportSections"]: continue
        s=wb.add_worksheet(name);s.write_row(0,0,head,h);[s.write_row(n+1,0,row) for n,row in enumerate(data)];s.set_column(0,len(head)-1,24)
    ai_blocks=rows(db,"SELECT section_key,content FROM report_ai_blocks WHERE session_id=?",(sid,))
    if ai_blocks:
        wrap=wb.add_format({"text_wrap":True,"valign":"top"})
        ai_sheet=wb.add_worksheet("Synthèse_IA"); ai_sheet.write_row(0,0,["Section","Contenu proposé par l'assistant IA, retenu par le modérateur"],h)
        for n,b in enumerate(ai_blocks): ai_sheet.write_row(n+1,0,[manifest["aiSectionLabels"].get(b['section_key'],b['section_key']),b['content']],wrap)
        ai_sheet.set_column(0,0,28); ai_sheet.set_column(1,1,110)
    domains=[d for d in a['domains'] if d.get('capacity') is not None]
    if PILImage is not None and domains:
        gs=wb.add_worksheet("Graphiques"); gs.set_column(0,0,4); row=1
        gs.write(0,0,"Vue synthétique",h)
        row=xlsx_add_image(gs,row,0,pil_chart_grid([
            docx_radar_chart(domains),
            docx_bars_chart(domains,"standard","Notes standardisées par domaine",with_mean=True),
            docx_bars_chart(domains,"graded","Notes graduées par domaine",with_mean=True),
            docx_grid_chart(domains),
        ]))
        if len(domains)>1:
            row+=1; gs.write(row,0,"Analyse comparative — cohorte des domaines",h); row+=1
            row=xlsx_add_image(gs,row,0,docx_cohort_chart(domains))
        for i in range(0,len(domains),4):
            group=domains[i:i+4]
            tiles=[]
            for d in group:
                inds=[ind for ind in d["indicators"] if ind.get("capacity") is not None]
                if inds:
                    caption=f"Capacité {pdf_fmt(d['capacity'])} · Consensus {pdf_c(d)} · {d['responses']} répondant(s)"
                    tiles.append((docx_bars_chart(inds,"standard",d["label"]),caption))
            if tiles:
                row+=1; gs.write(row,0,"Détail par domaine",h); row+=1
                row=xlsx_add_image(gs,row,0,pil_chart_grid(tiles))
        gs.activate()
    wb.close();return out.getvalue()

DOCX_FONT_CACHE = {}


def docx_font(size, bold=False):
    key = (size, bold)
    if key in DOCX_FONT_CACHE:
        return DOCX_FONT_CACHE[key]
    names = ["arialbd.ttf", "Arial Bold.ttf", "DejaVuSans-Bold.ttf", "LiberationSans-Bold.ttf"] if bold else \
            ["arial.ttf", "Arial.ttf", "DejaVuSans.ttf", "LiberationSans-Regular.ttf"]
    dirs = ["C:/Windows/Fonts", "/usr/share/fonts/truetype/dejavu", "/usr/share/fonts/truetype/liberation",
            "/System/Library/Fonts/Supplemental", "/System/Library/Fonts"]
    font = None
    for d in dirs:
        for n in names:
            try:
                font = ImageFont.truetype(f"{d}/{n}", size); break
            except Exception:
                continue
        if font:
            break
    if font is None:
        try:
            font = ImageFont.load_default(size=size)
        except TypeError:
            font = ImageFont.load_default()
    DOCX_FONT_CACHE[key] = font
    return font


def docx_text_w(draw, text, font):
    b = draw.textbbox((0, 0), text, font=font)
    return b[2] - b[0]


def docx_rotated_text(img, text, x, y, angle, font, fill, anchor_end=False):
    draw0 = ImageDraw.Draw(img)
    b = draw0.textbbox((0, 0), text, font=font)
    w, h = b[2] - b[0], b[3] - b[1]
    pad = 4
    txt_img = PILImage.new("RGBA", (w + pad * 2, h + pad * 2), (255, 255, 255, 0))
    ImageDraw.Draw(txt_img).text((pad, pad), text, font=font, fill=fill)
    rot = txt_img.rotate(angle, expand=1, resample=PILImage.BICUBIC)
    px = x - (rot.width if anchor_end else 0)
    img.paste(rot, (int(px), int(y)), rot)


def docx_dashed_line(draw, p1, p2, color, dash=8, gap=5, width=2):
    x1, y1 = p1; x2, y2 = p2
    length = math.hypot(x2 - x1, y2 - y1)
    if length == 0:
        return
    ux, uy = (x2 - x1) / length, (y2 - y1) / length
    d = 0
    while d < length:
        e = min(d + dash, length)
        draw.line([x1 + ux * d, y1 + uy * d, x1 + ux * e, y1 + uy * e], fill=color, width=width)
        d += dash + gap


def docx_bars_chart(items, mode, title, with_mean=False):
    """Vertical grouped column chart, raster twin of static/app.js bars() / pdf_bars_chart()."""
    grad = mode == "graded"
    data = [d for d in items if d.get("capacity") is not None]
    if with_mean and data:
        mc = sum(d["capacity"] for d in data) / len(data)
        cons_data = [d for d in data if d.get("consensus") is not None]
        ms = sum(d["consensus"] for d in cons_data) / len(cons_data) if cons_data else None
        gc = sum(d.get("gradedCapacity") or 0 for d in data) / len(data)
        gs = sum(d.get("gradedConsensus") or 0 for d in data) / len(data)
        data = data + [{"label": "Moyen", "code": "Moyen", "capacity": mc, "consensus": ms, "gradedCapacity": gc, "gradedConsensus": gs}]
    w = max(900, 90 + len(data) * 140) if data else 900
    h = 560
    img = PILImage.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    f_title, f_axis, f_legend, f_val, f_cat, f_cat_b, f_xtitle = docx_font(24, True), docx_font(15), docx_font(15), docx_font(14), docx_font(14), docx_font(14, True), docx_font(17, True)
    draw.text((w / 2, 30), title, font=f_title, fill="black", anchor="mm")
    if not data:
        draw.text((30, 66), "Pas encore de données disponibles.", font=f_axis, fill="black")
        return img
    left, right = 100, w - 30
    top, bottom = 72, h - 120
    plot_w, plot_h = right - left, bottom - top
    axis_min, axis_max = 20, 100
    scaleY = lambda v: bottom - (v - axis_min) / (axis_max - axis_min) * plot_h
    draw.rectangle([left, top, right, bottom], fill="#e9e9e9", outline="black")
    for v in (20, 40, 60, 80, 100):
        yy = scaleY(v)
        draw.line([left, yy, right, yy], fill="#bbbbbb")
        draw.text((left - 12, yy), str(v), font=f_axis, fill="black", anchor="rm")
    cap_label = "Capacité graduée" if grad else "Capacité standardisée"
    cons_label = "Consensus gradué" if grad else "Consensus standardisé"
    draw.rectangle([left, 40, left + 16, 56], fill="#9999FF")
    draw.text((left + 24, 48), cap_label, font=f_legend, fill="black", anchor="lm")
    draw.rectangle([left + 240, 40, left + 256, 56], fill="#993366")
    draw.text((left + 264, 48), cons_label, font=f_legend, fill="black", anchor="lm")
    n = len(data); step = plot_w / n
    for i, d in enumerate(data):
        x = left + i * step
        cap = d.get("gradedCapacity") if grad else d.get("capacity")
        cons = d.get("gradedConsensus") if grad else d.get("consensus")
        cv, sv = max(axis_min, cap or 0), max(axis_min, cons or 0)
        bw = step * 0.32
        y_cap, y_cons = scaleY(cv), scaleY(sv)
        draw.rectangle([x + step * .12, y_cap, x + step * .12 + bw, bottom], fill="#9999FF", outline="black")
        draw.rectangle([x + step * .12 + bw + 4, y_cons, x + step * .12 + bw + 4 + bw, bottom], fill="#993366", outline="black")
        cap_txt = "—" if cap is None else (str(round(cap)) if grad else f"{cap:.1f}")
        cons_txt = "—" if cons is None else (str(round(cons)) if grad else f"{cons:.1f}")
        draw.text((x + step * .12 + bw / 2, y_cap - 12), cap_txt, font=f_val, fill="black", anchor="mm")
        draw.text((x + step * .12 + bw + 4 + bw / 2, y_cons - 12), cons_txt, font=f_val, fill="black", anchor="mm")
        docx_rotated_text(img, pdf_short_label(d), x + step * .12 + bw + 2, bottom + 8, -40, f_cat_b if d.get("label") == "Moyen" else f_cat, "black")
    draw.text((left + plot_w / 2, h - 34), "Domaines de compétence", font=f_xtitle, fill="black", anchor="mm")
    return img


def docx_grid_chart(items):
    """Quadrant scatter with numbered markers, raster twin of graduatedGrid() / pdf_grid_chart()."""
    data = [d for d in items if d.get("capacity") is not None and d.get("consensus") is not None]
    w, h = 700, 620
    img = PILImage.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    f_title, f_axis, f_axis_title, f_quad, f_num, f_legend = docx_font(22, True), docx_font(14), docx_font(16, True), docx_font(13), docx_font(15, True), docx_font(13)
    draw.text((w / 2, 26), "Positionnement des domaines", font=f_title, fill="black", anchor="mm")
    if not data:
        draw.text((30, 60), "Pas encore de données disponibles pour la grille graduée.", font=f_axis, fill="black")
        return img
    left, right = 100, w - 40
    top, bottom = 60, h - 150
    plot_w, plot_h = right - left, bottom - top
    sx = lambda v: left + v / 100 * plot_w
    sy = lambda v: bottom - v / 100 * plot_h
    draw.rectangle([left, top, right, bottom], fill="#e9e9e9", outline="black")
    for v in (0, 20, 40, 60, 80, 100):
        draw.line([sx(v), top, sx(v), bottom], fill="#bbbbbb")
        draw.line([left, sy(v), right, sy(v)], fill="#bbbbbb")
        draw.text((sx(v), bottom + 16), str(v), font=f_axis, fill="black", anchor="mm")
        draw.text((left - 12, sy(v)), str(v), font=f_axis, fill="black", anchor="rm")
    draw.line([sx(50), top, sx(50), bottom], fill="#333333", width=2)
    draw.line([left, sy(50), right, sy(50)], fill="#333333", width=2)
    draw.text((left + plot_w / 2, bottom + 38), "Capacité", font=f_axis_title, fill="black", anchor="mm")
    docx_rotated_text(img, "Consensus", 22, top + plot_h / 2 + 45, 90, f_axis_title, "black")
    draw.text((left + 6, top + 8), "Faible capacité / consensus élevé", font=f_quad, fill="black")
    draw.text((right - 6 - docx_text_w(draw, "Capacité élevée / consensus élevé", f_quad), top + 8), "Capacité élevée / consensus élevé", font=f_quad, fill="black")
    draw.text((left + 6, bottom - 20), "Faible capacité / consensus faible", font=f_quad, fill="black")
    draw.text((right - 6 - docx_text_w(draw, "Capacité élevée / consensus faible", f_quad), bottom - 20), "Capacité élevée / consensus faible", font=f_quad, fill="black")
    avg_cap = sum(d["capacity"] for d in data) / len(data)
    avg_cons = sum(d["consensus"] for d in data) / len(data)
    order = sorted(range(len(data)), key=lambda i: data[i]["capacity"] + data[i]["consensus"])
    rank = {idx: r for r, idx in enumerate(order)}
    dirs = [(16, -14), (16, 18), (-16, -14), (-16, 18)]
    for i, d in enumerate(data):
        cx, cy = sx(d["capacity"]), sy(d["consensus"])
        r = 7
        draw.polygon([(cx, cy - r), (cx + r, cy), (cx, cy + r), (cx - r, cy)], fill="#000080")
        dx, dy = dirs[rank[i] % 4]
        if cx + dx > right - 16: dx = -abs(dx)
        elif cx + dx < left + 16: dx = abs(dx)
        if cy + dy < top + 12: dy = abs(dy)
        elif cy + dy > bottom - 12: dy = -abs(dy)
        draw.text((cx + dx, cy + dy), str(i + 1), font=f_num, fill="#000080", anchor="lm" if dx > 0 else "rm")
    ax, ay = sx(avg_cap), sy(avg_cons)
    label = "Moyenne générale"
    label_w = docx_text_w(draw, label, f_num)
    best = None
    for ddx, ddy, anchor in ((20, -16, "l"), (20, 22, "l"), (-20, -16, "r"), (-20, 22, "r")):
        lx, ly = ax + ddx, ay + ddy
        box_left = lx if anchor == "l" else lx - label_w
        box_right = box_left + label_w
        box_top, box_bottom = ly - 10, ly + 10
        in_bounds = box_left >= left + 2 and box_right <= right - 2 and box_top >= top + 2 and box_bottom <= bottom - 2
        min_dist = min(
            math.hypot(sx(d["capacity"]) - max(box_left, min(sx(d["capacity"]), box_right)),
                       sy(d["consensus"]) - max(box_top, min(sy(d["consensus"]), box_bottom)))
            for d in data
        )
        score = (in_bounds, min_dist)
        if best is None or score > best[0]:
            best = (score, lx, ly, anchor)
    _, lx, ly, anchor = best
    r = 7
    draw.polygon([(ax, ay - r), (ax + r, ay), (ax, ay + r), (ax - r, ay)], fill="white", outline="black", width=2)
    draw.text((lx, ly), label, font=f_num, fill="black", anchor="lm" if anchor == "l" else "rm")
    legend = " · ".join(f"{i + 1} {d['label']}" for i, d in enumerate(data))
    words, lines, cur = legend.split(" "), [], ""
    for word in words:
        trial = (cur + " " + word).strip()
        if not cur or docx_text_w(draw, trial, f_legend) <= w - 60:
            cur = trial
        else:
            lines.append(cur); cur = word
    if cur:
        lines.append(cur)
    ly2 = bottom + 74
    for line in lines:
        draw.text((30, ly2), line, font=f_legend, fill="#444444"); ly2 += 20
    return img


def docx_cohort_chart(items):
    """Standardised scores (bars) and graded scores (lines) across domains + the
    cohort mean as a 7th column, raster twin of the redesigned cohortChart()."""
    data = [d for d in items if d.get("capacity") is not None]
    w = max(900, 90 + (len(data) + 1) * 140) if data else 900
    h = 560
    img = PILImage.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    f_title, f_axis, f_legend, f_val, f_cat, f_cat_b, f_xtitle = docx_font(20, True), docx_font(15), docx_font(14), docx_font(14), docx_font(14), docx_font(14, True), docx_font(17, True)
    draw.text((w / 2, 24), "Notes standardisées et graduées — cohorte des domaines", font=f_title, fill="black", anchor="mm")
    if not data:
        draw.text((30, 60), "Pas encore de données disponibles pour l’analyse de cohorte.", font=f_axis, fill="black")
        return img
    mc = sum(d["capacity"] for d in data) / len(data)
    cons_data = [d for d in data if d.get("consensus") is not None]
    ms = sum(d["consensus"] for d in cons_data) / len(cons_data) if cons_data else None
    # Mirrors static/app.js cohort(): missing graded values count as 0, divided by
    # the full domain count (not just the domains that have a graded value).
    gc = sum(d.get("gradedCapacity") or 0 for d in data) / len(data)
    gs = sum(d.get("gradedConsensus") or 0 for d in data) / len(data)
    all_data = data + [{"label": "Moyen", "code": "Moyen", "capacity": mc, "consensus": ms, "gradedCapacity": gc, "gradedConsensus": gs}]
    left, right = 100, w - 30
    top, bottom = 108, h - 120
    plot_w, plot_h = right - left, bottom - top
    axis_min, axis_max = 20, 100
    scaleY = lambda v: bottom - (v - axis_min) / (axis_max - axis_min) * plot_h
    draw.rectangle([left, top, right, bottom], fill="#e9e9e9", outline="black")
    for v in (20, 40, 60, 80, 100):
        yy = scaleY(v)
        draw.line([left, yy, right, yy], fill="#bbbbbb")
        draw.text((left - 12, yy), str(v), font=f_axis, fill="black", anchor="rm")
    draw.rectangle([left, 56, left + 16, 72], fill="#9999FF")
    draw.text((left + 24, 64), "Capacité standardisée", font=f_legend, fill="black", anchor="lm")
    draw.rectangle([left + 240, 56, left + 256, 72], fill="#993366")
    draw.text((left + 264, 64), "Consensus standardisé", font=f_legend, fill="black", anchor="lm")
    draw.line([left, 86, left + 32, 86], fill="#800000", width=3)
    draw.text((left + 38, 86), "Capacité graduée", font=f_legend, fill="black", anchor="lm")
    draw.line([left + 220, 86, left + 252, 86], fill="#0000ff", width=3)
    draw.text((left + 258, 86), "Consensus gradué", font=f_legend, fill="black", anchor="lm")
    n = len(all_data); step = plot_w / n; bw = step * .32
    cap_pts, cons_pts = [], []
    for i, d in enumerate(all_data):
        x = left + i * step
        c_, s_ = d["capacity"], d["consensus"]
        cv, sv = max(axis_min, c_ or 0), max(axis_min, s_ or 0)
        y_cap, y_cons = scaleY(cv), scaleY(sv)
        draw.rectangle([x + step * .12, y_cap, x + step * .12 + bw, bottom], fill="#9999FF", outline="black")
        draw.rectangle([x + step * .12 + bw + 4, y_cons, x + step * .12 + bw + 4 + bw, bottom], fill="#993366", outline="black")
        lx = x + step * .12 + bw + 2
        docx_rotated_text(img, pdf_short_label(d), lx, bottom + 8, -40, f_cat_b if d["label"] == "Moyen" else f_cat, "black")
        if d.get("gradedCapacity") is not None:
            cap_pts.append((lx, scaleY(d["gradedCapacity"])))
        if d.get("gradedConsensus") is not None:
            cons_pts.append((lx, scaleY(d["gradedConsensus"])))
    if len(cap_pts) > 1:
        draw.line(cap_pts, fill="#800000", width=3)
    for px, py in cap_pts:
        draw.polygon([(px, py - 6), (px + 6, py + 5), (px - 6, py + 5)], fill="#800000")
    if len(cons_pts) > 1:
        draw.line(cons_pts, fill="#0000ff", width=3)
    for px, py in cons_pts:
        draw.line([px - 5, py - 5, px + 5, py + 5], fill="#0000ff", width=2)
        draw.line([px - 5, py + 5, px + 5, py - 5], fill="#0000ff", width=2)
    draw.text((left + plot_w / 2, h - 34), "Domaines", font=f_xtitle, fill="black", anchor="mm")
    return img


def docx_radar_chart(items):
    """Complementary synthesis radar, raster twin of radar() / pdf_radar_chart()."""
    data = [d for d in items if d.get("capacity") is not None]
    w, h = 560, 580
    img = PILImage.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    f_legend, f_label = docx_font(15), docx_font(13)
    if len(data) < 3:
        draw.text((30, 30), "Radar non disponible : au moins 3 domaines avec des données sont nécessaires.", font=f_legend, fill="black")
        return img
    n = len(data)
    cx, cy, r = w / 2, h / 2 + 16, 190
    draw.rectangle([20, 20, 34, 34], fill="#176b4b")
    draw.text((42, 27), "Capacité", font=f_legend, fill="black", anchor="lm")
    draw.rectangle([170, 20, 184, 34], fill="#536271")
    draw.text((192, 27), "Consensus", font=f_legend, fill="black", anchor="lm")
    pts_cap, pts_cons = [], []
    for i, d in enumerate(data):
        ang = -math.pi / 2 + i * 2 * math.pi / n
        ex, ey = cx + math.cos(ang) * r, cy + math.sin(ang) * r
        draw.line([cx, cy, ex, ey], fill="#aaaabb")
        lx, ly = cx + math.cos(ang) * (r + 20), cy + math.sin(ang) * (r + 20)
        draw.text((lx, ly), pdf_short_label(d)[:9], font=f_label, fill="black", anchor="mm")
        cv, sv = (d.get("capacity") or 0) / 100, (d.get("consensus") or 0) / 100
        pts_cap.append((cx + math.cos(ang) * r * cv, cy + math.sin(ang) * r * cv))
        pts_cons.append((cx + math.cos(ang) * r * sv, cy + math.sin(ang) * r * sv))
    draw.polygon(pts_cap, fill="#cfe3da", outline="#176b4b")
    draw.polygon(pts_cons, outline="#536271", width=2)
    return img


def docx_add_chart(doc, img):
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(6.3))


DOCX_NAVY = RGBColor(0x1B, 0x3A, 0x5C) if RGBColor else None
DOCX_MUTED = RGBColor(0x64, 0x74, 0x8B) if RGBColor else None


def docx_shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def docx_style_heading(heading, color=None):
    for run in heading.runs:
        run.font.color.rgb = color or DOCX_NAVY
    return heading


def docx_note(doc, text, size=9, color=None, italic=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.color.rgb = color or DOCX_MUTED
    return p


def docx_style_table(table, header_bg="1B3A5C"):
    table.style = "Table Grid"
    for cell in table.rows[0].cells:
        docx_shade_cell(cell, header_bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(table.rows[1:]):
        if i % 2 == 1:
            for cell in row.cells:
                docx_shade_cell(cell, "F2F5F8")
    return table


def docx_metrics_row(doc, metrics):
    """metrics: list of (label, value) tuples, rendered as a row of KPI cards like the web app."""
    t = doc.add_table(rows=2, cols=len(metrics))
    t.autofit = True
    for col, (label, value) in enumerate(metrics):
        vcell = t.cell(0, col)
        vcell.text = str(value)
        vp = vcell.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.runs[0]; vr.font.bold = True; vr.font.size = Pt(18); vr.font.color.rgb = DOCX_NAVY
        lcell = t.cell(1, col)
        lcell.text = label
        lp = lcell.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.runs[0]; lr.font.size = Pt(9); lr.font.color.rgb = DOCX_MUTED
        docx_shade_cell(vcell, "EAF1F9"); docx_shade_cell(lcell, "EAF1F9")
    return t


def pdf_level(v):
    """Mirrors static/app.js level(): qualitative reading of a capacity score."""
    if v is None: return "—"
    if v < 20: return ""
    if v <= 39: return "Loin en dessous de la moyenne"
    if v <= 59: return "En dessous de la moyenne"
    if v <= 70: return "Moyen"
    if v <= 80: return "Au-dessus de la moyenne"
    return "Bien au-dessus de la moyenne"


def pil_chart_grid(images, cols=2, pad=24, bg="white"):
    """Compose several chart images (as returned by the docx_*_chart helpers) into one
    dashboard image, so a report page shows several charts side by side instead of
    one chart per page. Each entry is either a PIL image, or an (image, caption) tuple
    whose caption is drawn above the chart (used for per-domain stat lines)."""
    items = [x if isinstance(x, tuple) else (x, None) for x in images if x is not None]
    if not items:
        return None
    cell_w = max(img.width for img, _ in items)
    cell_h = max(img.height for img, _ in items)
    caption_h = 34 if any(cap for _, cap in items) else 0
    rows = math.ceil(len(items) / cols)
    grid_w = cols * cell_w + pad * (cols + 1)
    grid_h = rows * (cell_h + caption_h) + pad * (rows + 1)
    canvas_img = PILImage.new("RGB", (grid_w, grid_h), bg)
    draw = ImageDraw.Draw(canvas_img)
    font = docx_font(20, True)
    for idx, (img, caption) in enumerate(items):
        r, col = divmod(idx, cols)
        x0 = pad + col * (cell_w + pad)
        y0 = pad + r * (cell_h + caption_h + pad)
        if caption:
            draw.text((x0 + cell_w / 2, y0 + caption_h / 2), caption, font=font, fill="black", anchor="mm")
        x = x0 + (cell_w - img.width) // 2
        y = y0 + caption_h + (cell_h - img.height) // 2
        canvas_img.paste(img, (x, y))
    return canvas_img


def xlsx_add_image(ws, row, col, img, scale=0.55):
    """Embed a PIL image in an xlsxwriter sheet at (row, col); returns the next free row."""
    if img is None:
        return row
    buf = BytesIO(); img.save(buf, format="PNG"); buf.seek(0)
    ws.insert_image(row, col, "chart.png", {"image_data": buf, "x_scale": scale, "y_scale": scale})
    return row + math.ceil(img.height * scale / 15) + 2


def report_docx(db, sid):
    """Mirrors the web app's 'Diagnostic terminé' page (the one 'Imprimer / Télécharger en
    PDF' prints) section for section: title/meta, KPI cards, Vue synthétique, Synthèse par
    domaine, Priorités retenues — same order, same headings, same table columns."""
    a = analysis(db, sid)
    if not a:
        raise ValueError("Session introuvable")
    session, g = a["session"], a["global"]
    manifest = restitution_manifest(template_payload(db, session["template_id"]))
    domains = [d for d in a["domains"] if d.get("capacity") is not None]
    priorities = qualitative_data(db, sid)["priorities"]
    doc = Document()

    docx_style_heading(doc.add_heading("Diagnostic terminé", level=0))
    docx_note(doc, f"{session['name']} · {a['participantCount']} participants · {a['completedCount']} validés · taux {round(a['completedCount']/a['participantCount']*100) if a['participantCount'] else 0}%", size=11, italic=False)
    doc.add_paragraph()
    docx_metrics_row(doc, [
        ("Capacité", pdf_fmt(g["capacity"])),
        ("Consensus", pdf_c(g)),
        ("Capacité graduée", pdf_fmt(g["gradedCapacity"])),
        ("Consensus gradué", pdf_fmt(g["gradedConsensus"])),
    ])

    profile_breakdown = participant_profile_breakdown(db, sid)
    if profile_breakdown:
        docx_style_heading(doc.add_heading("Profil des participants", level=2))
        pt = doc.add_table(rows=1, cols=3)
        for c, x in zip(pt.rows[0].cells, ["Champ", "Valeur", "N"]):
            c.text = x
        for field, value, n in profile_breakdown:
            row = pt.add_row().cells
            for c, v in zip(row, [field, value, n]):
                c.text = str(v)
        docx_style_table(pt)

    def add_findings_section():
        rows_ = findings_rows(a["findings"])
        if not rows_: return
        docx_style_heading(doc.add_heading("Forces / Fragilités / Points de vigilance", level=2))
        docx_note(doc, "Constats chiffrés issus des données ci-dessus — jamais présentés comme des causes.", italic=True)
        ft = doc.add_table(rows=1, cols=5)
        for c, x in zip(ft.rows[0].cells, ["Type", "Niveau", "Libellé", "Capacité", "Consensus"]):
            c.text = x
        for row_data in rows_:
            row = ft.add_row().cells
            for c, v in zip(row, row_data):
                c.text = "" if v is None else str(v)
        docx_style_table(ft)

    if not domains or PILImage is None:
        docx_style_heading(doc.add_heading("Synthèse par domaine", level=2))
        t = doc.add_table(rows=1, cols=6)
        for c, x in zip(t.rows[0].cells, ["Domaine", "Capacité", "Consensus", "Graduées", "Niveau", "Réponses"]):
            c.text = x
        for d in domains:
            row = t.add_row().cells
            graded = f"{d['gradedCapacity'] if d['gradedCapacity'] is not None else '—'} / {d['gradedConsensus'] if d['gradedConsensus'] is not None else '—'}"
            for c, v in zip(row, [d["label"], pdf_fmt(d["capacity"]), pdf_c(d), graded, pdf_level(d["capacity"]), d["responses"]]):
                c.text = str(v)
        docx_style_table(t)
        if not domains:
            doc.add_paragraph("Pas encore de données suffisantes pour la restitution graphique EPC.")
        elif PILImage is None:
            doc.add_paragraph("Graphiques indisponibles : le paquet optionnel Pillow n'est pas installé (voir README.md).")
        add_findings_section()
        docx_style_heading(doc.add_heading("Priorités retenues", level=2))
        doc.add_paragraph(f"{len(priorities)} priorité(s) sélectionnée(s)." if priorities else "Aucune priorité sélectionnée.")
        out = BytesIO(); doc.save(out); return out.getvalue()

    doc.add_page_break()
    docx_style_heading(doc.add_heading("Vue synthétique", level=2))
    docx_note(doc, "Radar global : vue complémentaire, ne remplace pas les graphiques EPC ci-dessous.")
    docx_add_chart(doc, pil_chart_grid([
        docx_radar_chart(domains),
        docx_bars_chart(domains, "standard", "Notes standardisées par domaine", with_mean=True),
        docx_bars_chart(domains, "graded", "Notes graduées par domaine", with_mean=True),
        docx_grid_chart(domains),
    ]))
    if len(domains) > 1:
        docx_add_chart(doc, docx_cohort_chart(domains))
    docx_note(doc, "20–39 : Loin en dessous · 40–59 : En dessous · 60–70 : Moyen · 71–80 : Au-dessus · 81–100 : Bien au-dessus", size=9, color=RGBColor(0x18, 0x6E, 0x42))

    doc.add_page_break()
    docx_style_heading(doc.add_heading("Synthèse par domaine", level=2))
    t = doc.add_table(rows=1, cols=6)
    for c, x in zip(t.rows[0].cells, ["Domaine", "Capacité", "Consensus", "Graduées", "Niveau", "Réponses"]):
        c.text = x
    for d in domains:
        row = t.add_row().cells
        graded = f"{d['gradedCapacity'] if d['gradedCapacity'] is not None else '—'} / {d['gradedConsensus'] if d['gradedConsensus'] is not None else '—'}"
        for c, v in zip(row, [d["label"], pdf_fmt(d["capacity"]), pdf_c(d), graded, pdf_level(d["capacity"]), d["responses"]]):
            c.text = str(v)
    docx_style_table(t)

    add_findings_section()
    docx_style_heading(doc.add_heading("Priorités retenues", level=2))
    doc.add_paragraph(f"{len(priorities)} priorité(s) sélectionnée(s)." if priorities else "Aucune priorité sélectionnée.")

    ai_blocks = rows(db, "SELECT section_key,content FROM report_ai_blocks WHERE session_id=?", (sid,))
    if ai_blocks:
        doc.add_page_break()
        docx_style_heading(doc.add_heading("Synthèse assistée par IA", level=2))
        docx_note(doc, "Propositions rédigées avec l'aide de l'assistant IA et explicitement retenues par le modérateur. "
            "Les données, scores et graphiques EPC ci-dessus restent la source primaire du diagnostic.", italic=False)
        for key in manifest["aiSectionLabels"]:
            block = next((b for b in ai_blocks if b["section_key"] == key), None)
            if not block: continue
            docx_style_heading(doc.add_heading(manifest["aiSectionLabels"][key], level=3))
            doc.add_paragraph(block["content"])

    out = BytesIO(); doc.save(out); return out.getvalue()

PDF_STOPWORDS = {"de", "des", "du", "la", "le", "les", "et", "en", "au", "aux", "à", "a", "l"}


def pdf_fmt(v):
    return "—" if v is None else f"{v:.1f}"


def pdf_c(obj):
    return "Non calculable" if obj.get("consensusNote") == "single_respondent" else pdf_fmt(obj.get("consensus"))


def pdf_short_label(item):
    """Mirrors static/app.js shortLabel(): configured code if it reads as a short
    alphabetic label, otherwise an abbreviation generated from the item's own name."""
    code = (item.get("code") or "").strip()
    if re.fullmatch(r"[A-Za-zÀ-ÿ]{2,10}", code):
        return code.upper()
    label = (item.get("label") or "").strip()
    words = [w for w in label.split() if w and w.lower().replace("’", "").replace("'", "") not in PDF_STOPWORDS]
    if not words:
        words = label.split()
    if len(words) >= 2:
        return "".join(w[0].upper() for w in words[:4])
    return label[:4].upper()


def qualitative_data(db, session_id: str):
    """Qualitative workshop data kept separate from EPC numerical results."""
    priorities = rows(db, """SELECT p.*,d.label AS domain_label,i.code AS indicator_code,
        i.label AS indicator_label,i.description AS indicator_description
        FROM priorities p JOIN domains d ON d.id=p.domain_id JOIN indicators i ON i.id=p.indicator_id
        WHERE p.session_id=? ORDER BY d.display_order,i.display_order""", (session_id,))
    return {
        "priorities": priorities,
        "analyses": rows(db, "SELECT * FROM priority_analyses WHERE session_id=?", (session_id,)),
        "entries": rows(db, "SELECT * FROM analysis_entries WHERE session_id=? ORDER BY created_at", (session_id,)),
        "recommendations": rows(db, "SELECT * FROM workshop_recommendations WHERE session_id=? ORDER BY created_at", (session_id,)),
        "trainingTopics": rows(db, "SELECT * FROM training_topics WHERE session_id=? ORDER BY created_at", (session_id,)),
    }

def report_data(db, session_id: str):
    a=analysis(db,session_id)
    if not a: return None
    meta=db.execute("SELECT * FROM session_report_meta WHERE session_id=?",(session_id,)).fetchone()
    template=template_payload(db,a["session"]["template_id"])
    return {"analysis":a,"template":template,"qualitative":qualitative_data(db,session_id),"meta":dict(meta) if meta else {"facilitator":"","audience":"","context":"","conclusion":""},"manifest":restitution_manifest(template),"profile":participant_profile_breakdown(db,session_id)}
